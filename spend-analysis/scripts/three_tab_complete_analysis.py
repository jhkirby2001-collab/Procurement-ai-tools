#!/usr/bin/env python3
"""
Complete Analysis of All Three Tabs
- PO14954: Amount Released Only
- PO30043: Amount Released Only
- Exhibit B: Full Comprehensive Analysis
"""

import pandas as pd
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows

# File paths
DATA_FILE = "/workspaces/Spend_Analysis_and_Categorization/spend-analysis/data/raw/correctedBottledWater raw data.xlsx"
OUTPUT_DIR = "/workspaces/Spend_Analysis_and_Categorization/spend-analysis/outputs/"
WORD_FILE = OUTPUT_DIR + "ALL_THREE_TABS_COMPLETE_ANALYSIS.docx"
EXCEL_FILE = OUTPUT_DIR + "ALL_THREE_TABS_COMPLETE_ANALYSIS.xlsx"

def clean_currency(series):
    """Convert currency strings to numeric values"""
    if series.dtype == 'object':
        cleaned = series.astype(str).str.replace('$', '').str.replace(',', '').str.strip()
        return pd.to_numeric(cleaned, errors='coerce')
    return pd.to_numeric(series, errors='coerce')

def format_currency(value):
    """Format number as currency"""
    if pd.isna(value):
        return "$0.00"
    return f"${value:,.2f}"

print("=" * 80)
print("READING DATA FROM ALL THREE TABS")
print("=" * 80)

# Read all three sheets
df_po1 = pd.read_excel(DATA_FILE, sheet_name='SPC49756A-PO14954')
df_po2 = pd.read_excel(DATA_FILE, sheet_name='SPC82197E-PO30043')
df_exhibitb = pd.read_excel(DATA_FILE, sheet_name='exhibit b ')

print(f"✓ PO14954: {len(df_po1)} rows")
print(f"✓ PO30043: {len(df_po2)} rows")
print(f"✓ Exhibit B: {len(df_exhibitb)} rows")

# ============================================================================
# PART 1: ANALYZE AMOUNT RELEASED FROM PO TABS
# ============================================================================

print("\n" + "=" * 80)
print("PART 1: ANALYZING AMOUNT RELEASED FROM PO TABS")
print("=" * 80)

# Clean Amount Released for PO1
if 'Amount Released' in df_po1.columns:
    df_po1['Amount Released Clean'] = clean_currency(df_po1['Amount Released'])
elif 'Contract Limit' in df_po1.columns:
    df_po1['Amount Released Clean'] = clean_currency(df_po1['Contract Limit'])
else:
    df_po1['Amount Released Clean'] = 0

# Clean Amount Released for PO2
if 'Amount Released' in df_po2.columns:
    df_po2['Amount Released Clean'] = clean_currency(df_po2['Amount Released'])
elif 'Contract Limit' in df_po2.columns:
    df_po2['Amount Released Clean'] = clean_currency(df_po2['Contract Limit'])
else:
    df_po2['Amount Released Clean'] = 0

# Calculate totals
po1_total = df_po1['Amount Released Clean'].sum()
po2_total = df_po2['Amount Released Clean'].sum()
po_grand_total = po1_total + po2_total

po1_items = len(df_po1[df_po1['Amount Released Clean'] > 0])
po2_items = len(df_po2[df_po2['Amount Released Clean'] > 0])
po_total_items = po1_items + po2_items

print(f"\nPO14954 Amount Released: {format_currency(po1_total)} ({po1_items} items)")
print(f"PO30043 Amount Released: {format_currency(po2_total)} ({po2_items} items)")
print(f"TOTAL Amount Released: {format_currency(po_grand_total)} ({po_total_items} items)")

# ============================================================================
# PART 2: FULL ANALYSIS OF EXHIBIT B
# ============================================================================

print("\n" + "=" * 80)
print("PART 2: FULL ANALYSIS OF EXHIBIT B")
print("=" * 80)

# Clean requisition amounts
df_exhibitb['Requisition Amount Clean'] = pd.to_numeric(df_exhibitb['Requisition Amount'], errors='coerce').fillna(0)

# Total requisition amount
exhibitb_total = df_exhibitb['Requisition Amount Clean'].sum()
exhibitb_count = len(df_exhibitb)
exhibitb_avg = df_exhibitb['Requisition Amount Clean'].mean()

print(f"\nTotal Requisition Amount: {format_currency(exhibitb_total)}")
print(f"Number of Requisitions: {exhibitb_count}")
print(f"Average Requisition: {format_currency(exhibitb_avg)}")

# Analysis by Vendor
vendor_analysis = df_exhibitb.groupby('Vendor Name').agg({
    'Requisition Amount Clean': ['sum', 'count', 'mean']
}).round(2)
vendor_analysis.columns = ['Total Amount', 'Count', 'Avg Amount']
vendor_analysis = vendor_analysis.sort_values('Total Amount', ascending=False)

print(f"\nTop 5 Vendors by Spend:")
for idx, (vendor, row) in enumerate(vendor_analysis.head().iterrows(), 1):
    print(f"  {idx}. {vendor}: {format_currency(row['Total Amount'])} ({int(row['Count'])} requisitions)")

# Analysis by Department
dept_analysis = df_exhibitb.groupby('Department Description').agg({
    'Requisition Amount Clean': ['sum', 'count', 'mean']
}).round(2)
dept_analysis.columns = ['Total Amount', 'Count', 'Avg Amount']
dept_analysis = dept_analysis.sort_values('Total Amount', ascending=False)

print(f"\nTop 5 Departments by Spend:")
for idx, (dept, row) in enumerate(dept_analysis.head().iterrows(), 1):
    print(f"  {idx}. {dept}: {format_currency(row['Total Amount'])} ({int(row['Count'])} requisitions)")

# Analysis by Procurement Phase
phase_analysis = df_exhibitb.groupby('Procurement Phase').agg({
    'Requisition Amount Clean': ['sum', 'count']
}).round(2)
phase_analysis.columns = ['Total Amount', 'Count']
phase_analysis = phase_analysis.sort_values('Total Amount', ascending=False)

print(f"\nRequisitions by Procurement Phase:")
for phase, row in phase_analysis.iterrows():
    print(f"  {phase}: {format_currency(row['Total Amount'])} ({int(row['Count'])} requisitions)")

# Analysis by Approval Status
approval_counts = df_exhibitb['Approved or Denied'].value_counts()
print(f"\nApproval Status:")
for status, count in approval_counts.items():
    print(f"  {status}: {count}")

# New Contract Requests
new_contract_counts = df_exhibitb['Request for New Contract/Mod'].value_counts()
print(f"\nNew Contract/Modification Requests:")
for status, count in new_contract_counts.items():
    print(f"  {status}: {count}")

# Timeline Analysis
df_exhibitb['Year'] = pd.to_datetime(df_exhibitb['Date Received-DPS']).dt.year
year_analysis = df_exhibitb.groupby('Year').agg({
    'Requisition Amount Clean': ['sum', 'count']
}).round(2)
year_analysis.columns = ['Total Amount', 'Count']
year_analysis = year_analysis.sort_index()

print(f"\nRequisitions by Year:")
for year, row in year_analysis.iterrows():
    print(f"  {int(year)}: {format_currency(row['Total Amount'])} ({int(row['Count'])} requisitions)")

# ============================================================================
# CREATE WORD DOCUMENT
# ============================================================================

print("\n" + "=" * 80)
print("CREATING MICROSOFT WORD DOCUMENT")
print("=" * 80)

doc = Document()

# Title
title = doc.add_heading('Complete Bottled Water Analysis', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Analysis of All Three Tabs: PO14954, PO30043, and Exhibit B')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(12)
subtitle_format.font.italic = True

# Date
date_para = doc.add_paragraph(f'Report Date: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# SECTION 1: EXECUTIVE SUMMARY
doc.add_heading('1. Executive Summary', 1)

doc.add_paragraph(
    f"This report provides a comprehensive analysis of bottled water procurement across "
    f"three data sources: two purchase orders (PO14954 and PO30043) showing amounts released, "
    f"and Exhibit B showing 214 requisitions with detailed procurement workflow data."
)

doc.add_heading('1.1 Amount Released from Purchase Orders', 2)

table1 = doc.add_table(rows=4, cols=3)
table1.style = 'Light Grid Accent 1'

headers1 = ['Contract', 'PO Number', 'Amount Released']
for i, header in enumerate(headers1):
    cell = table1.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].font.bold = True

table1.rows[1].cells[0].text = 'Contract 1'
table1.rows[1].cells[1].text = 'PO14954'
table1.rows[1].cells[2].text = format_currency(po1_total)

table1.rows[2].cells[0].text = 'Contract 2'
table1.rows[2].cells[1].text = 'PO30043'
table1.rows[2].cells[2].text = format_currency(po2_total)

table1.rows[3].cells[0].text = 'TOTAL'
table1.rows[3].cells[1].text = 'Both'
table1.rows[3].cells[2].text = format_currency(po_grand_total)
for cell in table1.rows[3].cells:
    cell.paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

doc.add_heading('1.2 Exhibit B Requisition Summary', 2)

table2 = doc.add_table(rows=4, cols=2)
table2.style = 'Light Grid Accent 1'

headers2 = ['Metric', 'Value']
for i, header in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].font.bold = True

table2.rows[1].cells[0].text = 'Total Requisition Amount'
table2.rows[1].cells[1].text = format_currency(exhibitb_total)

table2.rows[2].cells[0].text = 'Number of Requisitions'
table2.rows[2].cells[1].text = str(exhibitb_count)

table2.rows[3].cells[0].text = 'Average Requisition'
table2.rows[3].cells[1].text = format_currency(exhibitb_avg)

doc.add_paragraph()

doc.add_heading('1.3 Combined Total Spend', 2)
combined_total = po_grand_total + exhibitb_total
doc.add_paragraph(
    f"Total Amount Released (POs): {format_currency(po_grand_total)}\n"
    f"Total Requisition Amount (Exhibit B): {format_currency(exhibitb_total)}\n"
    f"COMBINED TOTAL: {format_currency(combined_total)}"
).runs[0].font.bold = True

doc.add_page_break()

# SECTION 2: EXHIBIT B DETAILED ANALYSIS
doc.add_heading('2. Exhibit B Detailed Analysis', 1)

doc.add_heading('2.1 Spending by Vendor', 2)
table3 = doc.add_table(rows=len(vendor_analysis.head(10))+1, cols=4)
table3.style = 'Light Grid Accent 1'

headers3 = ['Vendor Name', 'Total Amount', 'Requisitions', 'Avg Amount']
for i, header in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].font.bold = True

for idx, (vendor, row) in enumerate(vendor_analysis.head(10).iterrows(), 1):
    table3.rows[idx].cells[0].text = vendor
    table3.rows[idx].cells[1].text = format_currency(row['Total Amount'])
    table3.rows[idx].cells[2].text = str(int(row['Count']))
    table3.rows[idx].cells[3].text = format_currency(row['Avg Amount'])

doc.add_paragraph()

doc.add_heading('2.2 Spending by Department', 2)
table4 = doc.add_table(rows=len(dept_analysis.head(10))+1, cols=4)
table4.style = 'Light Grid Accent 1'

headers4 = ['Department', 'Total Amount', 'Requisitions', 'Avg Amount']
for i, header in enumerate(headers4):
    cell = table4.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].font.bold = True

for idx, (dept, row) in enumerate(dept_analysis.head(10).iterrows(), 1):
    table4.rows[idx].cells[0].text = dept
    table4.rows[idx].cells[1].text = format_currency(row['Total Amount'])
    table4.rows[idx].cells[2].text = str(int(row['Count']))
    table4.rows[idx].cells[3].text = format_currency(row['Avg Amount'])

doc.add_paragraph()

doc.add_heading('2.3 Procurement Phase Analysis', 2)
table5 = doc.add_table(rows=len(phase_analysis)+1, cols=3)
table5.style = 'Light Grid Accent 1'

headers5 = ['Procurement Phase', 'Total Amount', 'Requisitions']
for i, header in enumerate(headers5):
    cell = table5.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].font.bold = True

for idx, (phase, row) in enumerate(phase_analysis.iterrows(), 1):
    table5.rows[idx].cells[0].text = phase
    table5.rows[idx].cells[1].text = format_currency(row['Total Amount'])
    table5.rows[idx].cells[2].text = str(int(row['Count']))

doc.add_paragraph()

doc.add_heading('2.4 Timeline Analysis', 2)
table6 = doc.add_table(rows=len(year_analysis)+1, cols=3)
table6.style = 'Light Grid Accent 1'

headers6 = ['Year', 'Total Amount', 'Requisitions']
for i, header in enumerate(headers6):
    cell = table6.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].font.bold = True

for idx, (year, row) in enumerate(year_analysis.iterrows(), 1):
    table6.rows[idx].cells[0].text = str(int(year))
    table6.rows[idx].cells[1].text = format_currency(row['Total Amount'])
    table6.rows[idx].cells[2].text = str(int(row['Count']))

doc.add_page_break()

# SECTION 3: KEY FINDINGS AND RECOMMENDATIONS
doc.add_heading('3. Key Findings', 1)

doc.add_heading('3.1 PO Contract Findings', 2)
if po2_total > po1_total:
    pct_increase = ((po2_total - po1_total) / po1_total) * 100
    doc.add_paragraph(
        f"• Contract 2 (PO30043) is {pct_increase:.0f}% larger than Contract 1 (PO14954), "
        f"representing a ${po2_total - po1_total:,.2f} increase."
    )
doc.add_paragraph(
    f"• Total amount released across both contracts: {format_currency(po_grand_total)}"
)

doc.add_heading('3.2 Exhibit B Findings', 2)
doc.add_paragraph(f"• {exhibitb_count} requisitions totaling {format_currency(exhibitb_total)}")
doc.add_paragraph(f"• {len(vendor_analysis)} unique vendors")
doc.add_paragraph(f"• {len(dept_analysis)} departments making requisitions")

top_vendor = vendor_analysis.index[0]
top_vendor_pct = (vendor_analysis.iloc[0]['Total Amount'] / exhibitb_total) * 100
doc.add_paragraph(
    f"• Top vendor ({top_vendor}) represents {top_vendor_pct:.1f}% of total requisition spend"
)

approved_count = approval_counts.get('APPROVED', 0)
approval_rate = (approved_count / exhibitb_count) * 100 if exhibitb_count > 0 else 0
doc.add_paragraph(f"• Approval rate: {approval_rate:.1f}% ({approved_count} of {exhibitb_count})")

doc.add_heading('3.3 Combined Analysis', 2)
doc.add_paragraph(
    f"• Combined total spending (POs + Requisitions): {format_currency(combined_total)}"
)
doc.add_paragraph(
    f"• PO contracts represent {(po_grand_total/combined_total)*100:.1f}% of total spend"
)
doc.add_paragraph(
    f"• Exhibit B requisitions represent {(exhibitb_total/combined_total)*100:.1f}% of total spend"
)

doc.add_page_break()

# SECTION 4: RECOMMENDATIONS
doc.add_heading('4. Recommendations', 1)

doc.add_heading('4.1 Contract Management', 2)
doc.add_paragraph('• Investigate the significant growth from Contract 1 to Contract 2')
doc.add_paragraph('• Review pricing structure to ensure competitive rates')
doc.add_paragraph('• Consider consolidating vendors to improve pricing leverage')

doc.add_heading('4.2 Requisition Process', 2)
doc.add_paragraph('• Most requisitions are in Specification Development Phase - monitor progression')
doc.add_paragraph('• Review high-value requisitions for cost optimization opportunities')
doc.add_paragraph('• Standardize requisition amounts where possible to improve efficiency')

doc.add_heading('4.3 Vendor Management', 2)
doc.add_paragraph('• High vendor concentration - consider competitive bidding to reduce costs')
doc.add_paragraph('• Establish vendor performance metrics')
doc.add_paragraph('• Negotiate volume discounts based on combined spend')

# Save Word document
doc.save(WORD_FILE)
print(f"✓ Word document saved: {WORD_FILE}")

# ============================================================================
# CREATE EXCEL WORKBOOK
# ============================================================================

print("\n" + "=" * 80)
print("CREATING MICROSOFT EXCEL WORKBOOK")
print("=" * 80)

wb = Workbook()
wb.remove(wb.active)  # Remove default sheet

# Define styles
header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
header_font = Font(bold=True, color="FFFFFF", size=11)
total_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
total_font = Font(bold=True, size=11)
border = Border(
    left=Side(style='thin'),
    right=Side(style='thin'),
    top=Side(style='thin'),
    bottom=Side(style='thin')
)

# Sheet 1: Executive Summary
ws1 = wb.create_sheet("Executive Summary")
ws1.append(["COMPLETE BOTTLED WATER ANALYSIS"])
ws1.append(["Report Date:", datetime.now().strftime("%B %d, %Y")])
ws1.append([])

ws1.append(["AMOUNT RELEASED FROM PURCHASE ORDERS"])
ws1.append(["Contract", "PO Number", "Amount Released", "Line Items"])
ws1.append(["Contract 1", "PO14954", po1_total, po1_items])
ws1.append(["Contract 2", "PO30043", po2_total, po2_items])
ws1.append(["TOTAL", "Both", po_grand_total, po_total_items])
ws1.append([])

ws1.append(["EXHIBIT B REQUISITION SUMMARY"])
ws1.append(["Metric", "Value"])
ws1.append(["Total Requisition Amount", exhibitb_total])
ws1.append(["Number of Requisitions", exhibitb_count])
ws1.append(["Average Requisition", exhibitb_avg])
ws1.append([])

ws1.append(["COMBINED TOTAL SPEND"])
ws1.append(["Category", "Amount"])
ws1.append(["Amount Released (POs)", po_grand_total])
ws1.append(["Requisition Amount (Exhibit B)", exhibitb_total])
ws1.append(["COMBINED TOTAL", combined_total])

# Format Sheet 1
ws1.column_dimensions['A'].width = 30
ws1.column_dimensions['B'].width = 20
ws1.column_dimensions['C'].width = 20
ws1.column_dimensions['D'].width = 15

for row in ws1['A1:D1']:
    for cell in row:
        cell.font = Font(bold=True, size=14)

# Sheet 2: PO Analysis Detail
ws2 = wb.create_sheet("PO Analysis Detail")
ws2.append(["PO NUMBER", "AMOUNT RELEASED", "LINE ITEMS"])

# Apply header style
for cell in ws2[1]:
    cell.fill = header_fill
    cell.font = header_font
    cell.alignment = Alignment(horizontal='center')

ws2.append(["PO14954", po1_total, po1_items])
ws2.append(["PO30043", po2_total, po2_items])
ws2.append(["TOTAL", po_grand_total, po_total_items])

# Apply total row style
for cell in ws2[4]:
    cell.fill = total_fill
    cell.font = total_font

ws2.column_dimensions['A'].width = 15
ws2.column_dimensions['B'].width = 20
ws2.column_dimensions['C'].width = 15

# Sheet 3: Exhibit B - Vendor Analysis
ws3 = wb.create_sheet("Exhibit B - Vendors")
ws3.append(["VENDOR NAME", "TOTAL AMOUNT", "REQUISITIONS", "AVG AMOUNT"])

for cell in ws3[1]:
    cell.fill = header_fill
    cell.font = header_font
    cell.alignment = Alignment(horizontal='center')

for vendor, row in vendor_analysis.iterrows():
    ws3.append([vendor, row['Total Amount'], int(row['Count']), row['Avg Amount']])

ws3.column_dimensions['A'].width = 40
ws3.column_dimensions['B'].width = 20
ws3.column_dimensions['C'].width = 15
ws3.column_dimensions['D'].width = 20

# Sheet 4: Exhibit B - Department Analysis
ws4 = wb.create_sheet("Exhibit B - Departments")
ws4.append(["DEPARTMENT", "TOTAL AMOUNT", "REQUISITIONS", "AVG AMOUNT"])

for cell in ws4[1]:
    cell.fill = header_fill
    cell.font = header_font
    cell.alignment = Alignment(horizontal='center')

for dept, row in dept_analysis.iterrows():
    ws4.append([dept, row['Total Amount'], int(row['Count']), row['Avg Amount']])

ws4.column_dimensions['A'].width = 50
ws4.column_dimensions['B'].width = 20
ws4.column_dimensions['C'].width = 15
ws4.column_dimensions['D'].width = 20

# Sheet 5: Exhibit B - Timeline
ws5 = wb.create_sheet("Exhibit B - Timeline")
ws5.append(["YEAR", "TOTAL AMOUNT", "REQUISITIONS"])

for cell in ws5[1]:
    cell.fill = header_fill
    cell.font = header_font
    cell.alignment = Alignment(horizontal='center')

for year, row in year_analysis.iterrows():
    ws5.append([int(year), row['Total Amount'], int(row['Count'])])

ws5.column_dimensions['A'].width = 15
ws5.column_dimensions['B'].width = 20
ws5.column_dimensions['C'].width = 15

# Sheet 6: Exhibit B - Procurement Phase
ws6 = wb.create_sheet("Exhibit B - Procurement Phase")
ws6.append(["PROCUREMENT PHASE", "TOTAL AMOUNT", "REQUISITIONS"])

for cell in ws6[1]:
    cell.fill = header_fill
    cell.font = header_font
    cell.alignment = Alignment(horizontal='center')

for phase, row in phase_analysis.iterrows():
    ws6.append([phase, row['Total Amount'], int(row['Count'])])

ws6.column_dimensions['A'].width = 45
ws6.column_dimensions['B'].width = 20
ws6.column_dimensions['C'].width = 15

# Save Excel workbook
wb.save(EXCEL_FILE)
print(f"✓ Excel workbook saved: {EXCEL_FILE}")

print("\n" + "=" * 80)
print("ANALYSIS COMPLETE!")
print("=" * 80)
print(f"\nFiles created:")
print(f"1. {WORD_FILE}")
print(f"2. {EXCEL_FILE}")
print(f"\nBoth files are ready to download.")
