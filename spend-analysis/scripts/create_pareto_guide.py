#!/usr/bin/env python3
"""
Create a Word document explaining Pareto 80/20 Analysis
and how to interpret the Excel analysis file.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_color(doc, text, level, color_rgb=(0, 51, 102)):
    """Add a colored heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
    return heading

def add_styled_paragraph(doc, text, bold=False, italic=False, size=11):
    """Add a paragraph with custom styling."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return para

def add_bullet_point(doc, text, level=0):
    """Add a bullet point."""
    para = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        para.paragraph_format.left_indent = Inches(0.5 * level)
    return para

def add_table_with_style(doc, data, headers):
    """Add a styled table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Bold header text
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

def create_pareto_guide():
    """Create the Pareto 80/20 Analysis Guide document."""

    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ========================================================================
    # TITLE PAGE
    # ========================================================================
    title = doc.add_heading('Understanding Pareto 80/20 Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('A Guide to Interpreting Your Excel Analysis')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(
        'Chicago Spending Without Active Contracts\n'
        'Pareto Analysis Report\n\n'
        'January 2026'
    )
    info_run.font.size = Pt(12)
    info_run.font.color.rgb = RGBColor(64, 64, 64)

    doc.add_page_break()

    # ========================================================================
    # TABLE OF CONTENTS
    # ========================================================================
    add_heading_with_color(doc, 'Table of Contents', 1)
    doc.add_paragraph('1. What is Pareto 80/20 Analysis?', style='List Number')
    doc.add_paragraph('2. The History and Science Behind the Pareto Principle', style='List Number')
    doc.add_paragraph('3. Why Use Pareto Analysis for Spending?', style='List Number')
    doc.add_paragraph('4. Understanding Your Excel Analysis File', style='List Number')
    doc.add_paragraph('5. How to Read Each Worksheet', style='List Number')
    doc.add_paragraph('6. Key Findings from Your Analysis', style='List Number')
    doc.add_paragraph('7. How to Use This Analysis for Decision-Making', style='List Number')
    doc.add_paragraph('8. Common Questions and Answers', style='List Number')

    doc.add_page_break()

    # ========================================================================
    # SECTION 1: What is Pareto Analysis?
    # ========================================================================
    add_heading_with_color(doc, '1. What is Pareto 80/20 Analysis?', 1)

    add_styled_paragraph(doc,
        'The Pareto Principle, also known as the 80/20 Rule, is a powerful concept that states:',
        size=12
    )

    doc.add_paragraph()

    # Add a highlighted box
    box_para = doc.add_paragraph()
    box_run = box_para.add_run(
        '"Roughly 80% of effects come from 20% of causes"'
    )
    box_run.font.size = Pt(14)
    box_run.font.bold = True
    box_run.font.color.rgb = RGBColor(0, 51, 102)
    box_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_styled_paragraph(doc,
        'In simpler terms: A small number of factors typically account for the majority of results.',
        size=12
    )

    doc.add_paragraph()
    add_heading_with_color(doc, 'Real-World Examples:', 2)

    examples = [
        ('Business Sales', '80% of revenue comes from 20% of customers'),
        ('Software Bugs', '80% of errors come from 20% of the code'),
        ('Wealth Distribution', '80% of wealth is held by 20% of the population'),
        ('Productivity', '80% of results come from 20% of your efforts'),
        ('Inventory', '80% of sales come from 20% of products'),
    ]

    for category, example in examples:
        para = doc.add_paragraph()
        para.add_run(f'{category}: ').bold = True
        para.add_run(example)

    doc.add_paragraph()
    add_styled_paragraph(doc,
        'Important Note: The numbers don\'t have to be exactly 80/20. The principle is about '
        'identifying significant imbalances where a minority of inputs create the majority of outputs.',
        italic=True,
        size=11
    )

    doc.add_page_break()

    # ========================================================================
    # SECTION 2: History
    # ========================================================================
    add_heading_with_color(doc, '2. The History and Science Behind the Pareto Principle', 1)

    add_heading_with_color(doc, 'The Origin Story', 2)

    history_text = [
        'The principle is named after Italian economist Vilfredo Pareto, who in 1896 observed that '
        '80% of Italy\'s land was owned by 20% of the population.',

        'Pareto also noticed that 20% of the pea pods in his garden produced 80% of the peas. '
        'This observation led him to investigate similar patterns in economics and society.',

        'In the 1940s, management consultant Joseph Juran recognized the universal applicability '
        'of this principle and coined the term "Pareto Principle" while applying it to quality control.'
    ]

    for text in history_text:
        add_styled_paragraph(doc, text, size=11)
        doc.add_paragraph()

    add_heading_with_color(doc, 'Why It Matters in Government Spending', 2)

    government_text = [
        'In procurement and spending analysis, the Pareto Principle helps identify WHERE to focus '
        'limited resources for maximum impact.',

        'Rather than trying to fix every vendor relationship, every transaction, or every department '
        'at once, the 80/20 analysis shows you which specific areas will give you the biggest return '
        'on your efforts.',

        'This is especially critical in government where resources are limited and accountability is paramount.'
    ]

    for text in government_text:
        add_styled_paragraph(doc, text, size=11)
        doc.add_paragraph()

    doc.add_page_break()

    # ========================================================================
    # SECTION 3: Why Use Pareto for Spending
    # ========================================================================
    add_heading_with_color(doc, '3. Why Use Pareto Analysis for Spending?', 1)

    add_heading_with_color(doc, 'The Problem: Too Much Data, Limited Resources', 2)

    add_styled_paragraph(doc,
        'When analyzing spending without contracts, you might have:',
        size=11
    )

    doc.add_paragraph()

    problems = [
        'Hundreds of vendors to review',
        'Thousands of individual transactions',
        'Multiple departments with varying compliance levels',
        'Limited procurement staff time',
        'Pressure to act quickly'
    ]

    for problem in problems:
        add_bullet_point(doc, problem)

    doc.add_paragraph()
    add_styled_paragraph(doc,
        'Question: Where do you start? How do you prioritize?',
        bold=True,
        size=12
    )

    doc.add_paragraph()
    add_heading_with_color(doc, 'The Solution: Focus on the Vital Few', 2)

    add_styled_paragraph(doc,
        'Pareto analysis answers this question by identifying:',
        size=11
    )

    doc.add_paragraph()

    solutions = [
        'Which vendors account for the most spending (so you focus contract negotiations there)',
        'Which departments have the biggest compliance issues (so you provide targeted support)',
        'Which transaction sizes represent the greatest risk (so you implement controls)',
        'Which spending categories need immediate attention (so you prioritize category management)'
    ]

    for solution in solutions:
        add_bullet_point(doc, solution)

    doc.add_paragraph()

    # Add benefit box
    benefit_para = doc.add_paragraph()
    benefit_para.add_run('Result: ').bold = True
    benefit_para.add_run(
        'Instead of spreading your efforts thin across everything, you concentrate on the '
        'critical few items that will deliver 80% of the value.'
    )

    doc.add_page_break()

    # ========================================================================
    # SECTION 4: Understanding the Excel File
    # ========================================================================
    add_heading_with_color(doc, '4. Understanding Your Excel Analysis File', 1)

    add_styled_paragraph(doc,
        'The Excel file "PARETO_80_20_ANALYSIS.xlsx" contains multiple worksheets, each analyzing '
        'a different dimension of your spending data.',
        size=11
    )

    doc.add_paragraph()
    add_heading_with_color(doc, 'File Structure Overview', 2)

    add_styled_paragraph(doc,
        'The Excel file is organized into worksheets that answer specific questions:',
        size=11
    )

    doc.add_paragraph()

    worksheets = [
        ('Summary Sheet', 'High-level overview of all Pareto analyses'),
        ('Vendors_80_20', 'Which vendors account for most spending?'),
        ('Departments_80_20', 'Which departments have the most spending?'),
        ('Transactions_80_20', 'Which transaction sizes represent the biggest risk?'),
        ('Categories_80_20', 'Which spending categories dominate?'),
        ('Raw Data', 'The underlying data used for analysis'),
    ]

    table_data = []
    for sheet, description in worksheets:
        table_data.append([sheet, description])

    add_table_with_style(doc, table_data, ['Worksheet Name', 'Purpose'])

    doc.add_paragraph()

    add_heading_with_color(doc, 'Key Metrics You\'ll See', 2)

    add_styled_paragraph(doc,
        'Each worksheet contains similar metrics to help you understand the concentration. '
        'The most important metrics are Cumulative Spend and Cumulative %, which are essential '
        'for identifying the 80% threshold.',
        size=11
    )

    doc.add_paragraph()

    # Add detailed explanation of cumulative metrics
    add_heading_with_color(doc, 'Understanding Cumulative Spend and Cumulative %', 3)

    add_styled_paragraph(doc,
        'These are the MOST CRITICAL metrics in the Pareto analysis. Let\'s understand what they mean:',
        bold=True,
        size=11
    )

    doc.add_paragraph()

    # Cumulative Spend explanation
    cumul_spend_para = doc.add_paragraph()
    cumul_spend_para.add_run('Cumulative Spend: ').bold = True
    cumul_spend_para.add_run(
        'A running total that adds up spending as you move down the list. '
        'The data is sorted from highest to lowest, so cumulative spend shows you how much '
        'total spending is accounted for by the vendors/departments/transactions UP TO that row.'
    )

    doc.add_paragraph()

    # Cumulative % explanation
    cumul_pct_para = doc.add_paragraph()
    cumul_pct_para.add_run('Cumulative %: ').bold = True
    cumul_pct_para.add_run(
        'The cumulative spend expressed as a percentage of total spend. This tells you what '
        'percentage of ALL spending is accounted for by the items UP TO that row.'
    )

    doc.add_paragraph()

    # Example section
    add_heading_with_color(doc, 'Example: How Cumulative Metrics Work', 3)

    add_styled_paragraph(doc,
        'Let\'s use a simplified vendor example to illustrate. Assume total spending across all vendors is $10,000,000.',
        size=11
    )

    doc.add_paragraph()

    # Create example table
    example_data = [
        ['Vendor A', '$3,000,000', '$3,000,000', '30.0%'],
        ['Vendor B', '$2,500,000', '$5,500,000', '55.0%'],
        ['Vendor C', '$1,500,000', '$7,000,000', '70.0%'],
        ['Vendor D', '$1,200,000', '$8,200,000', '82.0%', '← 80% THRESHOLD'],
        ['Vendor E', '$800,000', '$9,000,000', '90.0%'],
        ['...176 more', '$1,000,000', '$10,000,000', '100.0%'],
    ]

    example_table = doc.add_table(rows=1, cols=5)
    example_table.style = 'Light Grid Accent 1'

    # Headers
    hdr_cells = example_table.rows[0].cells
    headers = ['Vendor', 'Individual Spend', 'Cumulative Spend', 'Cumulative %', 'Note']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data rows
    for row_data in example_data:
        row_cells = example_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            if i < len(row_data):
                row_cells[i].text = str(cell_data)
                # Highlight the 80% threshold row
                if '80% THRESHOLD' in str(cell_data):
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_paragraph()

    # Explanation of the example
    add_styled_paragraph(doc, 'Reading this table:', bold=True, size=11)
    doc.add_paragraph()

    example_bullets = [
        'Row 1: Vendor A spent $3M individually. Cumulative spend is $3M (just Vendor A so far). This is 30% of total.',
        'Row 2: Vendor B spent $2.5M individually. Cumulative spend is $5.5M (Vendor A + B combined). This is 55% of total.',
        'Row 3: Vendor C spent $1.5M individually. Cumulative spend is $7M (Vendors A + B + C combined). This is 70% of total.',
        'Row 4: Vendor D spent $1.2M individually. Cumulative spend is $8.2M (Vendors A + B + C + D). This is 82% of total.',
        'KEY INSIGHT: By Row 4, we\'ve crossed the 80% threshold! Just 4 vendors account for 82% of spending.',
        'Row 5+: The remaining 176 vendors only account for the last 18% of spending.'
    ]

    for bullet in example_bullets:
        para = doc.add_paragraph(bullet, style='List Bullet')
        if 'KEY INSIGHT' in bullet:
            para.runs[0].bold = True
            para.runs[0].font.color.rgb = RGBColor(0, 102, 0)

    doc.add_paragraph()

    # Why this matters section
    add_heading_with_color(doc, 'Why Cumulative Metrics Are Critical for Pareto Analysis', 3)

    add_styled_paragraph(doc,
        'Without cumulative metrics, you\'d have to manually add up amounts to find patterns. '
        'The cumulative columns do this math for you and make it easy to:',
        size=11
    )

    doc.add_paragraph()

    why_critical = [
        'Find the exact row where you hit 80% of spending - just scan the Cumulative % column for the first value over 80%',
        'Count how many items it takes to reach 80% - simply count the rows up to the 80% threshold',
        'Calculate efficiency - compare the count to 80% against total items to see your concentration level',
        'Visually see how quickly spending accumulates at the top - if cumulative % grows quickly at first (e.g., 30%, 55%, 70%), you have high concentration',
        'Make the business case - show stakeholders "just 4 vendors = 82% of spending" using the cumulative data'
    ]

    for item in why_critical:
        add_bullet_point(doc, item)

    doc.add_paragraph()

    # How to use in Excel
    add_heading_with_color(doc, 'How to Use These Metrics in Your Excel File', 3)

    how_to_use = [
        ('Step 1', 'Open any worksheet (Vendors_80_20, Departments_80_20, etc.)'),
        ('Step 2', 'Look at the "Cumulative %" column - this is sorted from highest to lowest spending'),
        ('Step 3', 'Scan down until you find the first row where Cumulative % exceeds 80%'),
        ('Step 4', 'This row will usually be highlighted - everything ABOVE this row is your priority focus area'),
        ('Step 5', 'Count the rows above this threshold - this tells you how many items to prioritize'),
        ('Step 6', 'Look at the summary statistics to see "Count to 80%" and "% to 80%" for quick reference')
    ]

    for step, instruction in how_to_use:
        para = doc.add_paragraph()
        para.add_run(f'{step}: ').bold = True
        para.add_run(instruction)

    doc.add_paragraph()

    # Other metrics section
    add_heading_with_color(doc, 'Other Important Metrics', 3)

    other_metrics = [
        ('Individual Amount/Spend', 'The actual spending amount for just that single vendor/department/transaction'),
        ('Count to 80%', 'Pre-calculated count of how many items are needed to reach 80% threshold'),
        ('% to 80%', 'What percentage of total items the "Count to 80%" represents (this is the key Pareto ratio)'),
        ('Rank', 'Position in the sorted list from highest to lowest')
    ]

    for metric, description in other_metrics:
        para = doc.add_paragraph()
        para.add_run(f'{metric}: ').bold = True
        para.add_run(description)

    doc.add_paragraph()

    # Summary box
    summary_para = doc.add_paragraph()
    summary_para.add_run('Bottom Line: ').bold = True
    summary_para.add_run(
        'The Cumulative % column is your roadmap. Find where it crosses 80%, and everything above '
        'that line is your priority. Everything below that line is secondary. The cumulative metrics '
        'transform raw spending data into actionable priorities.'
    )

    doc.add_page_break()

    # ========================================================================
    # SECTION 5: How to Read Each Worksheet
    # ========================================================================
    add_heading_with_color(doc, '5. How to Read Each Worksheet', 1)

    # --- Vendors Sheet ---
    add_heading_with_color(doc, '5.1 Vendors_80_20 Worksheet', 2)

    add_styled_paragraph(doc,
        'This worksheet shows vendor spending ranked from highest to lowest.',
        size=11
    )

    doc.add_paragraph()
    add_styled_paragraph(doc, 'What to Look For:', bold=True, size=11)
    doc.add_paragraph()

    vendor_lookfor = [
        'The highlighted row shows where you hit 80% of total spending',
        'Count how many vendors are above this line - these are your priority vendors',
        'Look at the "% to 80%" - this tells you what fraction of vendors need attention',
        'Check the top 5 vendors - they often represent 40-50% of spending alone'
    ]

    for item in vendor_lookfor:
        add_bullet_point(doc, item)

    doc.add_paragraph()
    add_styled_paragraph(doc, 'How to Use This:', bold=True, size=11)
    doc.add_paragraph()

    vendor_use = [
        'Start contract negotiations with vendors above the 80% line first',
        'These vendors should be your priority for master agreements',
        'Consider these vendors "critical" for risk management',
        'Vendors below the line are lower priority - address them later'
    ]

    for item in vendor_use:
        add_bullet_point(doc, item)

    doc.add_paragraph()
    doc.add_paragraph()

    # --- Departments Sheet ---
    add_heading_with_color(doc, '5.2 Departments_80_20 Worksheet', 2)

    add_styled_paragraph(doc,
        'This worksheet shows department spending ranked from highest to lowest.',
        size=11
    )

    doc.add_paragraph()
    add_styled_paragraph(doc, 'What to Look For:', bold=True, size=11)
    doc.add_paragraph()

    dept_lookfor = [
        'The highlighted row shows where you hit 80% of total spending',
        'Departments above this line need intensive contract compliance support',
        'Look for patterns - are these departments just large, or do they have poor processes?',
        'Check if multiple high-spending departments use the same vendors'
    ]

    for item in dept_lookfor:
        add_bullet_point(doc, item)

    doc.add_paragraph()
    add_styled_paragraph(doc, 'How to Use This:', bold=True, size=11)
    doc.add_paragraph()

    dept_use = [
        'Assign dedicated procurement liaisons to departments above the 80% line',
        'Implement monthly compliance reviews for these departments',
        'Provide training and process improvement support',
        'Consider why these departments have high spending without contracts'
    ]

    for item in dept_use:
        add_bullet_point(doc, item)

    doc.add_paragraph()
    doc.add_paragraph()

    # --- Transactions Sheet ---
    add_heading_with_color(doc, '5.3 Transactions_80_20 Worksheet', 2)

    add_styled_paragraph(doc,
        'This worksheet shows individual transactions ranked by dollar amount.',
        size=11
    )

    doc.add_paragraph()
    add_styled_paragraph(doc, 'What to Look For:', bold=True, size=11)
    doc.add_paragraph()

    trans_lookfor = [
        'The highlighted row shows where you hit 80% of total spending',
        'Note the dollar amount at this cutoff - this is your "high-value threshold"',
        'Count transactions above this line - these need immediate review',
        'Look for patterns in high-value transactions (same vendors, departments, categories?)'
    ]

    for item in trans_lookfor:
        add_bullet_point(doc, item)

    doc.add_paragraph()
    add_styled_paragraph(doc, 'How to Use This:', bold=True, size=11)
    doc.add_paragraph()

    trans_use = [
        'Set up automatic alerts for transactions above the high-value threshold',
        'Require mandatory contract review for high-value purchases',
        'Consider implementing a hard stop in your procurement system',
        'Review all transactions above the line for immediate contract action'
    ]

    for item in trans_use:
        add_bullet_point(doc, item)

    doc.add_paragraph()
    doc.add_paragraph()

    # --- Categories Sheet ---
    add_heading_with_color(doc, '5.4 Categories_80_20 Worksheet', 2)

    add_styled_paragraph(doc,
        'This worksheet shows spending by category (IT, Facilities, Professional Services, etc.).',
        size=11
    )

    doc.add_paragraph()
    add_styled_paragraph(doc, 'What to Look For:', bold=True, size=11)
    doc.add_paragraph()

    cat_lookfor = [
        'The highlighted row shows where you hit 80% of total spending',
        'Categories above this line need category management strategies',
        'Look for surprising categories - unexpected high spending may indicate miscategorization',
        'Compare categories to your organization\'s mission and priorities'
    ]

    for item in cat_lookfor:
        add_bullet_point(doc, item)

    doc.add_paragraph()
    add_styled_paragraph(doc, 'How to Use This:', bold=True, size=11)
    doc.add_paragraph()

    cat_use = [
        'Develop category-specific contract strategies for categories above the line',
        'Negotiate pre-approved pricing catalogs for high-spend categories',
        'Consider volume discount agreements within top categories',
        'Assign category managers to oversee top spending areas'
    ]

    for item in cat_use:
        add_bullet_point(doc, item)

    doc.add_paragraph()
    doc.add_paragraph()

    # --- Understanding Category Management ---
    add_heading_with_color(doc, 'Understanding Category Management: Why This Matters', 3)

    add_styled_paragraph(doc,
        'Category management is a strategic approach where you treat each major spending category '
        '(like IT, Facilities, Professional Services, Parts & Supplies) as its own strategic business unit '
        'with dedicated oversight and specialized strategies.',
        size=11
    )

    doc.add_paragraph()

    add_styled_paragraph(doc,
        'When the Pareto analysis identifies categories that account for 80% of spending, these categories '
        'have enough volume to justify a focused category management approach. Here\'s why this matters:',
        size=11
    )

    doc.add_paragraph()

    # Four key reasons
    add_styled_paragraph(doc, '1. Categories Represent Different Markets', bold=True, size=11)
    doc.add_paragraph()

    reason1_bullets = [
        'Parts & Supplies vendors operate differently than Professional Services vendors',
        'Each category has unique pricing models, contract structures, and negotiation strategies',
        'IT spending (software licenses, subscriptions) requires different approaches than Facilities (maintenance, repairs)',
        'Specialized knowledge is needed for each category to negotiate effectively'
    ]

    for bullet in reason1_bullets:
        add_bullet_point(doc, bullet)

    doc.add_paragraph()

    add_styled_paragraph(doc, '2. Volume Leverage Within Categories', bold=True, size=11)
    doc.add_paragraph()

    reason2_bullets = [
        'When you spend $17.4M in Parts & Supplies, that volume gives you negotiating power',
        'Create pre-approved catalogs of common items with negotiated prices',
        'Negotiate volume discounts across all vendors within the category',
        'Standardize specifications and reduce variation to increase leverage',
        'Instead of each department negotiating separately, pool your buying power'
    ]

    for bullet in reason2_bullets:
        add_bullet_point(doc, bullet)

    doc.add_paragraph()

    add_styled_paragraph(doc, '3. Pattern Recognition and Standardization', bold=True, size=11)
    doc.add_paragraph()

    reason3_bullets = [
        'Categories above the 80% line have enough spend to justify dedicated management resources',
        'Example: $11.4M in Facilities & Maintenance justifies hiring a Facilities Category Manager',
        'Develop master service agreements with standard terms across the category',
        'Create preferred vendor lists with pre-vetted, qualified suppliers',
        'Establish standardized service level agreements (SLAs) and quality metrics'
    ]

    for bullet in reason3_bullets:
        add_bullet_point(doc, bullet)

    doc.add_paragraph()

    add_styled_paragraph(doc, '4. Efficiency at Scale', bold=True, size=11)
    doc.add_paragraph()

    reason4_bullets = [
        'Without category management: Each of 50+ departments negotiates separately with vendors',
        'With category management: One expert negotiates on behalf of all departments',
        'Result: Better prices, consistent quality standards, reduced administrative burden',
        'Procurement staff focus on strategic vendor relationships instead of transactional purchasing',
        'Easier to track compliance and enforce contract usage'
    ]

    for bullet in reason4_bullets:
        add_bullet_point(doc, bullet)

    doc.add_paragraph()

    # Real-world example box
    add_heading_with_color(doc, 'Real-World Example: Parts & Supplies', 3)

    add_styled_paragraph(doc,
        'Your analysis shows Parts & Supplies as the #1 category at $17.4M (25.2% of total spending). '
        'Here\'s how category management creates value:',
        size=11
    )

    doc.add_paragraph()

    # Without scenario
    without_para = doc.add_paragraph()
    without_para.add_run('Without Category Management:\n').bold = True
    without_para.add_run(
        '• 50+ departments each order parts from various vendors\n'
        '• Each department negotiates their own prices\n'
        '• Same part might be purchased at different prices by different departments\n'
        '• No volume discounts because purchases are fragmented\n'
        '• High administrative cost - each purchase requires separate processing\n'
        '• Inconsistent quality because no standardization'
    )

    doc.add_paragraph()

    # With scenario
    with_para = doc.add_paragraph()
    with_para.add_run('With Category Management:\n').bold = True
    with_para.add_run(
        '• Create a centralized catalog of commonly purchased parts\n'
        '• Negotiate volume-based pricing (e.g., "City of Chicago buys $17M/year, give us 15% discount")\n'
        '• All departments order from the catalog at consistent, pre-negotiated rates\n'
        '• Streamlined procurement process reduces administrative costs\n'
        '• Quality standards enforced across all parts\n'
        '• Easy to track spending and identify additional savings opportunities'
    )

    doc.add_paragraph()

    # Result
    result_para = doc.add_paragraph()
    result_para.add_run('Result:\n').bold = True
    result_run = result_para.add_run(
        'Potential 10-20% cost savings on $17.4M = $1.7M to $3.5M annual savings, '
        'plus reduced administrative costs and improved quality control.'
    )
    result_run.font.color.rgb = RGBColor(0, 102, 0)

    doc.add_paragraph()

    # Strategic connection
    add_heading_with_color(doc, 'Connecting Category Analysis to Vendor Analysis', 3)

    add_styled_paragraph(doc,
        'Category management becomes even more powerful when combined with your vendor-level Pareto analysis:',
        size=11
    )

    doc.add_paragraph()

    connection_bullets = [
        'Start with category analysis: Identify that Parts & Supplies = $17.4M (top priority category)',
        'Then apply vendor Pareto within that category: Find that 20% of parts vendors account for 80% of the $17.4M',
        'Focus your category management efforts on those top vendors first',
        'Example: If 5 vendors account for $13.9M of the $17.4M Parts & Supplies spend, negotiate master agreements with those 5 first',
        'This "double Pareto" approach (category + vendor) maximizes efficiency: Focus on the right category, then focus on the right vendors within that category'
    ]

    for bullet in connection_bullets:
        add_bullet_point(doc, bullet)

    doc.add_paragraph()

    # Bottom line
    bottomline_para = doc.add_paragraph()
    bottomline_para.add_run('Bottom Line: ').bold = True
    bottomline_para.add_run(
        'Category management isn\'t just about organization - it\'s about using your purchasing volume '
        'as leverage to get better prices, better quality, and better service. The Pareto analysis shows you '
        'which categories have enough volume to make this approach worthwhile.'
    )

    doc.add_page_break()

    # ========================================================================
    # SECTION 6: Key Findings
    # ========================================================================
    add_heading_with_color(doc, '6. Key Findings from Your Analysis', 1)

    add_styled_paragraph(doc,
        'Based on the Pareto analysis of Chicago\'s $68.9M spending without active contracts, '
        'here are the critical findings:',
        size=11
    )

    doc.add_paragraph()

    # Finding 1 - Vendors
    add_heading_with_color(doc, 'Finding #1: Vendor Concentration (MOST CRITICAL)', 2)

    finding1_para = doc.add_paragraph()
    finding1_para.add_run('15% of vendors = 80% of spending').bold = True
    finding1_para.add_run('\n\n')
    finding1_para.add_run('Out of 207 total vendors, only 31 vendors (15%) account for $55.1M (80%) of spending.')

    doc.add_paragraph()

    key_vendors = [
        'Badger Meter Inc - $8.0M',
        'USALCO - $5.2M',
        'AT&T Corp - $4.3M',
        'Industrial Fence Inc - $4.0M',
        'Schindler Elevator Corp - $3.6M'
    ]

    doc.add_paragraph('Top 5 Vendors alone = $27.8M (40.4%):', style='List Bullet')
    for vendor in key_vendors:
        para = doc.add_paragraph(vendor, style='List Bullet 2')
        para.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()

    implication1 = doc.add_paragraph()
    implication1.add_run('Implication: ').bold = True
    implication1.add_run(
        'Focus contract negotiations on just 31 vendors to protect 80% of spending. '
        'This is far more concentrated than the typical 80/20 rule, making targeted action even more effective.'
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Finding 2 - Departments
    add_heading_with_color(doc, 'Finding #2: Department Concentration', 2)

    finding2_para = doc.add_paragraph()
    finding2_para.add_run('18.5% of departments = 80% of spending').bold = True
    finding2_para.add_run('\n\n')
    finding2_para.add_run('Out of 27 departments, only 5 departments (18.5%) account for $55.1M (80%) of spending.')

    doc.add_paragraph()

    key_depts = [
        'Water Management - $17.8M (25.9%)',
        'Aviation - $11.8M (17.2%)',
        'Fleet & Facility Management - $10.7M (15.5%)',
        'Technology & Innovation - $6.7M (9.7%)',
        'Streets & Sanitation - $5.8M (8.4%)'
    ]

    doc.add_paragraph('Top 5 Departments = $51.5M (74.7%):', style='List Bullet')
    for dept in key_depts:
        para = doc.add_paragraph(dept, style='List Bullet 2')
        para.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()

    implication2 = doc.add_paragraph()
    implication2.add_run('Implication: ').bold = True
    implication2.add_run(
        'Target just 5 departments for intensive contract compliance support. '
        'Assign dedicated procurement liaisons to these departments.'
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Finding 3 - Transactions
    add_heading_with_color(doc, 'Finding #3: Transaction Size Concentration (EXTREME)', 2)

    finding3_para = doc.add_paragraph()
    finding3_para.add_run('9.4% of transactions = 80% of spending').bold = True
    finding3_para.add_run('\n\n')
    finding3_para.add_run(
        'Out of 1,112 total transactions, only 105 transactions (9.4%) account for $55.1M (80%) of spending.'
    )

    doc.add_paragraph()

    trans_note = doc.add_paragraph('Key Point: ', style='List Bullet')
    trans_note.runs[0].bold = True
    trans_note.add_run('Top 36 transactions over $500K account for $40.9M (59.3%)')

    doc.add_paragraph()

    implication3 = doc.add_paragraph()
    implication3.add_run('Implication: ').bold = True
    implication3.add_run(
        'Focus on high-value transactions above $500K for maximum impact. '
        'Implement mandatory contract review and hard stops for large purchases.'
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Finding 4 - Categories
    add_heading_with_color(doc, 'Finding #4: Category Concentration', 2)

    finding4_para = doc.add_paragraph()
    finding4_para.add_run('41.7% of categories = 80% of spending').bold = True
    finding4_para.add_run('\n\n')
    finding4_para.add_run('Out of 12 categories, 5 categories (41.7%) account for $57.6M (83.6%) of spending.')

    doc.add_paragraph()

    key_cats = [
        'Parts & Supplies - $17.4M (25.2%)',
        'Facilities & Maintenance - $11.4M (16.5%)',
        'Professional Services - $9.6M (13.9%)',
        'Environmental Services - $6.9M (10.0%)',
        'Utilities - $6.2M (9.0%)'
    ]

    doc.add_paragraph('Top 5 Categories:', style='List Bullet')
    for cat in key_cats:
        para = doc.add_paragraph(cat, style='List Bullet 2')
        para.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()

    implication4 = doc.add_paragraph()
    implication4.add_run('Implication: ').bold = True
    implication4.add_run(
        'Implement category management strategies for these five areas. '
        'Negotiate pre-approved pricing catalogs and volume discounts.'
    )

    doc.add_page_break()

    # ========================================================================
    # SECTION 7: How to Use for Decision-Making
    # ========================================================================
    add_heading_with_color(doc, '7. How to Use This Analysis for Decision-Making', 1)

    add_styled_paragraph(doc,
        'The Pareto analysis provides a clear roadmap for prioritization. Here\'s how to put it into action:',
        size=11
    )

    doc.add_paragraph()

    # Phase 1
    add_heading_with_color(doc, 'PHASE 1: Immediate Action (0-30 Days)', 2)

    phase1_target = doc.add_paragraph()
    phase1_target.add_run('Target: ').bold = True
    phase1_target.add_run('Top 5 Vendors (40% of spend = $27.8M)')

    doc.add_paragraph()

    phase1_actions = [
        'Emergency contract awards or amendments for critical vendors',
        'Fast-track master agreements with Badger Meter, USALCO, AT&T, Industrial Fence, and Schindler',
        'Negotiate immediate compliance plans',
        'Set up recurring orders under new contracts'
    ]

    doc.add_paragraph('Actions:', style='List Bullet')
    for action in phase1_actions:
        para = doc.add_paragraph(action, style='List Bullet 2')
        para.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()
    doc.add_paragraph()

    # Phase 2
    add_heading_with_color(doc, 'PHASE 2: Short Term (30-90 Days)', 2)

    phase2_target = doc.add_paragraph()
    phase2_target.add_run('Target: ').bold = True
    phase2_target.add_run('Remaining 26 vendors to reach 80% threshold (additional $27.3M)')

    doc.add_paragraph()

    phase2_actions = [
        'Negotiate master agreements with remaining 26 priority vendors',
        'Establish blanket purchase orders where appropriate',
        'Set up preferred vendor programs',
        'Assign dedicated procurement liaisons to top 5 departments',
        'Implement department-specific contract tracking',
        'Begin monthly compliance reviews with high-spend departments'
    ]

    doc.add_paragraph('Actions:', style='List Bullet')
    for action in phase2_actions:
        para = doc.add_paragraph(action, style='List Bullet 2')
        para.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()
    doc.add_paragraph()

    # Phase 3
    add_heading_with_color(doc, 'PHASE 3: Medium Term (90-180 Days)', 2)

    phase3_target = doc.add_paragraph()
    phase3_target.add_run('Target: ').bold = True
    phase3_target.add_run('System-wide controls and category management')

    doc.add_paragraph()

    phase3_actions = [
        'Implement $500K hard stop in procurement system (prevents unauthorized large purchases)',
        'Mandatory contract review for all transactions above $100K',
        'CFO or senior approval required for transactions above $500K',
        'Develop category-specific strategies for top 5 spending categories',
        'Negotiate pre-negotiated pricing catalogs',
        'Implement volume discount agreements',
        'Assign category managers to oversee Parts & Supplies, Facilities, Professional Services, etc.'
    ]

    doc.add_paragraph('Actions:', style='List Bullet')
    for action in phase3_actions:
        para = doc.add_paragraph(action, style='List Bullet 2')
        para.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()

    # Why This Approach Works
    add_heading_with_color(doc, 'Why This Phased Approach Works', 2)

    why_works = [
        ('Quick Wins',
         'Phase 1 addresses 40% of spending with just 5 vendors - shows immediate progress'),
        ('Focused Resources',
         'Instead of trying to fix everything at once, you concentrate on high-impact areas'),
        ('Measurable Progress',
         'Each phase has clear targets and outcomes - easy to track and report'),
        ('Risk Reduction',
         'Systematically reduces the most significant risks first'),
        ('Efficient Use of Time',
         'Procurement staff focus on relationships that matter most')
    ]

    for title, description in why_works:
        para = doc.add_paragraph()
        para.add_run(f'{title}: ').bold = True
        para.add_run(description)

    doc.add_page_break()

    # ========================================================================
    # SECTION 8: Q&A
    # ========================================================================
    add_heading_with_color(doc, '8. Common Questions and Answers', 1)

    questions = [
        (
            'Q: Why not just address all vendors equally?',
            'A: With 207 vendors, equal attention would spread resources too thin. '
            'By focusing on 31 vendors first, you protect 80% of spending with only 15% of the effort. '
            'The remaining 176 vendors represent only 20% of spending - lower priority, lower risk.'
        ),
        (
            'Q: What if a small vendor is critical for operations?',
            'A: The Pareto analysis focuses on spending volume, not operational criticality. '
            'Use the analysis as a guide, but apply business judgment. If a $50K vendor provides '
            'critical services, prioritize them even if they\'re below the 80% threshold.'
        ),
        (
            'Q: Does this mean we ignore the other 80% of vendors?',
            'A: Not ignore - deprioritize. Address the critical 20% first (Phase 1-2), then systematically '
            'work through the remaining 80% (Phase 3+). The analysis helps you sequence your efforts, not '
            'eliminate work.'
        ),
        (
            'Q: How often should we run this analysis?',
            'A: Quarterly at minimum, monthly is better. As you address high-priority vendors, new ones may '
            'emerge. Regular analysis ensures you\'re always focused on current priorities.'
        ),
        (
            'Q: What if our concentration is different (e.g., 70/30 or 90/10)?',
            'A: The principle still applies. If it\'s 90/10 (more concentrated), your job is easier - '
            'focus on even fewer vendors. If it\'s 70/30 (less concentrated), you\'ll need to address '
            'a larger group to capture the same percentage of spending. Adjust your strategy accordingly.'
        ),
        (
            'Q: Can we use this analysis for other purposes?',
            'A: Absolutely! Apply the same principle to:\n'
            '  • Identifying top taxpayers for audit\n'
            '  • Focusing customer service on high-volume users\n'
            '  • Prioritizing infrastructure repairs by impact\n'
            '  • Allocating training resources to high-use systems'
        ),
        (
            'Q: What\'s the ROI of using this approach?',
            'A: Based on the analysis:\n'
            '  • Effort: 15% of vendor relationships (31 vendors)\n'
            '  • Return: 80% of risk mitigation ($55.1M protected)\n'
            '  • ROI: 5.3x return on effort\n\n'
            'This means for every hour spent on the critical 31 vendors, you get 5.3 times the value '
            'compared to spreading that hour across all 207 vendors equally.'
        ),
    ]

    for question, answer in questions:
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(question)
        q_run.font.bold = True
        q_run.font.size = Pt(11)
        q_run.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph()

        a_para = doc.add_paragraph(answer)
        a_para.paragraph_format.left_indent = Inches(0.5)

        doc.add_paragraph()

    doc.add_page_break()

    # ========================================================================
    # CONCLUSION
    # ========================================================================
    add_heading_with_color(doc, 'Conclusion: The Power of Focus', 1)

    conclusion_text = [
        'The Pareto Principle isn\'t just a mathematical curiosity - it\'s a practical tool for making '
        'better decisions with limited resources.',

        'Your analysis shows that Chicago\'s spending without contracts follows an even more extreme pattern '
        'than the traditional 80/20 rule in some areas. This concentration creates both risk and opportunity:',
    ]

    for text in conclusion_text:
        add_styled_paragraph(doc, text, size=11)
        doc.add_paragraph()

    risks_opps = [
        ('⚠️ Risk:', 'Just 31 vendors control 80% of spending - high dependency'),
        ('✓ Opportunity:', 'Focus on 31 vendors delivers 80% risk reduction - high efficiency'),
    ]

    for label, text in risks_opps:
        para = doc.add_paragraph()
        para.add_run(f'{label} ').bold = True
        para.add_run(text)

    doc.add_paragraph()

    final_text = [
        'By using this Excel analysis as your roadmap, you can:',
        '',
    ]

    for text in final_text:
        if text:
            add_styled_paragraph(doc, text, size=11)
        else:
            doc.add_paragraph()

    final_points = [
        'Make data-driven decisions about where to focus',
        'Demonstrate strategic thinking to stakeholders',
        'Achieve measurable progress quickly',
        'Use limited resources efficiently',
        'Reduce the most significant risks first'
    ]

    for point in final_points:
        add_bullet_point(doc, point)

    doc.add_paragraph()
    doc.add_paragraph()

    # Final message
    final_para = doc.add_paragraph()
    final_run = final_para.add_run(
        'Remember: Perfect is the enemy of good. Don\'t wait to address every vendor - '
        'start with the critical few, and you\'ll accomplish more than trying to do everything at once.'
    )
    final_run.font.italic = True
    final_run.font.size = Pt(12)
    final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Contact info
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run(
        '---\n\n'
        'Questions about this analysis?\n'
        'Contact: Procurement Department\n'
        'Analysis Date: January 2026'
    )
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    output_path = '/workspaces/Spend_Analysis_and_Categorization/spend-analysis/outputs/PARETO_ANALYSIS_GUIDE.docx'
    doc.save(output_path)
    print(f"Document created successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    create_pareto_guide()
