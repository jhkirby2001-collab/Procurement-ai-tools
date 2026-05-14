"""Build the Standard Operating Procedure document for the NIGP Classification System.

Generates one file in /workspaces/Procurement-ai-tools/outputs/:
  NIGP_SOP_JHK3.docx — Standard Operating Procedure, three sections (Operations,
                       Processing, Governance).

Style is borrowed from build_leadership_deliverables_JHK3.py so the SOP matches
NIGP_Executive_Brief_JHK3.docx and NIGP_Methodology_for_Leadership_JHK3.docx.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

# Reuse helpers + brand palette from the leadership-deliverables script.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from build_leadership_deliverables_JHK3 import (  # noqa: E402
    CHI_LT_BLUE_HEX,
    add_body,
    add_bullet,
    add_callout_box,
    add_h1,
    add_h2,
    add_title_banner,
    set_cell_shading,
    set_doc_margins,
    set_table_borders,
    style_table_header,
)

OUT_DIR = Path("/workspaces/Procurement-ai-tools/outputs")
OUT_PATH = OUT_DIR / "NIGP_SOP_JHK3.docx"

DOC_VERSION = "Version 1.0  •  14 May 2026"


def _shade_alt_rows(table, color_hex=CHI_LT_BLUE_HEX):
    for i, row in enumerate(table.rows[1:]):
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, color_hex)
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)


def _add_simple_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    style_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    _shade_alt_rows(table)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    return table


def build_sop():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_doc_margins(doc, cm=1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # -------------------------------------------------------------------
    # Title banner
    # -------------------------------------------------------------------
    add_title_banner(
        doc,
        title="Standard Operating Procedure",
        subtitle="NIGP Procurement Classification System",
        author_line=(
            "City of Chicago  •  Department of Procurement Services  •  "
            "Prepared by James H. Kirby III, CSCP, MS-SCM  •  " + DOC_VERSION
        ),
    )

    # -------------------------------------------------------------------
    # Purpose + Scope (short preamble)
    # -------------------------------------------------------------------
    add_h1(doc, "Purpose and Scope")
    add_body(
        doc,
        "This Standard Operating Procedure defines how the NIGP Procurement Classification "
        "System is operated, how new procurement data extracts are processed through it, and "
        "how the system and its taxonomy are governed over time. It is written for procurement "
        "analysts and managers responsible for day-to-day use and stewardship of the system.",
    )
    add_body(
        doc,
        "The procedures below assume no Python knowledge. Every step can be executed by a "
        "procurement analyst following the instructions verbatim. Where a step requires technical "
        "assistance, that is called out explicitly.",
    )

    add_callout_box(
        doc,
        label="DOCUMENT STRUCTURE",
        body_text=(
            "Section 1 — Classification System Operations (what the system is, where it lives, "
            "how to use the web app).  "
            "Section 2 — Data Processing Procedures (how to run a new batch, QA the output, "
            "triage the review queue, and maintain rules).  "
            "Section 3 — Taxonomy Governance (roles, annual reviews, accuracy audits, "
            "succession, and document control)."
        ),
    )

    # ===================================================================
    # SECTION 1 — OPERATIONS
    # ===================================================================
    add_h1(doc, "Section 1 — Classification System Operations")

    # ---- 1.1 System Overview ----
    add_h2(doc, "1.1  System Overview")
    add_body(
        doc,
        "The NIGP Procurement Classification System assigns a consistent commodity classification "
        "to every line of City procurement spend. The classifier reads each record's description "
        "text and Chicago FMPS account codes, applies a library of keyword and account-code rules, "
        "and writes the resulting classification back to the row. The same engine powers both the "
        "bulk historical batch run and the single-record web app used for live PO classification.",
    )
    add_h2(doc, "Three classification tiers (assigned on every record)")
    add_bullet(doc, "17 mutually exclusive buckets covering everything the City buys (Tier 1 — leadership-facing).", bold_lead="Business Categories.")
    add_bullet(doc, "138 distinct NIGP 3-digit Classes derived from the historical data (Tier 2 — sourcing- and audit-facing).", bold_lead="NIGP 3-digit Classes.")
    add_bullet(doc, "470 distinct NIGP 5-digit Class-Items (Tier 3 — used where the description supports specificity).", bold_lead="NIGP 5-digit Class-Items.")

    add_h2(doc, "Two operating modes")
    add_bullet(doc, "Processes a full procurement extract (~784,000 rows) end-to-end in roughly 15 minutes. Produces the full classified file, the review-queue subset, and a diagnostic coverage report.", bold_lead="Batch mode.")
    add_bullet(doc, "Procurement staff paste a single description into the web app and immediately receive the proposed Business Category, NIGP code, confidence, and reason.", bold_lead="Single-record mode (web app).")

    # ---- 1.2 System Components ----
    add_h2(doc, "1.2  System Components and File Locations")
    add_body(doc, "The system is hosted in a single GitHub repository. All files below are versioned and editable.")

    components = [
        ("classifier_JHK3.py", "spend-analysis/scripts/", "The classifier engine (batch + single-record)."),
        ("keyword_rules_DRAFT_JHK3.csv", "spend-analysis/data/reference/", "The combined rule file (hand-curated + AI-mined). Procurement-staff editable."),
        ("business_categories_JHK3.csv", "spend-analysis/data/reference/", "138-row NIGP-class to Business Category mapping."),
        ("business_categories_summary_JHK3.csv", "spend-analysis/data/reference/", "17-row Business Categories reference table."),
        ("NIGP_Mapping_JHK3.csv", "outputs/", "The full classified output file (one row per source record)."),
        ("NIGP_Mapping_Review_Queue_JHK3.csv", "outputs/", "The review-flagged subset for staff triage."),
        ("classifier_coverage_report_JHK3.csv", "spend-analysis/data/processed/", "Diagnostic — which rules fired how many times."),
    ]
    _add_simple_table(
        doc,
        headers=["File", "Location", "What it is"],
        rows=components,
        col_widths_cm=[5.5, 5.5, 6.0],
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_h2(doc, "External resources")
    add_bullet(doc, "GitHub repository (private to authorized staff). Source of truth for all scripts and rule files.", bold_lead="Repository:")
    add_bullet(doc, "https://chicago-nigp-classifier.streamlit.app/ — the single-record classification web app.", bold_lead="Web app URL:")
    add_bullet(doc, "Saved one-time AI mining output (spend-analysis/data/processed/ai_classified_unique_descriptions_JHK3.csv) — preserved for the AI-assist fallback layer. Do NOT regenerate.", bold_lead="AI mining output:")

    # ---- 1.3 Access and Credentials ----
    add_h2(doc, "1.3  Access and Credentials")
    add_bullet(doc, "Granted to the System Owner and designated procurement-staff maintainers. Read access for additional reviewers as needed.", bold_lead="GitHub repository access:")
    add_bullet(doc, "Shared with authorized procurement staff. Rotation is at the discretion of the System Owner.", bold_lead="Web app password:")
    add_bullet(doc, "James H. Kirby III, CSCP, MS-SCM. Responsible for access decisions, rule-quality oversight, and taxonomy integrity.", bold_lead="Designated System Owner:")

    add_h2(doc, "Procedure for granting access to new staff")
    add_bullet(doc, "Requester submits a written request to the System Owner identifying the role (read-only, rule maintainer, or triager) and the business justification.")
    add_bullet(doc, "System Owner approves or denies in writing.")
    add_bullet(doc, "If approved: System Owner adds the requester to the GitHub repository at the appropriate permission level and shares the web app password through an approved channel.")
    add_bullet(doc, "System Owner updates the Roles and Responsibilities table in Section 3.1 of this SOP.")
    add_bullet(doc, "Upon departure or role change: System Owner revokes access within five business days.")

    # ---- 1.4 Web App Operations ----
    add_h2(doc, "1.4  Web App Operations")
    add_body(doc, "The web app provides a no-installation interface to the classifier. Any authorized user can classify a single description, paste a batch, browse the taxonomy, or look up how any rule routes a given phrase.")

    add_h2(doc, "How to access")
    add_bullet(doc, "Navigate to https://chicago-nigp-classifier.streamlit.app/ in any modern browser.")
    add_bullet(doc, "Enter the password at the access gate. The app loads to the Classify page by default.")

    add_h2(doc, "The six pages")
    pages = [
        ("Classify", "Paste a single description; receive the Business Category, NIGP class, NIGP item, confidence, and reason."),
        ("Bulk Classify", "Paste many descriptions (one per line) or upload a CSV. Returns a downloadable CSV of classifications."),
        ("Procurement Taxonomy Logic", "Visual explainer of the three-tier hierarchy (Executive View → Sourcing View → Audit View)."),
        ("Methodology", "How the engine decides; confidence-badge logic; AI-use defensibility."),
        ("Business Categories", "Browseable list of all 17 Business Categories with definitions."),
        ("Rule Lookup", "Search across all rules to verify how any specific phrase routes."),
    ]
    _add_simple_table(
        doc,
        headers=["Page", "What it does"],
        rows=pages,
        col_widths_cm=[5.0, 12.0],
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_h2(doc, "What to do if the app is slow or unresponsive")
    add_bullet(doc, "The app runs on Streamlit Community Cloud, which puts free-tier apps into a sleep state after a period of inactivity. The first request after sleep can take 30 to 90 seconds to return.", bold_lead="Cold start behavior.")
    add_bullet(doc, "Refresh the page once and wait. If still unresponsive after two minutes, try an incognito or private window to rule out browser cache.", bold_lead="First action.")
    add_bullet(doc, "Notify the System Owner. The app may need a manual reboot from the Streamlit Cloud dashboard.", bold_lead="If still down.")
    add_bullet(doc, "If the URL fails to load entirely, fall back to running classifier_JHK3.py in single-record CLI mode from a project codespace.", bold_lead="Fallback.")

    add_h2(doc, "Contact for issues")
    add_bullet(doc, "James H. Kirby III, CSCP, MS-SCM — for all access, classification, taxonomy, or web app issues.", bold_lead="System Owner:")

    # ===================================================================
    # SECTION 2 — DATA PROCESSING
    # ===================================================================
    add_h1(doc, "Section 2 — Data Processing Procedures")

    # ---- 2.1 Processing a new extract ----
    add_h2(doc, "2.1  Processing a New Procurement Data Extract")
    add_body(doc, "Follow these steps each time a new procurement data extract is received and needs to be classified.")
    add_bullet(doc, "Receive the new data extract file (typically Excel .xlsx or CSV format).", bold_lead="Step 1.")
    add_bullet(doc, "Save the file to the raw data folder in the repository: spend-analysis/data/raw/. Use a clear, dated filename (for example, ey_raw_2026Q3.xlsx).", bold_lead="Step 2.")
    add_bullet(doc, "Convert the Excel/CSV file to parquet (the format the classifier reads). In a terminal, run a short Python command: import pandas as pd; pd.read_excel('spend-analysis/data/raw/<your file>.xlsx').to_parquet('spend-analysis/data/processed/ey_raw_cache.parquet'). If unsure, contact the System Owner — this is a one-line task but only runs successfully when the source file's columns match the expected schema.", bold_lead="Step 3.")
    add_bullet(doc, "Open classifier_JHK3.py in any text editor. Confirm the SRC_PARQUET line near the top of the file points to the parquet file you just created. Save and close.", bold_lead="Step 4.")
    add_bullet(doc, "Open a terminal in the project codespace and run the classifier in batch mode:  python /workspaces/Procurement-ai-tools/spend-analysis/scripts/classifier_JHK3.py --batch", bold_lead="Step 5.")
    add_bullet(doc, "Wait approximately 15 minutes for processing to complete. The terminal will print progress messages as it runs.", bold_lead="Step 6.")
    add_bullet(doc, "Verify the three output files were generated: NIGP_Mapping_JHK3.csv (full classified dataset), NIGP_Mapping_Review_Queue_JHK3.csv (review queue), and classifier_coverage_report_JHK3.csv (diagnostic).", bold_lead="Step 7.")
    add_bullet(doc, "Run the QA checks listed in Section 2.2 below before distributing the output to leadership or downstream consumers.", bold_lead="Step 8.")

    # ---- 2.2 Post-Run QA ----
    add_h2(doc, "2.2  Post-Run Quality Assurance Checks")
    add_body(doc, "After every batch run, verify the following before treating the output as final.")
    add_bullet(doc, "Total row count in the output matches the total row count in the input. Any discrepancy means rows were dropped — investigate before publishing.", bold_lead="Row-count integrity.")
    add_bullet(doc, "All 17 Business Categories (plus the explicit \"Unclassified — No Description\" tag) appear in the output. Any missing category means no rule fired for it in this batch — verify whether that's accurate or whether a rule has been inadvertently retired.", bold_lead="Category presence.")
    add_bullet(doc, "No Business_Category field is null. Every row should carry an assignment, even if only to \"Unclassified — No Description.\"", bold_lead="No nulls.")
    add_bullet(doc, "Open classifier_coverage_report_JHK3.csv. Rules that fired zero times may indicate retired commodity types or rule-pattern typos. Flag for the Rule File Maintainer.", bold_lead="Zero-fire rule review.")
    add_bullet(doc, "Check whether any new descriptions hit no rule (rows tagged human_review or AI-assist with low confidence). These may indicate new commodity types needing new rules.", bold_lead="Coverage gap review.")
    add_bullet(doc, "Open NIGP_Mapping_JHK3.csv and spot-check ten random rows. Confirm the description, the assigned Business Category, and the assigned NIGP class look reasonable together.", bold_lead="Spot check.")

    # ---- 2.3 Review queue triage ----
    add_h2(doc, "2.3  Review Queue Triage Procedure")
    add_body(
        doc,
        "The review queue contains rows where the classifier assigned a category with low confidence, "
        "or could not assign at all. Triage on a monthly or quarterly cadence to keep the queue from "
        "accumulating and to feed new rules back into the system.",
    )
    add_bullet(doc, "Open NIGP_Mapping_Review_Queue_JHK3.csv in Excel.", bold_lead="Step 1.")
    add_bullet(doc, "Sort by Confidence_Score ascending (lowest confidence first). Work the lowest-confidence rows down.", bold_lead="Step 2.")
    add_bullet(doc, "For each row, read the Description_Best field and the proposed Business_Category.", bold_lead="Step 3.")
    add_bullet(doc, "Make a decision per row:  (a) Correct classification — no action needed; optionally convert to a permanent keyword rule (see Section 2.4).  (b) Wrong classification — determine the correct Business Category and NIGP code; add a corrective keyword rule (see Section 2.4).  (c) Unclassifiable — description is too vague to classify from text alone; document and leave as review-flagged.", bold_lead="Step 4.")
    add_bullet(doc, "Track the number of rows triaged per session for reporting purposes. Triage volume becomes a leading indicator of taxonomy maturity over time.", bold_lead="Step 5.")

    # ---- 2.4 Adding or modifying rules ----
    add_h2(doc, "2.4  Adding or Modifying Keyword Rules")
    add_body(doc, "When a description needs a new rule or a correction, edit the rule file directly. No Python knowledge required.")
    add_bullet(doc, "Open spend-analysis/data/reference/keyword_rules_DRAFT_JHK3.csv in Excel.", bold_lead="Step 1.")
    add_bullet(doc, "Add a new row at the bottom of the file. The columns are described below.", bold_lead="Step 2.")
    add_bullet(doc, "Save the file (keep CSV format — do not save as .xlsx).", bold_lead="Step 3.")
    add_bullet(doc, "Re-run the classifier in batch mode (see Section 2.1, Steps 5 through 7) to apply the new rule to all rows.", bold_lead="Step 4.")
    add_bullet(doc, "Verify the rule fired correctly by checking classifier_coverage_report_JHK3.csv — search for your new pattern; the row-count column should be non-zero.", bold_lead="Step 5.")

    add_h2(doc, "Rule-row column reference")
    rule_columns = [
        ("pattern", "The description text to match (the keyword or phrase the classifier looks for)."),
        ("match_type", "How the pattern matches:  exact (full string),  starts_with (prefix),  contains (substring anywhere in the description)."),
        ("business_category", "The correct Business Category. Must exactly match one of the 17 categories in business_categories_summary_JHK3.csv."),
        ("nigp_class_3digit", "The correct NIGP 3-digit Class. Leave blank for the \"Grants & Pass-Through Funding\" category, which has no NIGP class."),
        ("nigp_item_5digit", "The NIGP 5-digit Class-Item code, where the description supports that level of specificity. Leave blank if not specific enough."),
        ("nigp_match_level", "Set to broad for normal cases.  Set to exact only for full-string rules where you are confident in the 5-digit assignment."),
        ("notes", "Author initials, date added, and a short justification for the rule. Important for audit traceability."),
    ]
    _add_simple_table(
        doc,
        headers=["Column", "What to enter"],
        rows=rule_columns,
        col_widths_cm=[4.5, 12.5],
    )

    # ---- 2.5 Retiring rules ----
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_h2(doc, "2.5  Retiring Outdated Rules")
    add_body(doc, "When a rule is no longer relevant — for example, a discontinued commodity type or a duplicate that has been superseded — retire it as follows.")
    add_bullet(doc, "Open keyword_rules_DRAFT_JHK3.csv. Locate the rule row to retire.", bold_lead="Step 1.")
    add_bullet(doc, "Either delete the row entirely, or move it to a separate retired_rules.csv archive file in the same folder (preferred — preserves history for audit).", bold_lead="Step 2.")
    add_bullet(doc, "In the archive's notes column, add the retirement date, your initials, and a short reason for retirement.", bold_lead="Step 3.")
    add_bullet(doc, "Re-run the classifier in batch mode to apply the change.", bold_lead="Step 4.")
    add_bullet(doc, "Verify the retired rule no longer appears in classifier_coverage_report_JHK3.csv.", bold_lead="Step 5.")

    # ===================================================================
    # SECTION 3 — GOVERNANCE
    # ===================================================================
    add_h1(doc, "Section 3 — Taxonomy Governance")

    # ---- 3.1 Roles ----
    add_h2(doc, "3.1  Roles and Responsibilities")
    roles = [
        ("System Owner", "Overall accountability for the classification system, rule quality, and taxonomy integrity. Approves all access requests and rule-policy changes.", "James H. Kirby III"),
        ("Rule File Maintainer", "Day-to-day additions, corrections, and retirements of keyword rules. Owns the keyword_rules_DRAFT_JHK3.csv file.", "[To be designated]"),
        ("Review Queue Triager", "Monthly or quarterly triage of low-confidence classifications. Feeds triage decisions back into the rule base.", "[To be designated]"),
        ("Taxonomy Reviewer", "Annual review of the 17 Business Categories for continued fit with reporting needs.", "DPS Leadership"),
    ]
    _add_simple_table(
        doc,
        headers=["Role", "Responsibility", "Current Holder"],
        rows=roles,
        col_widths_cm=[3.8, 9.5, 3.7],
    )

    # ---- 3.2 Annual Taxonomy Review ----
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_h2(doc, "3.2  Annual Taxonomy Review")
    add_body(doc, "Once per year, the Taxonomy Reviewer (DPS leadership) reviews the 17 Business Categories.")
    add_bullet(doc, "Pull the current year's classified output (NIGP_Mapping_JHK3.csv).", bold_lead="Step 1.")
    add_bullet(doc, "Review the Business Category distribution. Are any categories too large, too small, or no longer relevant?", bold_lead="Step 2.")
    add_bullet(doc, "Review any new commodity types that appeared during the year. Do they fit existing categories or do they require a new category?", bold_lead="Step 3.")
    add_bullet(doc, "Decision: confirm the current 17 categories, or propose additions, merges, or splits.", bold_lead="Step 4.")
    add_bullet(doc, "If changes are made: update business_categories_JHK3.csv and business_categories_summary_JHK3.csv. Re-run the classifier in batch mode to apply.", bold_lead="Step 5.")
    add_bullet(doc, "Document the review decision and date. Archive the prior categories file before overwriting.", bold_lead="Step 6.")

    # ---- 3.3 Accuracy Audit ----
    add_h2(doc, "3.3  Annual Classification Accuracy Audit")
    add_body(doc, "Once per year, conduct an independent accuracy check to verify the classifier is still producing correct results.")
    add_bullet(doc, "Pull 100 classified rows at random from the most recent batch output.", bold_lead="Step 1.")
    add_bullet(doc, "Have a procurement analyst who was NOT involved in building the classifier review each row independently.", bold_lead="Step 2.")
    add_bullet(doc, "For each row, the analyst records:  (a) Agree with the Business Category assignment? Yes or No.  (b) If No, the proposed correction.", bold_lead="Step 3.")
    add_bullet(doc, "Calculate the agreement rate. The target is 90% or higher.", bold_lead="Step 4.")
    add_bullet(doc, "Any systematic disagreements (for example, the same category mismatch on multiple rows) become candidates for new or corrected rules.", bold_lead="Step 5.")
    add_bullet(doc, "Document the audit results, the agreement rate, and any corrective actions taken.", bold_lead="Step 6.")
    add_bullet(doc, "Track the agreement rate year over year. It should improve as the rule base matures.", bold_lead="Step 7.")

    # ---- 3.4 NIGP Catalog Expansion ----
    add_h2(doc, "3.4  NIGP Catalog Expansion")
    add_body(
        doc,
        "The current working catalog contains 138 NIGP 3-digit Classes and 470 NIGP 5-digit "
        "Class-Items derived from the historical data. The full NIGP standard contains "
        "approximately 9,000 codes.",
    )
    add_h2(doc, "Recommendation")
    add_body(doc, "Consider procuring a full NIGP catalog license from Periscope Holdings. This would:")
    add_bullet(doc, "Enable more specific 5-digit classifications where current descriptions support them.")
    add_bullet(doc, "Support benchmarking against the full national standard used by peer public-sector procurement entities.")
    add_bullet(doc, "Future-proof the taxonomy for commodity types not yet represented in Chicago's historical data.")

    # ---- 3.5 Succession ----
    add_h2(doc, "3.5  Succession and Knowledge Transfer")
    add_body(doc, "If the System Owner role transfers to a new individual, follow this procedure to ensure continuity.")
    add_bullet(doc, "The incoming owner reads this SOP in full.", bold_lead="Step 1.")
    add_bullet(doc, "The incoming owner is granted access to the GitHub repository and the Streamlit web app.", bold_lead="Step 2.")
    add_bullet(doc, "The outgoing owner walks through one complete batch-processing cycle with the incoming owner, demonstrating each step in Section 2.1.", bold_lead="Step 3.")
    add_bullet(doc, "The incoming owner independently processes one data extract under the outgoing owner's observation.", bold_lead="Step 4.")
    add_bullet(doc, "The outgoing owner transfers all credentials and access. The web app password is rotated as part of the handoff.", bold_lead="Step 5.")
    add_bullet(doc, "Update the Roles and Responsibilities table in Section 3.1 to reflect the new owner.", bold_lead="Step 6.")

    # ---- 3.6 Document Control ----
    add_h2(doc, "3.6  Document Control")
    doc_control = [
        ("This SOP", "Annually or upon system change", "System Owner"),
        ("Methodology document (NIGP_Methodology_for_Leadership_JHK3.docx)", "Upon methodology change", "System Owner"),
        ("Executive Brief (NIGP_Executive_Brief_JHK3.docx)", "Upon significant system update", "System Owner"),
        ("Rule file (keyword_rules_DRAFT_JHK3.csv)", "Ongoing as rules are added or retired", "Rule File Maintainer"),
        ("Business Categories reference files", "Annually at taxonomy review", "Taxonomy Reviewer"),
    ]
    _add_simple_table(
        doc,
        headers=["Document", "Review Frequency", "Owner"],
        rows=doc_control,
        col_widths_cm=[8.0, 5.5, 3.5],
    )

    # -------------------------------------------------------------------
    # End-of-document footer
    # -------------------------------------------------------------------
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("End of Standard Operating Procedure.")
    run.font.italic = True
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT_PATH)
    return OUT_PATH


def main():
    out = build_sop()
    print(f"Built: {out}")


if __name__ == "__main__":
    main()
