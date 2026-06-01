"""Build the personal-toolbox case study / capability reference for the NIGP project.

Generates one file in /workspaces/Procurement-ai-tools/outputs/:
  NIGP_Case_Study_JHK3.docx  — 2-page case study + capability reference (STAR format)

PURPOSE (2026-06-01): this is James's PERSONAL PROFESSIONAL TOOLBOX asset — a
durable capability reference he carries forward as he leaves this position but
stays with the City of Chicago, moving into contract administration & supply
chain. It is INTERNAL-CONTEXT: it names the City of Chicago and frames the work
toward contract administration / supply chain, in the same class as the
Executive Brief, Methodology, and SOP (internal `outputs/` artifacts that carry
City / DPS context). Per locked decision #3, keep this OUT of the public repo —
it is not a public-facing surface like streamlit_app.py / README / HOW_TO_USE.

Styling helpers are imported from build_leadership_deliverables_JHK3 so the
document matches the look of the Executive Brief, Methodology, and SOP.
Idempotent — re-run to regenerate.
"""
import sys
from pathlib import Path

# Allow `from build_leadership_deliverables_JHK3 import ...` when run from anywhere.
sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from build_leadership_deliverables_JHK3 import (
    GRAY_RGB,
    CHI_LT_BLUE_HEX,
    add_body,
    add_bullet,
    add_callout_box,
    add_h1,
    add_h2,
    add_title_banner,
    set_cell_shading,
    set_doc_margins,
    set_table_borders,
    style_table_header,
)

OUT_DIR = Path("/workspaces/Procurement-ai-tools/outputs")

LIVE_APP_URL = "https://chicago-nigp-classifier.streamlit.app/"
REPO_URL = "https://github.com/jhkirby2001-collab/Procurement-ai-tools"

# Headline facts (locked production metrics, post-resolver 2026-05-14).
ROWS_TOTAL = "784,556"
COVERAGE_ROWS = [
    ("Keyword rule (246 hand-curated + 6,766 AI-mined)", "688,044", "87.7%"),
    ("AI-assist fallback (saved one-time AI output, no runtime API)", "93,518", "11.9%"),
    ("Account-code pattern (subgrant / pass-through disbursements)", "2,028", "0.3%"),
    ("Unclassified — no usable description (explicitly tagged)", "966", "0.1%"),
    ("TOTAL — every record mapped", "784,556", "100%"),
]

SKILLS_ROWS = [
    ("Procurement domain",
     "NIGP commodity-code alignment; a 17-category MECE spend taxonomy; category "
     "management; identification of subgrant / pass-through disbursements vs. commodity "
     "buys; audit-defensible classification."),
    ("Data & AI engineering",
     "Python and pandas at 784K-row scale; LLM-assisted pattern mining (Anthropic "
     "Claude); JSON-schema-bounded prompts with closed enumerations; a deterministic, "
     "three-tier rules-engine classification pipeline."),
    ("Product & delivery",
     "A deployed Streamlit web application; externalized CSV rule files maintainable by "
     "non-technical staff; leadership-ready Word / Excel deliverables; a full "
     "methodology document and per-record audit trail."),
    ("Judgment & governance",
     "One-time AI cost held to ~$27; a rules-only, fully auditable production runtime "
     "with no black-box dependency; provenance metadata on every classification "
     "decision."),
]


def build_case_study():
    doc = Document()
    set_doc_margins(doc, cm=1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title_banner(
        doc,
        title="NIGP-Sourced Procurement Category Mapper",
        subtitle="Project Case Study & Capability Reference — for Contract Administration & Supply Chain",
        author_line=(
            "City of Chicago  •  Department of Procurement Services  •  "
            "James H. Kirby III, CSCP, MS-SCM  •  Personal capability reference, June 2026"
        ),
    )

    add_callout_box(
        doc,
        label="AT A GLANCE",
        body_text=(
            f"{ROWS_TOTAL} City procurement records classified at 100% coverage  •  "
            "0 records left in a manual-review backlog  •  "
            "a 17-category MECE taxonomy aligned to the public NIGP standard  •  "
            "AI used once (~$27 total); the production runtime is rules-only and "
            "fully auditable  •  delivered as a live, deployed web tool."
        ),
    )

    # ---- Challenge ----
    add_h1(doc, "The Challenge")
    add_body(
        doc,
        "The City of Chicago, like most public agencies, did not maintain its own "
        "commodity classification of what it buys. Without one, the most basic strategic "
        "questions go unanswered: where does the spend actually go, which commodities are "
        "bought across multiple departments without coordination, and can any single "
        "purchase be traced and defended in an audit?",
    )
    add_body(
        doc,
        "The raw material was a 784,556-row City of Chicago procurement transaction "
        "history — a consultant-delivered extract covering AP activity years 2017, 2020, "
        "2021, and 2023, with messy descriptions, inconsistent formatting, and no owned "
        "taxonomy. The goal: turn it into a clean, auditable, industry-standard category "
        "structure the City could own and reuse.",
    )

    # ---- What I built ----
    add_h1(doc, "What I Built")
    add_bullet(
        doc,
        "Every record is assigned a leadership-friendly Business Category, a 3-digit NIGP "
        "Class, and — where the description supports it — a 5-digit NIGP Item. 17 "
        "categories → 138 classes → 470 items.",
        bold_lead="A three-level taxonomy.",
    )
    add_bullet(
        doc,
        "A dual-mode classification engine: one core function that runs both a full "
        "historical batch and a single paste-a-description lookup for live work.",
        bold_lead="A reusable engine.",
    )
    add_bullet(
        doc,
        "A deployed, password-gated web application where staff can classify a "
        "description, bulk-classify a file, browse the taxonomy logic, and look up the "
        "exact rule behind any decision.",
        bold_lead="A live web tool.",
    )
    add_bullet(
        doc,
        "All classification rules live in plain CSV files, so a non-technical analyst can "
        "add or retire a rule without touching code.",
        bold_lead="Staff-editable rules.",
    )

    # ---- Approach ----
    add_h1(doc, "My Approach — A Defensible AI Architecture")
    add_body(
        doc,
        "The defining design choice was how to use AI without sacrificing auditability. "
        "AI was used exactly once, at build time, to mine recurring patterns from the "
        "long tail of descriptions no hand-written rule covered. Its output was harvested "
        "into frozen rules and the model was never called again. The production "
        "classifier that does the day-to-day work is 100% deterministic rules — no API "
        "key, no internet dependency, no per-classification charge.",
    )
    add_bullet(
        doc,
        "17 mutually exclusive, collectively exhaustive Business Categories built around "
        "what the City buys — not how much it spends — so the structure stays stable as "
        "budgets move.",
        bold_lead="MECE category design.",
    )
    add_bullet(
        doc,
        "The classifier reads description text plus FMPS account/object/fund codes only. "
        "Vendor name and consultant-supplied NIGP codes are deliberately excluded — the "
        "classification is independently derived and defensible.",
        bold_lead="Disciplined inputs.",
    )
    add_bullet(
        doc,
        "Every classified record carries the exact rule that fired, a confidence level, "
        "and a machine-readable reason. Any single decision can be reconstructed "
        "end-to-end.",
        bold_lead="Full audit trail.",
    )
    add_callout_box(
        doc,
        label="WHY THIS MATTERS",
        body_text=(
            "Most \"AI did it\" projects can't tell you why a given row was classified the "
            "way it was. This one can — for all 784,556 of them. AI bought coverage of the "
            "long tail once; deterministic rules keep the running system transparent, "
            "cheap, and reviewable."
        ),
    )

    # ---- Results ----
    add_h1(doc, "The Results")
    add_body(
        doc,
        f"Run end-to-end against all {ROWS_TOTAL} records, the engine maps 100% of rows "
        "with zero left in a human-review queue. Coverage breaks down by method as "
        "follows:",
    )
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "Classification method"
    hdr[1].text = "Records"
    hdr[2].text = "Share"
    style_table_header(table.rows[0])
    for i, (method, rows, share) in enumerate(COVERAGE_ROWS):
        r = table.add_row().cells
        r[0].text = method
        r[1].text = rows
        r[2].text = share
        for cell in r:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        if method.startswith("TOTAL"):
            for cell in r:
                set_cell_shading(cell, CHI_LT_BLUE_HEX)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
        elif i % 2 == 0:
            for cell in r:
                set_cell_shading(cell, CHI_LT_BLUE_HEX)
    for row in table.rows:
        row.cells[0].width = Cm(11.0)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(2.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ---- BRIDGE: Contract Administration & Supply Chain ----
    add_h1(doc, "How This Transfers to Contract Administration & Supply Chain")
    add_body(
        doc,
        "The taxonomy and the classification method are not single-use. They are reusable "
        "infrastructure for the functions I am moving into — and the same engine can be "
        "re-pointed at a contract portfolio or a new spend extract without re-running AI "
        "or engaging a consultant.",
    )
    add_h2(doc, "In contract administration")
    add_bullet(
        doc,
        "Apply the same 17-category → NIGP-class structure to active contracts so every "
        "contract carries a defensible commodity classification.",
        bold_lead="Commodity-code the contract portfolio.",
    )
    add_bullet(
        doc,
        "Surface where the same commodity is bought under multiple contracts and vendors "
        "— the candidates for consolidation, bundling, or cooperative purchasing.",
        bold_lead="Detect overlap and consolidation opportunities.",
    )
    add_bullet(
        doc,
        "The per-record audit trail — rule that fired, confidence, reason — is the same "
        "evidence model a clean contract file needs for review and compliance.",
        bold_lead="Audit-ready contract records.",
    )
    add_bullet(
        doc,
        "Group contracts by commodity class to plan sourcing waves and avoid fragmented, "
        "one-off renewals.",
        bold_lead="Renewal and category planning.",
    )
    add_h2(doc, "In supply chain & category management")
    add_bullet(
        doc,
        "A commodity taxonomy is the backbone of category management — spend visibility by "
        "commodity is the prerequisite for any strategic sourcing program.",
        bold_lead="Category management foundation.",
    )
    add_bullet(
        doc,
        "Identify the same commodity bought across departments, aggregate the volume, and "
        "leverage scale.",
        bold_lead="Demand aggregation.",
    )
    add_bullet(
        doc,
        "A commodity-level view of who supplies what supports supplier rationalization and "
        "reduces fragmented tail spend.",
        bold_lead="Supplier rationalization.",
    )
    add_bullet(
        doc,
        "The rules-engine + one-time-AI method re-points at any new spend or contract "
        "dataset — a repeatable capability, not a one-time deliverable.",
        bold_lead="Repeatable method.",
    )

    # ---- Capabilities table ----
    add_h1(doc, "Capabilities This Project Adds to My Professional Toolbox")
    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "Dimension"
    hdr[1].text = "What this project demonstrates"
    style_table_header(table.rows[0])
    for i, (dim, detail) in enumerate(SKILLS_ROWS):
        r = table.add_row().cells
        r[0].text = dim
        r[1].text = detail
        for cell in r:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        for run in r[0].paragraphs[0].runs:
            run.font.bold = True
        if i % 2 == 0:
            for cell in r:
                set_cell_shading(cell, CHI_LT_BLUE_HEX)
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(12.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ---- See it ----
    add_h1(doc, "See It / Explore the Work")
    add_bullet(doc, LIVE_APP_URL + " (password-gated).", bold_lead="Live tool:")
    add_bullet(doc, REPO_URL + " (code, rule files, and methodology).", bold_lead="Source:")

    # ---- Footer ----
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "James H. Kirby III, CSCP, MS-SCM  —  City of Chicago, Department of Procurement "
        "Services.  Built with Python, pandas, the Anthropic Claude API (one-time), and "
        "Streamlit."
    )
    run.font.italic = True
    run.font.color.rgb = GRAY_RGB
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out = OUT_DIR / "NIGP_Case_Study_JHK3.docx"
    doc.save(out)
    return out


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = build_case_study()
    print(f"Built: {out}")


if __name__ == "__main__":
    main()
