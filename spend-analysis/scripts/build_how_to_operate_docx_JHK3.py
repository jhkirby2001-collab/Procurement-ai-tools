"""Build the leadership-facing "How to Operate" Word document for the
NIGP-Sourced Procurement Category Mapper.

Generates:
  outputs/HOW_TO_USE_NIGP_Mapping_JHK3.docx

This is the Word companion to outputs/HOW_TO_USE_NIGP_Mapping_JHK3.md — the
detailed build + operate + trust guide written for leadership and staff.

Style/palette are reused from build_leadership_deliverables_JHK3.py so this
matches the other leadership artifacts. Content is DE-BRANDED (independent
tool framing) to match the public-facing Markdown guide — no agency branding
in the author line or body.

Portable: the output directory is derived from this file's location, so the
script runs both in a codespace and in a fresh clone.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

# Reuse helpers + brand palette from the leadership-deliverables script.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from build_leadership_deliverables_JHK3 import (  # noqa: E402
    CHI_LT_BLUE_HEX,
    add_body,
    add_bullet,
    add_callout_box,
    add_h1,
    add_h2,
    add_title_banner,
    set_cell_shading,
    set_doc_margins,
    set_table_borders,
    style_table_header,
)

# Portable output path: repo_root/outputs
REPO_ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = REPO_ROOT / "outputs"
OUT_PATH = OUT_DIR / "HOW_TO_USE_NIGP_Mapping_JHK3.docx"

DOC_VERSION = "Version 2026-08-07  •  Expanded operate-and-trust edition"
AUTHOR_LINE = (
    "Independent public-sector procurement tool  •  "
    "Prepared by James H. Kirby III, CSCP, MS-SCM  •  " + DOC_VERSION
)


def _shade_alt_rows(table, color_hex=CHI_LT_BLUE_HEX):
    for i, row in enumerate(table.rows[1:]):
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, color_hex)
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)


def _table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    style_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    _shade_alt_rows(table)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    return table


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_doc_margins(doc, cm=1.8)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ------------------------------------------------------------------ Banner
    add_title_banner(
        doc,
        title="How the Mapper Works — and How to Operate It",
        subtitle="NIGP-Sourced Procurement Category Mapper",
        author_line=AUTHOR_LINE,
    )
    add_body(
        doc,
        "Who this is for: leadership who need to understand and stand behind the tool's "
        "output, and procurement staff who operate it day to day.",
    )

    # ------------------------------------------------------------ 0. Summary
    add_callout_box(
        doc,
        label="THE ONE-PARAGRAPH SUMMARY",
        body_text=(
            "This tool takes a plain-English procurement description — \"elevator maintenance,\" "
            "\"bulk rock salt,\" \"chain link fence repair\" — and assigns a standardized "
            "commodity classification: a business-friendly Business Category, an industry-standard "
            "NIGP 3-digit Class, and (where the text is specific enough) a 5-digit NIGP Item. It "
            "was validated on 784,556 historical purchase-order and invoice-line records and maps "
            "100% of them. Every answer carries the exact rule that produced it, a confidence "
            "level, and a plain-language reason. No AI model is called when you use it — AI was "
            "used once, at build time, to help write rules, then frozen. That is what makes the "
            "output defensible in an audit."
        ),
    )

    # ------------------------------------------------------------ 1. Problem
    add_h1(doc, "1.  What problem this solves")
    add_body(
        doc,
        "Most public agencies never build their own classification of what they buy. Without one, "
        "the basic strategic questions can't be answered with confidence: where spend actually "
        "goes by commodity, what is bought across departments without coordination, and whether "
        "any single purchase can be traced, categorized, and defended to an auditor.",
    )
    add_body(
        doc,
        "This mapper is that missing classification layer. It is independent — it does not rely on "
        "any prior consultant's labels — and it is reusable: it re-points at any future spend "
        "extract or contract portfolio without re-running AI or hiring anyone.",
    )

    # ------------------------------------------------------------ 2. Taxonomy
    add_h1(doc, "2.  The taxonomy — three levels on every row")
    add_body(
        doc,
        "Every classified record carries three levels of classification, serving two audiences at "
        "once: leadership reads the top level; auditors and sourcing analysts use the NIGP codes "
        "underneath.",
    )
    add_bullet(doc, "Level 1 — Business Category (17 buckets — the plain-English answer)", bold_lead="Level 1")
    add_bullet(doc, "Level 2 — NIGP 3-digit Class (industry-standard commodity class)", bold_lead="Level 2")
    add_bullet(doc, "Level 3 — NIGP 5-digit Item (specific item, when the text supports it)", bold_lead="Level 3")

    add_h2(doc, "The 17 Business Categories")
    _table(
        doc,
        ["#", "Business Category", "Typical contents"],
        [
            ["1", "Facilities Operations & Maintenance", "HVAC, plumbing, lighting, doors, electrical supplies"],
            ["2", "Public Safety, Uniforms & PPE", "Uniforms, fire equipment, gloves, masks, body armor, ammunition"],
            ["3", "Construction Materials", "Concrete, asphalt, steel, pipe, rebar, pavement marking"],
            ["4", "Vehicles & Fleet", "Vehicles, fleet parts, repairs, tires, fuel"],
            ["5", "Professional & Administrative Services", "Management, security guards, translation, legal/records"],
            ["6", "Office, Print & Marketing", "Office supplies, printing, advertising, signage"],
            ["7", "Equipment Rental & Leasing", "Equipment rentals, copier leases, traffic-control rentals"],
            ["8", "Janitorial, Sanitation & Waste", "Cleaning supplies, paper products, carpet, dumpsters, disposal"],
            ["9", "Heavy Equipment & Machinery", "Loaders, sweepers, lifts, forklifts, sprayers (purchased)"],
            ["10", "IT, Telecom & Audio/Visual", "Network services, telephony, A/V, printers, toner"],
            ["11", "Landscaping, Grounds & Irrigation", "Landscape services, irrigation, fencing, window washing, snow removal"],
            ["12", "Chemicals & Water Treatment", "Industrial chemicals, water treatment, deicers / road salt"],
            ["13", "Medical & Health Services", "Medical gases, defibrillators/AEDs, exam services"],
            ["14", "Animal Care & Veterinary", "Veterinary supplies, animal feed, beekeeping"],
            ["15", "Furniture & Furnishings", "Office furniture, mattresses, fitness equipment"],
            ["16", "Construction & Trades Services", "Bridge construction, expansion-joint repair, tradesmen"],
            ["17", "Grants & Pass-Through Funding", "Subgrant disbursements (not commodity purchases)"],
        ],
        col_widths_cm=[0.9, 5.6, 9.5],
    )
    add_callout_box(
        doc,
        label="WHY CATEGORY 17 IS SEPARATE",
        body_text=(
            "Subgrants are financial transfers to subrecipients, not commodity buys. They have "
            "different governance and audit treatment, so mixing them into \"professional services\" "
            "would obscure real structure. This category has no NIGP class because NIGP is a "
            "commodity framework with no \"subgrants\" code."
        ),
    )

    # ------------------------------------------------------------ 3. How built
    add_h1(doc, "3.  How the mapper was built — the part that earns your confidence")

    add_h2(doc, "3.1  The source data")
    add_body(
        doc,
        "The taxonomy and rules were derived from a 784,556-row public-sector procurement extract "
        "(87 columns, AP activity years 2017, 2020, 2021, 2023). Every record was classified "
        "independently from its description text — not re-labeled from any prior categorization. "
        "The source file carried partial prior commodity codes on ~30% of rows; the mapper does "
        "NOT use those as inputs. They are preserved only for optional after-the-fact cross-checks.",
    )

    add_h2(doc, "3.2  What the classifier looks at — and what it does not")
    add_bullet(doc, "Description text — reads up to four description fields and picks the first substantive one.", bold_lead="Uses:")
    add_bullet(doc, "Account / object / fund codes — the agency's own financial codes, as a backup signal.", bold_lead="Uses:")
    add_bullet(doc, "Vendor name — the same vendor sells across many categories; including it would introduce error.", bold_lead="Excluded:")
    add_bullet(doc, "Spend dollars — the tool answers what was bought, not how much was spent.", bold_lead="Excluded:")
    add_bullet(doc, "Prior consultant codes — building an owned classification means not inheriting them.", bold_lead="Excluded:")

    add_h2(doc, "3.3  The rule base — how classifications are decided")
    add_body(
        doc,
        "Classification runs on a library of rules stored as plain CSV files that procurement staff "
        "can read and edit without touching code. As of this document:",
    )
    _table(
        doc,
        ["Rule source", "Count", "What it is"],
        [
            ["Hand-curated keyword rules", "280", "Written and owned by the author. Highest-trust, high-volume rules."],
            ["AI-mined keyword rules", "6,766", "From a one-time AI pattern-mining pass. Full provenance. Frozen."],
            ["Account-code patterns", "6", "Map specific financial account codes to a category (subgrant accounts)."],
            ["Total rules", "7,046", "Plus one resolver that consumes saved AI output (see 3.4)."],
        ],
        col_widths_cm=[5.0, 2.2, 8.8],
    )
    add_body(
        doc,
        "Each keyword rule says: if the description matches this pattern, assign this Business "
        "Category and NIGP class. Matching supports exact (whole string), starts_with (prefix), "
        "and contains (substring), evaluated in that priority order. The first match wins.",
    )

    add_h2(doc, "3.4  The one and only use of AI — bounded, one-time, and frozen")
    add_callout_box(
        doc,
        label="THE HEADLINE",
        body_text=(
            "The production tool is rules-only. No AI model is called when you classify anything. "
            "AI was used once, at build time, to read the long tail of ~30,000 descriptions no "
            "hand-written rule covered and propose patterns. Total one-time cost was about $27."
        ),
    )
    add_body(doc, "Why this AI use is defensible, in six points:")
    add_bullet(doc, "One-time, not ongoing. Runs never call AI — no API key, no internet, no per-classification charge.", bold_lead="1.")
    add_bullet(doc, "Bounded. The model could only choose from the 17 approved categories and 138 NIGP classes — it cannot invent a code.", bold_lead="2.")
    add_bullet(doc, "Transparent. Every AI-mined rule records the model, confidence, source frequency, and reasoning.", bold_lead="3.")
    add_bullet(doc, "Reviewable & editable. Staff can demote or delete any AI-mined rule; it has no special status.", bold_lead="4.")
    add_bullet(doc, "Replaceable. Delete all AI content and the tool still runs on curated rules + account patterns. Coverage drops; nothing breaks.", bold_lead="5.")
    add_bullet(doc, "Confidence-labeled, never laundered. AI-derived answers are tagged AI-high / AI-medium / AI-low so you can filter or exclude them.", bold_lead="6.")

    # ------------------------------------------------------------ 4. Tiers
    add_h1(doc, "4.  How a row gets classified — the four tiers")
    add_body(
        doc,
        "Every record runs through four tiers in order. The first tier that matches wins, and the "
        "row records which tier produced its answer — that is what creates a clean, per-row audit "
        "trail.",
    )
    _table(
        doc,
        ["Tier", "Method label", "What it means"],
        [
            ["1", "keyword_rule", "Matched a curated or AI-mined keyword rule. Strongest evidence."],
            ["2", "account_pattern", "Description too thin; a financial account code resolved it."],
            ["3", "ai_assist", "No rule matched; saved one-time AI output supplied the answer. No new API call."],
            ["4", "no_description", "No usable text anywhere; honestly tagged rather than guessed."],
        ],
        col_widths_cm=[1.2, 3.4, 11.4],
    )
    add_h2(doc, "Coverage from the last full production run (2026-05-14)")
    _table(
        doc,
        ["Tier", "Rows", "Share"],
        [
            ["Tier 1 — Keyword rule", "688,044", "87.7%"],
            ["Tier 2 — Account-code pattern", "2,028", "0.3%"],
            ["Tier 3 — AI-assist resolver (saved output)", "93,518", "11.9%"],
            ["Tier 4 — Unclassified — No Description", "966", "0.1%"],
            ["Total mapped to a Business Category", "784,556", "100.0%"],
            ["Terminal human-review queue", "0", "0.0%"],
        ],
        col_widths_cm=[8.6, 3.6, 3.8],
    )
    add_callout_box(
        doc,
        label="NOTE ON FRESHNESS",
        body_text=(
            "34 curated rules were added on 2026-08-07 (see Section 6). Those improve the "
            "interactive tool immediately, but the batch percentages above only change on the next "
            "full batch refresh. The \"100% mapped\" headline is unaffected — the new rules mainly "
            "shift a small number of rows from Tier 3 (AI-assist) to Tier 1 (keyword rule), which "
            "strengthens the evidence behind those rows. Note also: \"0 review queue\" is not the "
            "same as \"0 error\" — Tier 3 rows rest on an AI proposal and are labeled so, so you can "
            "always filter to the strongest-evidence tiers when accuracy matters most."
        ),
    )

    # ------------------------------------------------------------ 5. Operate
    add_h1(doc, "5.  How to operate it — the web app (primary) and CLI (advanced)")
    add_body(doc, "There are two ways to use the tool. Most people should use the web app.")
    add_h2(doc, "5A.  The web app — eight pages")
    _table(
        doc,
        ["Page", "What it does"],
        [
            ["Classify", "Paste one description; get category, NIGP class, confidence, and the rule that fired."],
            ["Bulk Classify", "Upload a spreadsheet of descriptions; download it back with a category on every row."],
            ["Spend Report", "Upload a spend file; it classifies, then builds a spend analysis (spend-by-category, Pareto 80/20, top vendors, consolidation) + downloadable Excel."],
            ["Spend Report Methodology", "How the spend analysis works, how to run it, how to read each report, and its limits."],
            ["Procurement Taxonomy Logic", "Visual walkthrough of the three levels and the four tiers."],
            ["Methodology", "The full build narrative and defensibility argument."],
            ["Business Categories", "The 17 categories and what rolls into each."],
            ["Rule Lookup", "Search all 7,046 rules to see exactly how any phrase routes and why."],
        ],
        col_widths_cm=[4.6, 12.2],
    )
    add_body(
        doc,
        "The address and password are maintained by the project author — request access from James. "
        "Everyday workflows: use Classify to answer \"what category is this?\"; use Rule Lookup to "
        "answer \"why did it decide that?\"",
    )
    add_h2(doc, "5B.  The Spend Report — adaptive spend analysis from any file")
    add_body(
        doc,
        "The Spend Report page takes the tool one step past classification: upload a spend file "
        "in any format and it builds a full spend analysis you can download as one Excel "
        "workbook. It is adaptive — it inspects every column (name and content), figures out "
        "which is the description, amount, vendor, department, date, and any breakdown "
        "dimensions, then runs every analysis the data supports and skips the ones it can't, "
        "telling you why. You do not have to shape your file to fit the tool. Steps: (1) upload "
        "a CSV or Excel file; (2) confirm the detected columns (only change a box if a guess is "
        "wrong; only Description is required); (3) click Generate, then download the workbook.",
    )
    add_body(
        doc,
        "Depending on what the file contains, it produces: a data-driven executive summary "
        "(findings + recommended steps), spend by category, Pareto 80/20, top vendors, vendor "
        "concentration/risk (HHI), tail-spend, single- vs multi-source, spend trend over time, "
        "spend by department, a category-by-department matrix, vendor consolidation/"
        "fragmentation, and a spend breakdown for any dimension it finds (contract vs "
        "non-contract, supplier diversity, status). Each becomes a tab in a brand-colored Excel "
        "workbook with charts. All deterministic arithmetic — no AI, and the file stays in the "
        "session. Classification is keyword-rules-only, so coverage depends on description "
        "quality; the coverage note always shows how much of the file was classified. The "
        "in-app Spend Report Methodology page documents each report in full.",
    )
    add_h2(doc, "5C.  The command line (advanced / technical users)")
    add_body(
        doc,
        "From a terminal in the project environment, classify a single description with:  "
        "python spend-analysis/scripts/classifier_JHK3.py --describe \"PASTE DESCRIPTION HERE\". "
        "It prints the Business Category, NIGP code, confidence, and the rule that fired.",
    )

    # ------------------------------------------------------------ 6. Fix
    add_h1(doc, "6.  How to fix or improve a classification — permanently")
    add_body(
        doc,
        "This is the maintenance loop, and it is deliberately simple: teach it once, and every "
        "future similar description is right. To add a rule, open "
        "keyword_rules_DRAFT_JHK3.csv in Excel and add a row with: the pattern text to match; "
        "the match type (contains / starts_with / exact); the correct Business Category (one of the "
        "17, spelled exactly); the correct NIGP class (verified against the reference table — never "
        "guessed); an optional 5-digit item; a match level (broad for normal cases); and a short "
        "note. Save — it takes effect immediately in the web app and CLI.",
    )
    add_callout_box(
        doc,
        label="WORKED EXAMPLE — THE \"FENCE\" FIX (2026-08-07)",
        body_text=(
            "A user typed \"fence\" and got no result. Investigation showed there was no rule for "
            "the bare word, and the one related rule pointed at the wrong category. The actual data "
            "settled it: every fence row was a landscape salt-fence service, and the NIGP item for "
            "chain-link fence/guardrail repair sits under class 988 — Landscaping, Grounds & "
            "Irrigation. Three rules were added/aligned to that verified home. \"fence,\" "
            "\"fencing,\" and \"chain link fence repair\" now all classify correctly. That is the "
            "whole maintenance loop: find the gap, verify the correct NIGP home against the "
            "reference table, and close it with a one-line rule."
        ),
    )

    # ------------------------------------------------------------ 7. Recent
    add_h1(doc, "7.  Recent improvement — the interactive robustness pass (2026-08-07)")
    add_body(
        doc,
        "The \"fence\" report surfaced a broader pattern: the interactive tool only answers when "
        "typed words match a rule, and many short common terms had no rule (even though the full "
        "784K batch was already 100% mapped, because real descriptions are long and detailed). A "
        "45-word test of everyday procurement terms found only 17 of 45 (38%) returned a result.",
    )
    add_body(
        doc,
        "In response, 34 curated rules were added — each mapped to its correct NIGP class verified "
        "against the reference table, deliberately NOT auto-matched (which would produce "
        "confident-but-wrong answers). Coverage on those common terms rose to 35 of 45 (78%), and "
        "the plain-English regression test moved to a clean 72/72 with zero mis-categorizations. "
        "Runtime stayed rules-only, no API key. Newly covered terms include gloves, uniforms, "
        "ammunition, PPE/masks, paper towels, carpet, road salt/deicer, translation, security "
        "guards, defibrillator, batteries, electrical wire, forklift, toner, window cleaning, snow "
        "removal, fuel, and pavement marking.",
    )
    add_body(
        doc,
        "Honest limit: a handful of terms (e.g. hand sanitizer, traffic signals, surveying, "
        "architecture services, scaffolding) still have no home, because the working NIGP "
        "reference contains only the ~470 codes that actually appeared in the source data. "
        "Covering those would require importing the full official NIGP dictionary — a separate, "
        "deliberate decision.",
    )

    # ------------------------------------------------------------ 8. Reading
    add_h1(doc, "8.  Reading a result — what every field means")
    _table(
        doc,
        ["Field", "What it tells you"],
        [
            ["Business_Category", "The main answer — one of the 17 buckets."],
            ["NIGP_Class_3digit", "The industry-standard 3-digit commodity class."],
            ["NIGP_Item_5digit", "The specific 5-digit item, when the text supports it (else blank)."],
            ["NIGP_Match_Level", "exact = specific item • broad = 3-digit class only • review = flagged."],
            ["Classification_Confidence", "High / Medium / Low (or AI-high/medium/low for Tier 3)."],
            ["Classification_Method", "Which of the four tiers produced the answer."],
            ["Classification_Reason", "The exact rule that fired — the audit trail."],
        ],
        col_widths_cm=[5.0, 11.8],
    )

    # ------------------------------------------------------------ 9. Limits
    add_h1(doc, "9.  Limitations — stated plainly")
    add_bullet(doc, "Single-source origin. The taxonomy came from one dataset; new extracts may reveal commodity types not represented yet.")
    add_bullet(doc, "Description-quality dependent. Thin descriptions are resolved at Tier 3 with a low-confidence tag; empty ones are tagged, not guessed.")
    add_bullet(doc, "Partial NIGP working set. The ~138 classes / ~470 items in use are a subset of the full ~9,000-code NIGP catalog.")
    add_bullet(doc, "No vendor or dollar signal, by design. The answer reflects what was bought, not who sold it or how much it cost.")

    # ------------------------------------------------------------ 10. Governance
    add_h1(doc, "10.  Keeping it trustworthy over time — governance")
    add_bullet(doc, "Rule-file ownership. One named owner for the curated rules; add rules as new commodities appear. The base has grown 148 to 246 to 280.", bold_lead="1.")
    add_bullet(doc, "AI-tier audit cadence. Quarterly, sample AI-medium/AI-low rows; where the AI got a recurring pattern right, promote it into a curated rule.", bold_lead="2.")
    add_bullet(doc, "Annual taxonomy review. Confirm the 17 categories still fit reporting needs.", bold_lead="3.")
    add_bullet(doc, "Blind audit benchmark. Once a year, independently classify 100 random rows blind and compare; track the agreement rate by tier.", bold_lead="4.")

    add_body(
        doc,
        "Questions direct to James H. Kirby III, CSCP, MS-SCM, project author. The full methodology "
        "and audit trail are in METHODOLOGY_JHK3.md, and every classification decision can be "
        "traced to the exact rule that produced it.",
    )

    doc.save(str(OUT_PATH))
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
