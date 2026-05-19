"""Build leadership-facing deliverables for the NIGP project.

Generates three files in /workspaces/Procurement-ai-tools/outputs/:
  1. NIGP_Summary_for_Leadership_JHK3.xlsx   — single-tab Excel summary
  2. NIGP_Executive_Brief_JHK3.docx          — 2-page Word brief
  3. NIGP_Methodology_for_Leadership_JHK3.docx — leadership-friendly methodology

City of Chicago brand colors used throughout:
  - Chicago Light Blue (flag stripes / Lake Michigan): #41B6E6
  - Chicago Red (the four stars):                       #DA291C
  - Chicago Navy (deep accent):                         #002F6C
"""
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path("/workspaces/Procurement-ai-tools/outputs")

# -- Brand palette ---------------------------------------------------------
CHI_BLUE_HEX = "41B6E6"
CHI_RED_HEX = "DA291C"
CHI_NAVY_HEX = "002F6C"
CHI_LT_BLUE_HEX = "D6EEF9"   # 15% tint for table banding / callout boxes
WHITE_HEX = "FFFFFF"
BLACK_HEX = "000000"
GRAY_HEX = "595959"

CHI_BLUE_RGB = RGBColor(0x41, 0xB6, 0xE6)
CHI_RED_RGB = RGBColor(0xDA, 0x29, 0x1C)
CHI_NAVY_RGB = RGBColor(0x00, 0x2F, 0x6C)
GRAY_RGB = RGBColor(0x59, 0x59, 0x59)

# -- Final production results (2026-05-14: post review-queue elimination pass) --
HEADLINE = {
    "rows_total": 784_556,
    # 100% mapped — every row carries a Business Category (or an explicit "no description" tag)
    "rows_mapped": 784_556,
    "mapped_pct": "100.0%",
    "review_queue": 0,
    "review_pct": "0.0%",
    # 4-tier classification breakdown (sums to rows_total)
    "n_keyword_rule": 688_044,
    "pct_keyword_rule": "87.7%",
    "n_ai_assist": 93_518,
    "pct_ai_assist": "11.9%",
    "n_account_pattern": 2_028,
    "pct_account_pattern": "0.3%",
    "n_no_description": 966,
    "pct_no_description": "0.1%",
    # Aggregate "auto-classified" = rule + account + AI-assist (any deterministic or AI source)
    "auto_classified": 783_590,
    "auto_pct": "99.9%",
    "runtime": "~15 min batch + 5 sec resolver",
    "rules_total": 7_012,
    "rules_curated": 246,
    "rules_ai": 6_766,
    "categories": 17,
    "nigp_classes": 138,
}

CATEGORY_DISTRIBUTION = [
    ("Grants & Pass-Through Funding", 222_109, "28.3%"),
    ("Professional & Administrative Services", 108_576, "13.8%"),
    ("Equipment Rental & Leasing", 86_156, "11.0%"),
    ("Office, Print & Marketing", 62_319, "7.9%"),
    ("Facilities Operations & Maintenance", 54_899, "7.0%"),
    ("IT, Telecom & Audio/Visual", 51_744, "6.6%"),
    ("Vehicles & Fleet", 47_858, "6.1%"),
    ("Construction Materials", 29_466, "3.8%"),
    ("Janitorial, Sanitation & Waste", 23_903, "3.0%"),
    ("Landscaping, Grounds & Irrigation", 23_789, "3.0%"),
    ("Public Safety, Uniforms & PPE", 19_257, "2.5%"),
    ("Construction & Trades Services", 15_136, "1.9%"),
    ("Medical & Health Services", 12_784, "1.6%"),
    ("Heavy Equipment & Machinery", 12_166, "1.6%"),
    ("Chemicals & Water Treatment", 6_093, "0.8%"),
    ("Animal Care & Veterinary", 4_847, "0.6%"),
    ("Furniture & Furnishings", 2_488, "0.3%"),
    ("Unclassified — No Description", 966, "0.1%"),
]

CATEGORY_DEFINITIONS = [
    ("Facilities Operations & Maintenance",
     "HVAC, lighting, plumbing components, doors, pumps, building hardware, batteries, airport facility upkeep."),
    ("Public Safety, Uniforms & PPE",
     "Police/fire uniforms, badges, fire equipment, prison equipment, safety footwear, SCBA breathing systems, lockup food."),
    ("Construction Materials",
     "Concrete, asphalt, steel, rebar, water/sewer pipe, manhole structures, brick — raw materials installed into infrastructure."),
    ("Vehicles & Fleet",
     "Vehicles, parts, repair services, fuel-related taxes/fees, watercraft."),
    ("Professional & Administrative Services",
     "Management services, security guards, armored car, records storage, weather forecasting, testing, logistics."),
    ("Office, Print & Marketing",
     "Office consumables, paper, signs/posters, advertising/PR, printing, graphic-art supplies, promotional items."),
    ("Equipment Rental & Leasing",
     "Short-term rentals and multi-year leases of equipment, copiers, plotters, stages, linens, traffic-control gear."),
    ("Janitorial, Sanitation & Waste",
     "Cleaning supplies, trash liners, paper towels, floor machines, dumpsters, portable-toilet servicing."),
    ("Heavy Equipment & Machinery",
     "Purchased loaders, sweepers, lifts, sprayers, conveyors, outdoor power equipment (excludes rentals)."),
    ("IT, Telecom & Audio/Visual",
     "Network services, telephone equipment, telecom cable, in-car video, A/V maintenance, printer supplies."),
    ("Landscaping, Grounds & Irrigation",
     "Landscape services, irrigation design/repair, integrated pest management, topsoil, beekeeping."),
    ("Medical & Health Services",
     "Medical gases, AEDs, lab spectrometers, medical exams, counseling and behavioral-health services."),
    ("Chemicals & Water Treatment",
     "Chlorine, sodium hexametaphosphate, phosphoric acid, ice/snow melt chemicals, runway deicer, hazmat handling."),
    ("Animal Care & Veterinary",
     "Live animals, animal feed, litter, veterinary supplies and services, beekeeping (animal-husbandry side)."),
    ("Furniture & Furnishings",
     "Office and institutional furniture, mattresses, fitness equipment, furniture repair/installation."),
    ("Construction & Trades Services",
     "Skilled construction labor: bridge construction, expansion-joint repair, contracted tradesmen."),
    ("Grants & Pass-Through Funding",
     "Subgrant disbursements to community-based organizations (DFSS, CDPH, BACP, DOH, DPD, DCASE) — not a commodity purchase."),
]


# =========================================================================
# 1) EXCEL — single-tab leadership summary
# =========================================================================
def build_excel():
    wb = Workbook()
    ws = wb.active
    ws.title = "Leadership Summary"
    ws.sheet_view.showGridLines = False

    blue_fill = PatternFill("solid", fgColor=CHI_BLUE_HEX)
    navy_fill = PatternFill("solid", fgColor=CHI_NAVY_HEX)
    red_fill = PatternFill("solid", fgColor=CHI_RED_HEX)
    lt_fill = PatternFill("solid", fgColor=CHI_LT_BLUE_HEX)

    title_font = Font(name="Calibri", size=18, bold=True, color=WHITE_HEX)
    subtitle_font = Font(name="Calibri", size=11, italic=True, color=WHITE_HEX)
    section_font = Font(name="Calibri", size=12, bold=True, color=WHITE_HEX)
    metric_value_font = Font(name="Calibri", size=20, bold=True, color=CHI_RED_HEX)
    metric_label_font = Font(name="Calibri", size=10, color=GRAY_HEX)
    table_header_font = Font(name="Calibri", size=11, bold=True, color=WHITE_HEX)
    body_font = Font(name="Calibri", size=11, color=BLACK_HEX)
    footnote_font = Font(name="Calibri", size=9, italic=True, color=GRAY_HEX)

    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    right = Alignment(horizontal="right", vertical="center")

    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # Column widths
    widths = {"A": 4, "B": 42, "C": 16, "D": 16, "E": 60}
    for col, w in widths.items():
        ws.column_dimensions[col].width = w

    # Title banner
    ws.merge_cells("B2:E3")
    c = ws["B2"]
    c.value = "City of Chicago  —  Department of Procurement Services"
    c.font = title_font
    c.fill = navy_fill
    c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
    ws.row_dimensions[2].height = 22
    ws.row_dimensions[3].height = 22

    ws.merge_cells("B4:E4")
    c = ws["B4"]
    c.value = "NIGP Procurement Taxonomy — Final Production Results, May 2026 (100% Mapped)"
    c.font = subtitle_font
    c.fill = blue_fill
    c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
    ws.row_dimensions[4].height = 20

    # Headline metrics row (3 metric tiles)
    ws.row_dimensions[6].height = 8  # spacer

    metric_row_label = 7
    metric_row_value = 8
    tiles = [
        (f"{HEADLINE['rows_total']:,}", "Records classified (AP activity 2017 – 2023)"),
        (HEADLINE["mapped_pct"], "Rows fully mapped to a Business Category"),
        (HEADLINE["review_pct"], "Rows remaining in review queue"),
    ]
    columns_for_tiles = [("B", "B"), ("C", "D"), ("E", "E")]
    for (val, lab), (start, end) in zip(tiles, columns_for_tiles):
        ws.merge_cells(f"{start}{metric_row_value}:{end}{metric_row_value}")
        ws.merge_cells(f"{start}{metric_row_label}:{end}{metric_row_label}")
        v_cell = ws[f"{start}{metric_row_value}"]
        l_cell = ws[f"{start}{metric_row_label}"]
        v_cell.value = val
        v_cell.font = metric_value_font
        v_cell.alignment = center
        v_cell.fill = lt_fill
        l_cell.value = lab
        l_cell.font = metric_label_font
        l_cell.alignment = center
        l_cell.fill = lt_fill
    ws.row_dimensions[metric_row_label].height = 22
    ws.row_dimensions[metric_row_value].height = 36

    # Section header — Category distribution
    ws.row_dimensions[10].height = 8

    ws.merge_cells("B11:E11")
    c = ws["B11"]
    c.value = "Final Business Category Distribution (auto-classified rows)"
    c.font = section_font
    c.fill = blue_fill
    c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
    ws.row_dimensions[11].height = 22

    # Table headers
    headers = ["#", "Business Category", "Rows", "% of Classified", "What's in it"]
    for i, h in enumerate(headers):
        col = get_column_letter(2 + i)
        cell = ws[f"{col}12"]
        cell.value = h
        cell.font = table_header_font
        cell.fill = navy_fill
        cell.alignment = center
        cell.border = border
    ws.row_dimensions[12].height = 24

    def_lookup = {n: d for n, d in CATEGORY_DEFINITIONS}
    row = 13
    for idx, (name, count, pct) in enumerate(CATEGORY_DISTRIBUTION, start=1):
        ws[f"B{row}"] = idx
        ws[f"C{row}"] = name
        ws[f"D{row}"] = count
        ws[f"E{row}"] = pct
        ws[f"F{row}"] = def_lookup.get(name, "")
        # We laid out as B..E plus E for description — fix by using separate column for desc
        row += 1

    # Reset — rewrite table with definition column F
    # Clear the columns we just touched 13..end
    end_row = 13 + len(CATEGORY_DISTRIBUTION) - 1
    for r in range(13, end_row + 1):
        for col in ("B", "C", "D", "E", "F"):
            ws[f"{col}{r}"].value = None

    # Re-do widths — push description to its own column F, keep summary on B-E
    ws.column_dimensions["B"].width = 4
    ws.column_dimensions["C"].width = 42
    ws.column_dimensions["D"].width = 14
    ws.column_dimensions["E"].width = 16
    ws.column_dimensions["F"].width = 75

    # Re-write headers with the corrected layout
    headers = ["#", "Business Category", "Rows", "% of Classified", "What's in it"]
    cols = ["B", "C", "D", "E", "F"]
    for col, h in zip(cols, headers):
        cell = ws[f"{col}12"]
        cell.value = h
        cell.font = table_header_font
        cell.fill = navy_fill
        cell.alignment = center
        cell.border = border

    # Data rows with banded fill
    for i, (name, count, pct) in enumerate(CATEGORY_DISTRIBUTION):
        r = 13 + i
        band = lt_fill if i % 2 == 0 else PatternFill("solid", fgColor=WHITE_HEX)
        values = [(i + 1, center), (name, left), (count, right), (pct, center), (def_lookup.get(name, ""), left)]
        for col, (val, align) in zip(cols, values):
            cell = ws[f"{col}{r}"]
            cell.value = val
            cell.font = body_font
            cell.alignment = align
            cell.fill = band
            cell.border = border
            if col == "D":
                cell.number_format = "#,##0"
        ws.row_dimensions[r].height = 28

    # Total row
    total_row = 13 + len(CATEGORY_DISTRIBUTION)
    ws.merge_cells(f"B{total_row}:C{total_row}")
    cell = ws[f"B{total_row}"]
    cell.value = "TOTAL — Auto-Classified"
    cell.font = Font(name="Calibri", size=11, bold=True, color=WHITE_HEX)
    cell.fill = red_fill
    cell.alignment = Alignment(horizontal="left", vertical="center", indent=1)
    cell.border = border

    cell = ws[f"D{total_row}"]
    cell.value = HEADLINE["auto_classified"]
    cell.font = Font(name="Calibri", size=11, bold=True, color=WHITE_HEX)
    cell.fill = red_fill
    cell.alignment = right
    cell.number_format = "#,##0"
    cell.border = border

    cell = ws[f"E{total_row}"]
    cell.value = HEADLINE["auto_pct"]
    cell.font = Font(name="Calibri", size=11, bold=True, color=WHITE_HEX)
    cell.fill = red_fill
    cell.alignment = center
    cell.border = border

    cell = ws[f"F{total_row}"]
    cell.value = f"All {HEADLINE['rows_total']:,} rows classified — 0 rows remain in review queue (incl. {HEADLINE['n_no_description']:,} empty-description rows explicitly tagged \"Unclassified\")."
    cell.font = Font(name="Calibri", size=11, bold=True, italic=True, color=WHITE_HEX)
    cell.fill = red_fill
    cell.alignment = left
    cell.border = border
    ws.row_dimensions[total_row].height = 28

    # Methodology footnotes block
    foot_start = total_row + 2
    ws.merge_cells(f"B{foot_start}:F{foot_start}")
    c = ws[f"B{foot_start}"]
    c.value = "How These Numbers Were Produced"
    c.font = section_font
    c.fill = blue_fill
    c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
    ws.row_dimensions[foot_start].height = 22

    footnotes = [
        f"Source file:  EY raw data extract — {HEADLINE['rows_total']:,} purchase-order and invoice line records covering AP activity years 2017, 2020, 2021, 2023.",
        "Classifier inputs:  description text + Chicago FMPS account/object/fund codes only.  Vendor name and EY-supplied NIGP codes are NOT used.",
        f"Rule base:  {HEADLINE['rules_total']:,} rules total — {HEADLINE['rules_curated']} hand-curated by procurement leadership and {HEADLINE['rules_ai']:,} AI-mined long-tail patterns.",
        "AI use:  one-time, build-time only.  Production runtime is rules-only — no API key, no internet, no recurring cost, no per-classification charge.",
        f"Runtime:  {HEADLINE['runtime']} end-to-end on a standard workstation.  Repeatable on any future Chicago procurement extract.",
        f"Taxonomy:  three-level — {HEADLINE['categories']} Business Categories  →  {HEADLINE['nigp_classes']} NIGP 3-digit Classes  →  470 NIGP 5-digit Class-Items.",
        "Audit trail:  every classified row carries the exact rule that fired and a confidence score.  Rule files are CSVs — editable by procurement staff, no Python required.",
    ]
    for i, fn in enumerate(footnotes):
        r = foot_start + 1 + i
        ws.merge_cells(f"B{r}:F{r}")
        cell = ws[f"B{r}"]
        cell.value = "•  " + fn
        cell.font = footnote_font
        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True, indent=1)
        ws.row_dimensions[r].height = 18

    # Document footer
    foot_doc = foot_start + 1 + len(footnotes) + 1
    ws.merge_cells(f"B{foot_doc}:F{foot_doc}")
    cell = ws[f"B{foot_doc}"]
    cell.value = (
        "Prepared by James H. Kirby III, CSCP, MS-SCM  |  Department of Procurement Services  |  "
        "Final production run: 30 April 2026  |  Methodology v1.1"
    )
    cell.font = Font(name="Calibri", size=9, italic=True, color=GRAY_HEX)
    cell.alignment = Alignment(horizontal="left", vertical="center", indent=1)
    ws.row_dimensions[foot_doc].height = 18

    # Print setup — fit to one page wide
    ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_options.horizontalCentered = True
    ws.page_margins.left = 0.4
    ws.page_margins.right = 0.4
    ws.page_margins.top = 0.5
    ws.page_margins.bottom = 0.5

    out = OUT_DIR / "NIGP_Summary_for_Leadership_JHK3.xlsx"
    wb.save(out)
    return out


# =========================================================================
# Word helpers (shared by brief + methodology)
# =========================================================================
def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_horizontal_rule(paragraph, color_hex=CHI_BLUE_HEX, size=12):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def style_heading(paragraph, color_rgb, size_pt, bold=True):
    for run in paragraph.runs:
        run.font.color.rgb = color_rgb
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.name = "Calibri"


def add_bullet(doc, text, bold_lead=None, size=11):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        run = p.add_run(bold_lead)
        run.font.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.color.rgb = CHI_NAVY_RGB
        run2 = p.add_run("  " + text)
        run2.font.name = "Calibri"
        run2.font.size = Pt(size)
    else:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.color.rgb = CHI_NAVY_RGB
    add_horizontal_rule(p, CHI_BLUE_HEX, size=18)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = CHI_BLUE_RGB
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_body(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_title_banner(doc, title, subtitle, author_line):
    """Big colored title block at top of document."""
    # Title (navy, 22pt)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = CHI_NAVY_RGB
    p.paragraph_format.space_after = Pt(0)

    # Subtitle (blue, 13pt)
    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = CHI_BLUE_RGB
    run.font.italic = True
    p.paragraph_format.space_after = Pt(0)
    add_horizontal_rule(p, CHI_RED_HEX, size=18)

    # Author line (gray, small)
    p = doc.add_paragraph()
    run = p.add_run(author_line)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_RGB
    p.paragraph_format.space_after = Pt(12)


def add_callout_box(doc, label, body_text):
    """Single-row table styled as a colored callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, CHI_LT_BLUE_HEX)
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    run = p.add_run(label)
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = CHI_RED_RGB
    p.paragraph_format.space_after = Pt(2)

    p2 = cell.add_paragraph()
    run = p2.add_run(body_text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = CHI_NAVY_RGB
    # padding via cell margins
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "120")
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)

    # spacer paragraph after the box
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def style_table_header(row, color_hex=CHI_NAVY_HEX):
    for cell in row.cells:
        set_cell_shading(cell, color_hex)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.name = "Calibri"
                r.font.size = Pt(10)


def set_table_borders(table, color_hex="BFBFBF"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)


def set_doc_margins(doc, cm=2.0):
    for section in doc.sections:
        section.top_margin = Cm(cm)
        section.bottom_margin = Cm(cm)
        section.left_margin = Cm(cm)
        section.right_margin = Cm(cm)


# =========================================================================
# 2) WORD — 2-page Executive Brief
# =========================================================================
def build_executive_brief():
    doc = Document()
    set_doc_margins(doc, cm=1.8)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title_banner(
        doc,
        title="Chicago's First In-House Procurement Taxonomy",
        subtitle="Executive Brief — NIGP-Aligned Classification Engine",
        author_line=(
            "City of Chicago  •  Department of Procurement Services  •  "
            "Prepared by James H. Kirby III, CSCP, MS-SCM  •  30 April 2026  •  Methodology v1.1"
        ),
    )

    # The Bottom Line
    add_h1(doc, "The Bottom Line")
    add_body(
        doc,
        "Until now, Chicago has not maintained its own commodity classification of what the City buys. "
        "This project delivers the City's first internally-owned procurement taxonomy and a reusable "
        "classification engine that runs on every historical and future procurement extract — "
        "with no recurring vendor cost and a complete audit trail on every record.",
    )

    add_callout_box(
        doc,
        label="HEADLINE OUTCOME — Production Run, 14 May 2026 (100% Mapped)",
        body_text=(
            f"784,556 records classified across City spend (AP activity years 2017, 2020, 2021, 2023)  •  "
            f"100% mapped — 0 rows remain in review queue  •  "
            f"87.7% via deterministic keyword rule  •  "
            f"11.9% via AI-assist fallback (one-time AI mining harvested at build, no runtime API cost)  •  "
            f"0.3% via Chicago account-code pattern  •  "
            f"0.1% empty-description rows explicitly tagged \"Unclassified\"."
        ),
    )

    # What we built
    add_h1(doc, "What We Built")
    add_bullet(doc, "Three classification levels are assigned to every record: a leadership-friendly Business Category, a 3-digit NIGP Class, and (where description supports it) a 5-digit NIGP Item.", bold_lead="Three-level taxonomy.")
    add_bullet(doc, "17 mutually exclusive Business Categories cover everything Chicago buys — from Facilities Operations to Public Safety to Grants & Pass-Through Funding.", bold_lead="17 Business Categories.")
    add_bullet(doc, "Aligned to the public NIGP commodity framework used by U.S. public procurement entities. Enables peer benchmarking, audit defensibility, and future portability.", bold_lead="NIGP-aligned.")
    add_bullet(doc, "Same classifier supports bulk historical files and single-record paste-a-description for live PO work.", bold_lead="Two operating modes.")
    add_bullet(doc, "Rules live in editable CSV files. Procurement staff add or retire rules without touching code.", bold_lead="Procurement-staff editable.")

    # How the engine decides
    add_h1(doc, "How the Engine Decides")
    add_bullet(doc, "Description text from the PO and invoice fields, plus Chicago's own FMPS account/object/fund codes.", bold_lead="Inputs used:")
    add_bullet(doc, "Vendor name (a vendor can sell across many categories) and EY-supplied NIGP codes (another consultant's work product, not Chicago's authoritative judgment).", bold_lead="Inputs deliberately excluded:")
    add_bullet(doc, "First a hand-curated keyword rule, then an AI-mined long-tail keyword rule, then a Chicago account-code pattern. Anything still unmatched is filled at the AI-assist fallback layer from the saved one-time AI mining output (no runtime API calls).", bold_lead="Decision pipeline:")
    add_bullet(doc, "Every classified record carries the exact rule that fired, a confidence score, and a review flag. Procurement staff can trace any single decision end-to-end.", bold_lead="Audit trail:")

    # AI use — defensible
    add_h1(doc, "AI Use — Limited, Transparent, One-Time")
    add_body(
        doc,
        "The production classifier Chicago operates is rules-only. There is no AI in the runtime path — no API key, "
        "no internet dependency, no per-classification charge. AI was used once during the build, in a bounded way, "
        "to mine recurring patterns from the long tail of descriptions where no hand-curated rule applied. The AI's "
        "output was harvested into rules and frozen. Every AI-mined rule is auditable and editable.",
    )

    # Results table — concise
    add_h1(doc, "Production Results — Classification Method Breakdown")
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "Classification Method"
    hdr[1].text = "Rows"
    hdr[2].text = "Share"
    style_table_header(table.rows[0])

    rows_data = [
        ("Keyword rule (hand-curated + AI-mined)", f"{HEADLINE['n_keyword_rule']:,}", HEADLINE['pct_keyword_rule']),
        ("AI-assist fallback (one-time AI output, no runtime API)", f"{HEADLINE['n_ai_assist']:,}", HEADLINE['pct_ai_assist']),
        ("Chicago account-code pattern (subgrant disbursements)", f"{HEADLINE['n_account_pattern']:,}", HEADLINE['pct_account_pattern']),
        ("Unclassified — No Description (explicitly tagged)", f"{HEADLINE['n_no_description']:,}", HEADLINE['pct_no_description']),
        ("TOTAL", f"{HEADLINE['rows_total']:,}", HEADLINE['mapped_pct']),
    ]
    for outcome, n, share in rows_data:
        r = table.add_row().cells
        r[0].text = outcome
        r[1].text = n
        r[2].text = share

    for row_i, row in enumerate(table.rows[1:], start=1):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
            if row_i % 2 == 1:
                set_cell_shading(cell, CHI_LT_BLUE_HEX)

    # Spacer
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # What this means for leadership
    add_h1(doc, "What This Means for DPS Leadership")
    add_bullet(doc, "Spend $ totals, vendor consolidation analysis, and category benchmarking are now possible against a Chicago-owned, internally-defensible classification — not against a consultant's work product.", bold_lead="Reporting unlocked.")
    add_bullet(doc, "The classifier runs on any future Chicago procurement extract with no code changes. Procurement staff edit the CSV rule files as new commodity types appear.", bold_lead="Repeatable.")
    add_bullet(doc, "AI was used once, with bounded inputs and full provenance metadata on every harvested rule. Any auditor can reconstruct how every record was classified.", bold_lead="Audit-defensible.")
    add_bullet(doc, "Every row is mapped to a Business Category. The 966 \"Unclassified — No Description\" rows are explicitly tagged so downstream consumers can filter on them rather than treat them as silent failures.", bold_lead="Maintainable.")

    out = OUT_DIR / "NIGP_Executive_Brief_JHK3.docx"
    doc.save(out)
    return out


# =========================================================================
# 3) WORD — Methodology, leadership-friendly
# =========================================================================
def build_methodology():
    doc = Document()
    set_doc_margins(doc, cm=2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Cover banner ----
    add_title_banner(
        doc,
        title="Procurement Taxonomy & Classification Methodology",
        subtitle="A Plain-Language Methodology Document for Department of Procurement Services Leadership",
        author_line=(
            "City of Chicago  •  Department of Procurement Services  •  "
            "Prepared by James H. Kirby III, CSCP, MS-SCM  •  Document version 1.1, finalized 30 April 2026"
        ),
    )

    # ---- How to read this document ----
    add_h1(doc, "How to Read This Document")
    add_bullet(doc, "What was built, why, and how it works — written for non-technical leadership.", bold_lead="Purpose:")
    add_bullet(doc, "13 short sections plus a one-page decisions appendix at the end.", bold_lead="Structure:")
    add_bullet(doc, "Each section opens with a one-line takeaway. Skim the headings; read the bullets where you need detail.", bold_lead="How to skim:")
    add_bullet(doc, "Review the Executive Summary and Section 7 (Production Results). Everything else is supporting detail.", bold_lead="Most important sections:")

    # ---- Executive Summary ----
    add_h1(doc, "Executive Summary")
    add_body(
        doc,
        "Chicago has not historically maintained its own commodity classification system for procurement spend. "
        "This project delivers the City's first internally-owned commodity taxonomy and a reusable, defensible "
        "classification engine that runs on every historical and future Chicago procurement extract — with no "
        "recurring vendor cost and a full audit trail.",
    )
    add_callout_box(
        doc,
        label="HEADLINE OUTCOME — 14 May 2026",
        body_text=(
            "784,556 records classified across City spend (AP activity years 2017, 2020, 2021, 2023)  •  "
            "100% mapped — 0 rows in review queue  •  "
            "87.7% via deterministic keyword rule  •  11.9% via AI-assist fallback (one-time AI mining)  •  "
            "0.3% via Chicago account-code pattern  •  0.1% empty-description rows explicitly tagged."
        ),
    )
    add_h2(doc, "What's different from prior consulting work")
    add_bullet(doc, "Chicago now owns its own classification. The earlier EY consulting deliverable is a useful raw data source, but its NIGP labels are EY's work product — not Chicago's authoritative judgment.", bold_lead="In-house ownership.")
    add_bullet(doc, "Every classified record carries the exact rule that fired, a confidence score, and a review flag. Any auditor can reconstruct any decision.", bold_lead="Full audit trail.")
    add_bullet(doc, "Procurement staff can edit rules in CSV files. No software developer is required to add a new rule or retire an outdated one.", bold_lead="Staff-editable.")
    add_bullet(doc, "The classifier runs on any future Chicago procurement extract with no code modifications.", bold_lead="Repeatable.")

    # ---- 1. Project Objective ----
    add_h1(doc, "1.  Project Objective")
    add_body(doc, "Build a transparent, tier-labeled classification engine that:")
    add_bullet(doc, "the public-standard NIGP commodity framework, for inter-agency comparability and audit defensibility.", bold_lead="Uses")
    add_bullet(doc, "a Chicago-specific Business Category rollup on top, designed for leadership reporting.", bold_lead="Layers")
    add_bullet(doc, "in two modes — bulk historical files, and single-record live PO work.", bold_lead="Operates")
    add_bullet(doc, "transparency over a single \"auto-classified\" headline. Every row carries an explicit tier-source label, so leadership and auditors can filter by deterministic-rule, account-pattern, or AI-assist tier rather than reading one blended accuracy number.", bold_lead="Prioritizes")
    add_bullet(doc, "fully repeatable on future raw extracts. Rule files are CSVs; procurement staff own them.", bold_lead="Is")

    # ---- 2. Source Data ----
    add_h1(doc, "2.  Source Data")
    add_h2(doc, "What was analyzed")
    add_bullet(doc, "EY raw data extract (`ey raw data.xlsx`, 326 MB).", bold_lead="File:")
    add_bullet(doc, "784,556 purchase-order and invoice line records.", bold_lead="Records:")
    add_bullet(doc, "AP activity years 2017, 2020, 2021, 2023 (all rows).", bold_lead="Time span:")
    add_bullet(doc, "EY's prior NIGP classification (~30% of rows). The classifier does NOT use these labels — Chicago is building its own.", bold_lead="EY-supplied labels:")
    add_h2(doc, "Data quality checks")
    add_bullet(doc, "Line-level descriptions are populated on virtually all rows. Only 967 of 784,556 (0.1%) have no usable description anywhere.")
    add_bullet(doc, "17.8% of high-level PO Description values are uninformative (\"Misc\", \"Per contract\"). The classifier falls back to the line-level description fields, which carry substantive content.")
    add_bullet(doc, "Two date fields contained typos or open-ended placeholders. Documented; no impact on classification.")

    # ---- 3. Three-Level Taxonomy ----
    add_h1(doc, "3.  The Three-Level Taxonomy")
    add_body(doc, "Every classified record carries three levels — designed to serve both leadership reporting and sourcing/audit needs.")

    add_h2(doc, "Level 1 — Business Category (17 buckets, custom Chicago rollup)")
    add_body(doc, "Mutually exclusive, collectively exhaustive, and immediately recognizable to leadership and audit reviewers.")
    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "Business Category"
    hdr[1].text = "What's in it"
    style_table_header(table.rows[0])
    for i, (name, desc) in enumerate(CATEGORY_DEFINITIONS):
        r = table.add_row().cells
        r[0].text = name
        r[1].text = desc
        for cell in r:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        if i % 2 == 0:
            for cell in r:
                set_cell_shading(cell, CHI_LT_BLUE_HEX)
    # Set first column narrower
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(11.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_body(
        doc,
        "Note: \"Grants & Pass-Through Funding\" is not a commodity purchase — it is the financial transfer of dollars to subrecipients (DFSS social services, CDPH public health, BACP business development, DOH housing, DPD planning, DCASE arts/culture). It has different procurement governance and audit treatment than commodity buying, so it has its own bucket. NIGP does not have a subgrants code; this category therefore has no NIGP class.",
    )

    add_h2(doc, "Level 2 — NIGP 3-digit Class")
    add_bullet(doc, "Each Business Category maps to one or more NIGP 3-digit Classes (138 distinct classes appear in the EY data).")
    add_bullet(doc, "NIGP is a public, inter-agency-compatible commodity framework used widely across U.S. public procurement.")
    add_bullet(doc, "Aligning to NIGP enables peer benchmarking, audit defensibility, and future portability.")

    add_h2(doc, "Level 3 — NIGP 5-digit Class-Item")
    add_bullet(doc, "When a description supports specificity, the classifier assigns a 5-digit Class-Item code (470 distinct codes in our working catalog).")
    add_bullet(doc, "When confidence at 5-digit is insufficient, the row carries the broader 3-digit Class only — flagged via the `NIGP_Match_Level` field as \"broad\" instead of \"exact.\"")

    # ---- 4. Inputs ----
    add_h1(doc, "4.  What the Classifier Looks At — and What It Ignores")
    add_h2(doc, "Inputs the classifier USES")
    add_bullet(doc, "Description text from four PO and invoice fields. The classifier picks the first non-blank, substantive description across all four.", bold_lead="Description text:")
    add_bullet(doc, "Used as supplemental signal where description text alone is ambiguous (most consequentially, the 220xxx subgrant accounts).", bold_lead="Chicago FMPS account/object/fund codes:")
    add_h2(doc, "Inputs the classifier does NOT use — by design")
    add_bullet(doc, "The same vendor often sells across many categories (e.g., a generic supplier selling IT, office, and janitorial on one contract). Vendor-based inference introduces misclassification risk.", bold_lead="Vendor name:")
    add_bullet(doc, "These represent EY's prior classification work, not Chicago's authoritative judgment. Chicago is independently classifying every record from description text.", bold_lead="EY-supplied NIGP codes:")

    # ---- 5. Decision Pipeline ----
    add_h1(doc, "5.  How a Single Record Is Classified")
    add_body(doc, "The classifier evaluates each record through a deterministic four-tier pipeline. The first tier that fires assigns the classification — subsequent tiers don't run for that record. This produces a clean audit trail: every classified row carries the exact source that drove the decision.")

    add_h2(doc, "Tier 1 — Keyword rules on description text")
    add_bullet(doc, "246 hand-curated rules drafted by procurement leadership.")
    add_bullet(doc, "6,766 AI-mined rules harvested once during the build (more on AI use in Section 6).")
    add_bullet(doc, "Match types: exact (full string), starts-with (prefix), contains (substring). Priority: exact > starts-with > contains. First match wins.")

    add_h2(doc, "Tier 2 — Account-code patterns (only if Tier 1 didn't fire)")
    add_bullet(doc, "Maps Chicago FMPS account codes to (Business Category, NIGP Class).")
    add_bullet(doc, "Most consequential rule: the 220xxx subgrant-disbursement account series. Empirical analysis confirmed these accounts are 93–100% used for subgrant disbursements regardless of description text quality.")

    add_h2(doc, "Tier 3 — AI-assist resolver (only if Tiers 1 and 2 didn't fire)")
    add_bullet(doc, "Looks up Description_Best in the saved 2026-04-30 AI mining output (30,342-row CSV) and applies the model's Business Category and NIGP class.")
    add_bullet(doc, "No new API call. The classifier and resolver operate on saved files. The original AI charge of $27 has already been paid; this tier consumes more of the same output.")
    add_bullet(doc, "Records are tagged Classification_Method = ai_assist and Classification_Confidence = AI-high / AI-medium / AI-low so leadership and auditors can filter by sub-tier.")

    add_h2(doc, "Tier 4 — Unclassified — No Description (only if Tiers 1, 2, 3 didn't fire)")
    add_bullet(doc, "Records with no usable description text in any of the four description fields. The data itself is the limitation, not the classifier.")
    add_bullet(doc, "Explicitly tagged Business_Category = \"Unclassified — No Description\" rather than silently dropped. Downstream consumers can filter on this tag.")

    # ---- 6. AI Use ----
    add_h1(doc, "6.  Use of AI — Limited, Transparent, One-Time")
    add_callout_box(
        doc,
        label="DEFENSIBILITY UPFRONT",
        body_text=(
            "The production classifier Chicago operates is rules-only. No AI in the runtime path. No API key, "
            "no internet dependency, no recurring cost, no per-classification charge. Any City staff member can "
            "run it from their workstation."
        ),
    )
    add_body(doc, "AI was used ONCE during the initial build for one bounded purpose: to mine recurring patterns from the long tail of descriptions where no hand-curated rule applied. The AI's output was harvested into the keyword rules file and then frozen.")

    add_h2(doc, "What the AI did and didn't see")
    add_bullet(doc, "Anthropic Claude Haiku 4.5.", bold_lead="Model:")
    add_bullet(doc, "A single batch run over 30,342 unique long-tail descriptions.", bold_lead="Invocation:")
    add_bullet(doc, "Description text plus a system prompt listing the 17 Business Categories and 138 NIGP classes. Vendor name, EY codes, and Chicago account codes were NOT passed to the AI.", bold_lead="What the AI saw:")
    add_bullet(doc, "Output was constrained by JSON schema with a closed enumeration. The model could not invent codes outside the catalog.", bold_lead="Output bounds:")

    add_h2(doc, "How AI proposals are applied")
    add_bullet(doc, "(6,419 proposals, 21.2%) → auto-promoted into the rule file.", bold_lead="High confidence:")
    add_bullet(doc, "(11,645 proposals, 38.4%) → promoted only if the description recurred in 5+ rows. 1,998 met the bar; remainder applied at fill time via AI-assist fallback.", bold_lead="Medium confidence:")
    add_bullet(doc, "(12,278 proposals, 40.5%) → not promoted as rules. Applied at fill time via AI-assist fallback for descriptions no rule matches.", bold_lead="Low confidence:")
    add_bullet(doc, "6,766 AI-mined rules in the rule file + 93,518 rows filled at the AI-assist fallback layer from the same one-time AI output (no new API calls).", bold_lead="Net result:")

    add_h2(doc, "Why this AI use is defensible")
    add_bullet(doc, "Production runs do not call AI. The single 2026-04-30 mining run produced a CSV; every downstream consumption is a local file read. No API key, internet access, or per-classification charge is required at runtime.", bold_lead="One-time, not ongoing.")
    add_bullet(doc, "The AI cannot output codes outside the 17 categories and 138 classes Chicago's procurement leadership approved.", bold_lead="Bounded.")
    add_bullet(doc, "Every AI-mined rule and every AI-assist resolver assignment carries provenance metadata — model name, confidence level, source row count, and the model's reasoning.", bold_lead="Transparent.")
    add_bullet(doc, "Procurement staff can edit, demote, or remove any AI-mined rule. AI rules have no special status — they're evaluated alongside hand-curated rules. AI-assist resolver assignments can be promoted into the curated rule file at any time.", bold_lead="Reviewable.")
    add_bullet(doc, "If Chicago later decides to remove all AI-mined content, both the rule file and the resolver source CSV can be deleted. The classifier still operates on hand-curated rules + account patterns alone — coverage drops, but the system does not break.", bold_lead="Replaceable.")
    add_bullet(doc, "AI-assist resolver assignments are tagged AI-high / AI-medium / AI-low in the output file. Leadership and auditors can filter or exclude any confidence tier they choose. The AI tier is not hidden inside a generic \"auto-classified\" bucket.", bold_lead="Confidence-labeled, not laundered.")

    # ---- 7. Production Results ----
    add_h1(doc, "7.  Production Run Results — 14 May 2026 (100% Mapped)")
    add_body(doc, "The classifier was run end-to-end against all 784,556 rows of the EY file, followed by a post-classifier resolver pass that merges the one-time AI mining output into any rows no keyword rule matched. Result: 100% of rows mapped, 0 rows in review queue.")

    add_h2(doc, "Coverage by classification method")
    table = doc.add_table(rows=1, cols=4)
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "Method"
    hdr[1].text = "Rows"
    hdr[2].text = "Share"
    hdr[3].text = "Notes"
    style_table_header(table.rows[0])
    coverage_rows = [
        ("Keyword rule", f"{HEADLINE['n_keyword_rule']:,}", HEADLINE['pct_keyword_rule'], "246 hand-curated + 6,766 AI-mined"),
        ("AI-assist fallback", f"{HEADLINE['n_ai_assist']:,}", HEADLINE['pct_ai_assist'], "Filled from saved one-time AI output; zero runtime API cost"),
        ("Account-code pattern", f"{HEADLINE['n_account_pattern']:,}", HEADLINE['pct_account_pattern'], "Subgrant 220xxx accounts"),
        ("Unclassified — No Description", f"{HEADLINE['n_no_description']:,}", HEADLINE['pct_no_description'], "Empty Description_Best; explicitly tagged"),
    ]
    for i, (a, b, c, d) in enumerate(coverage_rows):
        r = table.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c
        r[3].text = d
        for cell in r:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        if i % 2 == 0:
            for cell in r:
                set_cell_shading(cell, CHI_LT_BLUE_HEX)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_h2(doc, "Curated vs. AI contribution")
    add_bullet(doc, "246 curated rules + 6,766 AI-mined rules together hit 688,044 rows (87.7%). The hand-curated set targets the highest-frequency patterns; AI-mined rules handle the recurring long tail.", bold_lead="Keyword rules:")
    add_bullet(doc, "93,518 rows filled at the resolver layer from the saved one-time AI mining output. These are the AI calls that didn't promote into rules (medium/low confidence or singleton descriptions). No new API calls — the AI output from 2026-04-30 is reused.", bold_lead="AI-assist fallback:")
    add_bullet(doc, "2,028 rows resolved by 220xxx Chicago subgrant account-code patterns; 966 rows tagged \"Unclassified — No Description\" because Description_Best was empty/null.", bold_lead="Other tiers:")

    add_h2(doc, "Final Business Category distribution")
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "Business Category"
    hdr[1].text = "Rows"
    hdr[2].text = "% of Total"
    style_table_header(table.rows[0])
    for i, (name, count, pct) in enumerate(CATEGORY_DISTRIBUTION):
        r = table.add_row().cells
        r[0].text = name
        r[1].text = f"{count:,}"
        r[2].text = pct
        for cell in r:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        if i % 2 == 0:
            for cell in r:
                set_cell_shading(cell, CHI_LT_BLUE_HEX)

    add_body(doc, "")
    add_body(
        doc,
        "Zero rows remain in the human-review queue. 966 rows (0.1%) carry the explicit "
        "\"Unclassified — No Description\" tag because Description_Best was empty/null, with no FMPS "
        "account-code pattern matching either. These are formally classified — just to an honest \"no description\" "
        "category — so downstream consumers can filter on them rather than treat them as silent failures.",
    )

    # ---- 8. Confidence ----
    add_h1(doc, "8.  Confidence and Audit Trail")
    add_body(doc, "Every classified record carries a confidence triple plus a reason — a complete decision record.")
    add_bullet(doc, "exact (5-digit Class-Item assignable), broad (3-digit Class only), review (best broader category but flagged), or empty (no match — Grants category, or no description).", bold_lead="NIGP_Match_Level:")
    add_bullet(doc, "High / Medium / Low for keyword-rule rows; AI-high / AI-medium / AI-low for resolver-filled rows. The AI- prefix makes the AI tier filterable on every row.", bold_lead="Classification_Confidence:")
    add_bullet(doc, "0.0 to 1.0 numeric value reflecting the rule or resolver match quality.", bold_lead="Confidence_Score:")
    add_bullet(doc, "Records the exact rule pattern that fired (or the resolver lookup that filled it) — a complete audit trail end-to-end.", bold_lead="Classification_Reason:")
    add_body(doc, "Through 2026-04-30, the classifier flagged any row scoring below 0.75 as Review_Flag = Yes, producing a 17.8% review queue. After the AI-assist resolver was added 2026-05-14, the field is preserved in the schema for backward compatibility and defaults to \"No\" — every row is mapped, including the 0.1% Unclassified — No Description residue.")

    # ---- 9. Output ----
    add_h1(doc, "9.  What Procurement Staff Receive")
    add_h2(doc, "Two primary deliverable files plus one diagnostic, every batch run")
    add_bullet(doc, "Every record with its full classification, post-resolver. 100% mapped.", bold_lead="outputs/NIGP_Mapping_JHK3.csv:")
    add_bullet(doc, "Subset of Review_Flag = Yes rows. 0 rows since the 2026-05-14 resolver pass; retained in schema for forward compatibility.", bold_lead="outputs/NIGP_Mapping_Review_Queue_JHK3.csv:")
    add_bullet(doc, "Diagnostic of which rules fired how many times — used to identify rule gaps.", bold_lead="data/processed/classifier_coverage_report_JHK3.csv:")
    add_h2(doc, "Lean output design")
    add_bullet(doc, "Output is ~16 columns, not all 87 raw columns. The original 87-column file is preserved untouched and can be joined back if a kitchen-sink view is needed for ad-hoc analysis.")

    # ---- 10. Single-record mode ----
    add_h1(doc, "10.  Single-Record Mode for Live PO Work")
    add_body(doc, "The same classifier core powers a paste-a-description tool. A procurement analyst can paste a single PO or requisition description and immediately receive the proposed Business Category, NIGP code, confidence, and reason.")
    add_bullet(doc, "This is the foundation for future integration with Chicago's DPS PO-creation workflow — classifying new POs at point-of-creation rather than retrospectively.")

    # ---- 11. Limitations ----
    add_h1(doc, "11.  Honest Limitations")
    add_bullet(doc, "Taxonomy was derived from one consulting deliverable. New extracts may reveal commodity types not represented in the EY file.", bold_lead="Single-source dataset.")
    add_bullet(doc, "The classifier ignores transaction date. Year-over-year inflation, contract restructuring, and reorganizations don't affect classification consistency.", bold_lead="Time-agnostic by design.")
    add_bullet(doc, "Classification reflects what was bought, not who sold it. Vendor-based views can be produced separately from the classified output.", bold_lead="No vendor signal.")
    add_bullet(doc, "Thin but non-empty descriptions (\"Misc supplies\", \"Per contract\") that recurred in the source file received an AI proposal during the 2026-04-30 mining run and are now resolved at Tier 3 with an explicit AI-low or AI-medium tag — auditors can filter on that tag where accuracy concerns warrant it. Rows with no usable description across any of the four description fields are tagged \"Unclassified — No Description\" (Tier 4) rather than being assigned a category.", bold_lead="Description quality dependent.")
    add_bullet(doc, "The 138 NIGP classes derived from the EY file are a subset of the full ~9,000-code NIGP standard. Future work should consider licensing the full catalog from Periscope Holdings.", bold_lead="Working NIGP catalog is partial.")

    # ---- 12. Governance recommendations ----
    add_h1(doc, "12.  Governance and Maintenance Recommendations")
    add_bullet(doc, "Procurement leadership reviews the 17 Business Categories annually for continued fit with reporting needs.", bold_lead="Annual taxonomy review.")
    add_bullet(doc, "Designate a procurement analyst as owner of `keyword_rules_DRAFT_JHK3.csv`. New rules added as new commodity types appear; outdated rules retired. Curated-rule count has grown from 148 (2026-04-30) to 246 (2026-05-14) on this principle.", bold_lead="Rule-file ownership.")
    add_bullet(doc, "With zero rows in the terminal review queue, staff time redirects to confidence-tier auditing. On a monthly or quarterly cadence, sample rows where Classification_Method = ai_assist and Classification_Confidence = AI-medium / AI-low. For descriptions that recur in volume and the AI tier got right, promote the pattern into the curated rule file (Tier 1). This is the new mechanism by which the rule base grows and Tier 3 shrinks over time.", bold_lead="AI-assist tier audit cadence.")
    add_bullet(doc, "When DPS integration is built, analyst overrides of single-record proposals become candidate rule data. Rule base grows; classifier improves without re-running AI.", bold_lead="Single-record feedback loop.")
    add_bullet(doc, "Expands the working catalog from 138 codes to the full ~9,000-code NIGP standard.", bold_lead="Consider procuring a Periscope NIGP license.")
    add_bullet(doc, "Once a year, sample 100 classified rows at random; have a procurement analyst classify them blind. Stratify the sample across tiers (e.g., 25 keyword-rule, 25 account-pattern, 25 AI-medium, 25 AI-low) so the audit reads agreement separately by tier.", bold_lead="Annual stratified audit benchmark.")

    # ---- 13. Sources ----
    add_h1(doc, "13.  Sources and References")
    add_bullet(doc, "Originating organization for the NIGP Commodity/Services Code framework — https://www.nigp.org/", bold_lead="NIGP — Institute for Public Procurement.")
    add_bullet(doc, "Operating context for this project — https://www.chicago.gov/city/en/depts/dps.html", bold_lead="City of Chicago Department of Procurement Services.")
    add_bullet(doc, "Source system for account/object/fund codes used as supplemental classification signal.", bold_lead="City of Chicago FMPS (Financial Management & Purchasing System).")
    add_bullet(doc, "Provider of the AI service used in the one-time pattern-mining phase — https://www.anthropic.com/", bold_lead="Anthropic — Claude API.")
    add_bullet(doc, "Data processing library — https://pandas.pydata.org/", bold_lead="Pandas.")

    # ---- Appendix A — Decisions Locked ----
    add_h1(doc, "Appendix A — Decisions Locked During Project Build")
    add_body(doc, "For audit traceability, the following design decisions were made during taxonomy construction and locked in writing.")

    decisions = [
        ("1", "Three-level taxonomy on every record: Business Category → NIGP Class → NIGP Item.",
         "Serves both executive dashboards (Business Category) and sourcing/audit needs (NIGP codes)."),
        ("2", "AI assist allowed once during build, with strict guardrails.",
         "Rules-only would leave a 20–40% review pile; bounded one-time AI use shrinks it while remaining defensible."),
        ("3", "Classifier inputs: description text + Chicago FMPS account/object/fund codes only. Vendor and EY-supplied NIGP codes are NOT inputs.",
         "Chicago must own its own classification; vendor isn't a reliable signal; EY codes are another consultant's work product."),
        ("4", "Lean ~16-column output instead of preserving all 87 raw columns.",
         "Raw file preserved separately; downstream analysis doesn't need 87 columns."),
        ("5", "Dual-mode classifier (batch + single-record), same core function.",
         "Future-state: procurement analysts paste PO descriptions into a tool, get instant category + confidence."),
        ("6", "Externalized rule files (CSVs).",
         "Procurement staff can update rules without touching Python code."),
        ("7", "Feedback loop: analyst overrides become new rule data.",
         "AI/rule usage shrinks over time as the rule base grows."),
        ("8", "Classify all 784,556 rows regardless of date.",
         "Classifier is time-agnostic; date filtering is a downstream concern."),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Decision"
    hdr[2].text = "Rationale"
    style_table_header(table.rows[0])
    for i, (n, dec, rat) in enumerate(decisions):
        r = table.add_row().cells
        r[0].text = n
        r[1].text = dec
        r[2].text = rat
        for cell in r:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        if i % 2 == 0:
            for cell in r:
                set_cell_shading(cell, CHI_LT_BLUE_HEX)
    for row in table.rows:
        row.cells[0].width = Cm(1.0)
        row.cells[1].width = Cm(8.0)
        row.cells[2].width = Cm(8.0)

    # End of document line
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("End of methodology document.")
    run.font.italic = True
    run.font.color.rgb = GRAY_RGB
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out = OUT_DIR / "NIGP_Methodology_for_Leadership_JHK3.docx"
    doc.save(out)
    return out


# =========================================================================
def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    xlsx = build_excel()
    brief = build_executive_brief()
    method = build_methodology()
    print("Built:")
    print(f"  {xlsx}")
    print(f"  {brief}")
    print(f"  {method}")


if __name__ == "__main__":
    main()
