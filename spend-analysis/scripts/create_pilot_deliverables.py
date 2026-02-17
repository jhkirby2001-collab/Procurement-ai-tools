#!/usr/bin/env python3
"""
Create both deliverables for the drinking water vendor consolidation pilot:
1. One-page business case (Word .docx)
2. Comprehensive savings analysis workbook (Excel .xlsx)
"""

import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.utils import get_column_letter

# City of Chicago colors
NAVY_BLUE = RGBColor(0, 51, 102)  # #003366
LIGHT_BLUE = RGBColor(0, 113, 188)  # #0071BC

# Load the processed data
vendor_dept_spend = pd.read_csv('/workspaces/Spend_Analysis_and_Categorization/spend-analysis/data/processed/pilot_vendor_dept_spend_matrix.csv', index_col=0)
vendor_summary = pd.read_csv('/workspaces/Spend_Analysis_and_Categorization/spend-analysis/data/processed/pilot_vendor_summary.csv', index_col=0)
dept_summary = pd.read_csv('/workspaces/Spend_Analysis_and_Categorization/spend-analysis/data/processed/pilot_dept_summary.csv', index_col=0)

# Constants from verified data
TOTAL_5YR_SPEND = 700404.52
ANNUAL_RUN_RATE = TOTAL_5YR_SPEND / 5
TOTAL_TRANSACTIONS = 240
ANNUAL_TRANSACTIONS = TOTAL_TRANSACTIONS / 5
UNIQUE_VENDORS = 4
UNIQUE_DEPARTMENTS = 13

print("=" * 80)
print("CREATING VENDOR CONSOLIDATION PILOT DELIVERABLES")
print("=" * 80)
print()

# ============================================================================
# DELIVERABLE 1: ONE-PAGE BUSINESS CASE (WORD DOCUMENT)
# ============================================================================

print("Creating Deliverable 1: One-Page Business Case (.docx)...")
print()

doc = Document()

# Set document margins (narrow)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Header
header = doc.add_paragraph()
header_run = header.add_run("Department of Procurement Services — Strategic Sourcing Initiative")
header_run.font.size = Pt(8)
header_run.font.color.rgb = NAVY_BLUE
header.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Title
title = doc.add_heading('Vendor Consolidation Pilot: Drinking Water Services', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = NAVY_BLUE
title_run.font.size = Pt(16)
title_run.font.bold = True

# Subtitle
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('90-Day Proof of Concept for Immediate Cost Savings')
subtitle_run.font.size = Pt(11)
subtitle_run.font.color.rgb = LIGHT_BLUE
subtitle_run.font.italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # spacing

# Section 1: The Opportunity
heading1 = doc.add_heading('The Opportunity', level=2)
heading1.runs[0].font.color.rgb = NAVY_BLUE
heading1.runs[0].font.size = Pt(12)

p1 = doc.add_paragraph()
p1.add_run("Thirteen City departments independently purchase bottled water and water cooler delivery services through Exhibit B (non-contract spend). Four different vendors are providing the identical commodity service across 240 transactions over five years, creating significant administrative overhead and fragmented pricing. Three departments—Office of Public Safety Administration, Chicago Police, and Chicago Fire—are each using multiple vendors for the same commodity, demonstrating the exact fragmentation pattern a consolidated contract eliminates. This is a textbook commodity consolidation: water delivery is an interchangeable service with zero operational risk to standardize.")
p1.paragraph_format.space_after = Pt(8)

# Section 2: Current State
heading2 = doc.add_heading('Current State', level=2)
heading2.runs[0].font.color.rgb = NAVY_BLUE
heading2.runs[0].font.size = Pt(12)

# Create table
table1 = doc.add_table(rows=6, cols=2)
table1.style = 'Light Grid Accent 1'

# Populate current state table
table1.rows[0].cells[0].text = "Metric"
table1.rows[0].cells[1].text = "Current State"
table1.rows[1].cells[0].text = "Annual Spend"
table1.rows[1].cells[1].text = f"~${ANNUAL_RUN_RATE:,.0f}"
table1.rows[2].cells[0].text = "Vendors Used"
table1.rows[2].cells[1].text = "4"
table1.rows[3].cells[0].text = "Departments Buying Independently"
table1.rows[3].cells[1].text = "13"
table1.rows[4].cells[0].text = "Annual Transactions"
table1.rows[4].cells[1].text = "~48"
table1.rows[5].cells[0].text = "Administrative Cost per Transaction"
table1.rows[5].cells[1].text = "Est. $150-250*"

# Format table header
for cell in table1.rows[0].cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = NAVY_BLUE

footnote = doc.add_paragraph()
footnote_run = footnote.add_run("*Note: Transaction cost estimate based on industry benchmarks from NIGP for government purchase order processing")
footnote_run.font.size = Pt(8)
footnote_run.font.italic = True
footnote.paragraph_format.space_after = Pt(8)

# Section 3: Using Areas
heading3 = doc.add_heading('Using Areas — Who Is Buying What From Whom', level=2)
heading3.runs[0].font.color.rgb = NAVY_BLUE
heading3.runs[0].font.size = Pt(12)

intro3 = doc.add_paragraph()
intro3.add_run("This table shows the 13 departments, their vendor relationships, and total spend. ").bold = False
intro3.add_run("Three departments (bold) are using 2-3 vendors each for the same commodity").bold = True
intro3.add_run(" — the City is processing multiple POs, multiple invoices, and managing multiple vendor relationships for bottled water within the same department clusters.")
intro3.paragraph_format.space_after = Pt(6)

# Create vendor-department matrix table
table2 = doc.add_table(rows=15, cols=7)
table2.style = 'Light Grid Accent 1'

# Header row
headers = ["Department", "DS Services", "BLUETRITON", "Hinckley", "BrewSmart", "Total Spend", "Vendors"]
for i, header in enumerate(headers):
    cell = table2.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = NAVY_BLUE
    cell.paragraphs[0].runs[0].font.size = Pt(9)

# Department data (sorted by spend descending)
dept_data = [
    ("59-Chicago Fire Department", "✓", "✓", "", "", "$487,713", "2", True),
    ("85-Chicago Dept of Aviation", "", "✓", "", "", "$88,091", "1", False),
    ("81-Dept of Streets & Sanitation", "", "✓", "", "", "$38,168", "1", False),
    ("50-Dept of Family & Support Services", "", "✓", "", "", "$35,145", "1", False),
    ("51-Office of Public Safety Admin", "✓", "✓", "✓", "", "$16,310", "3", True),
    ("88-Dept of Water Management", "", "✓", "", "", "$15,688", "1", False),
    ("57-Chicago Police Department", "✓", "✓", "✓", "", "$10,288", "3", True),
    ("84-Chicago Dept of Transportation", "", "✓", "", "", "$4,569", "1", False),
    ("73-Chicago Animal Care & Control", "", "✓", "", "", "$2,447", "1", False),
    ("31-Department of Law", "", "", "", "✓", "$666", "1", False),
    ("70-Business Affairs & Consumer Protection", "", "✓", "", "", "$501", "1", False),
    ("35-Dept of Procurement Services", "", "✓", "", "", "$422", "1", False),
    ("60-Civilian Office of Police Accountability", "", "✓", "", "", "$397", "1", False),
    ("TOTAL", "$488,674", "$197,044", "$14,020", "$666", "$700,405", "", False),
]

for row_idx, dept_row in enumerate(dept_data, start=1):
    is_bold = dept_row[7]
    for col_idx, value in enumerate(dept_row[:7]):
        cell = table2.rows[row_idx].cells[col_idx]
        cell.text = value
        if is_bold or row_idx == len(dept_data):  # Bold for multi-vendor depts and total row
            cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)

# Key insight paragraph
insight = doc.add_paragraph()
insight_run = insight.add_run("Key Insight: ")
insight_run.font.bold = True
insight.add_run("Public Safety Administration, CPD, and CFD (bold rows) collectively use 2-3 vendors each. Chicago Fire alone spent $487,713 and uses two different vendors for identical drinking water services. This fragmentation means the City is processing 28 separate transactions for CFD water instead of a single quarterly contract invoice.")
insight.paragraph_format.space_after = Pt(8)

# Section 4: Proposed Pilot
heading4 = doc.add_heading('Proposed Pilot', level=2)
heading4.runs[0].font.color.rgb = NAVY_BLUE
heading4.runs[0].font.size = Pt(12)

p4 = doc.add_paragraph("Consolidate all 13 departments to 1-2 preferred vendors through a single citywide contract. Route all drinking water purchases through the contract instead of individual Exhibit B requisitions. The pilot requires a 90-day implementation timeline with no new budget—the CPO can authorize under existing procurement authority. This eliminates maverick spend, standardizes pricing, and dramatically reduces administrative burden.")
p4.paragraph_format.space_after = Pt(8)

# Section 5: Projected Impact
heading5 = doc.add_heading('Projected Impact', level=2)
heading5.runs[0].font.color.rgb = NAVY_BLUE
heading5.runs[0].font.size = Pt(12)

table3 = doc.add_table(rows=5, cols=3)
table3.style = 'Light Grid Accent 1'

# Header
table3.rows[0].cells[0].text = "Benefit"
table3.rows[0].cells[1].text = "Conservative (5%)"
table3.rows[0].cells[2].text = "Moderate (10%)"

# Data
table3.rows[1].cells[0].text = "Annual Price Savings"
table3.rows[1].cells[1].text = "$5,953"
table3.rows[1].cells[2].text = "$11,907"

table3.rows[2].cells[0].text = "Transaction Reduction"
table3.rows[2].cells[1].text = "~48 → ~12 per year"
table3.rows[2].cells[2].text = "~48 → ~12 per year"

table3.rows[3].cells[0].text = "Admin Cost Savings"
table3.rows[3].cells[1].text = "$5,400 - $9,000"
table3.rows[3].cells[2].text = "$5,400 - $9,000"

table3.rows[4].cells[0].text = "Total Annual Benefit"
table3.rows[4].cells[1].text = "$11,353 - $14,953"
table3.rows[4].cells[2].text = "$17,307 - $20,907"

# Format header
for cell in table3.rows[0].cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = NAVY_BLUE

# Bold totals row
for cell in table3.rows[4].cells:
    cell.paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# Section 6: Why This Pilot
heading6 = doc.add_heading('Why This Pilot', level=2)
heading6.runs[0].font.color.rgb = NAVY_BLUE
heading6.runs[0].font.size = Pt(12)

p6_1 = doc.add_paragraph(style='List Bullet')
p6_1.add_run("Zero operational risk: ").bold = True
p6_1.add_run("Water delivery is a commodity—service quality is identical across vendors")

p6_2 = doc.add_paragraph(style='List Bullet')
p6_2.add_run("Proves the model: ").bold = True
p6_2.add_run("If consolidation works here, the same approach applies to the 508 tail vendors representing $59M in Exhibit B spend")

p6_3 = doc.add_paragraph(style='List Bullet')
p6_3.add_run("Reduces maverick spend: ").bold = True
p6_3.add_run("Moves $140K/year from non-contract (Exhibit B) to managed contract spend, directly addressing EY Efficiency Review recommendations")

doc.add_paragraph()

# Section 7: Recommended Next Steps
heading7 = doc.add_heading('Recommended Next Steps', level=2)
heading7.runs[0].font.color.rgb = NAVY_BLUE
heading7.runs[0].font.size = Pt(12)

steps = [
    ("CPO approves 90-day pilot", "for drinking water services consolidation"),
    ("Issue RFQ", "to 2-3 qualified water delivery vendors for citywide contract"),
    ("Direct all 13 departments", "to use new contract — report savings at 90 days"),
    ("Evaluate Phase 2 sustainability transition:", "Use consolidated contract data to conduct cost-benefit analysis for replacing bottled water with filtration systems/water stations. Aligns with national municipal trend toward eliminating single-use plastics and reducing environmental impact while achieving long-term cost savings")
]

for i, (bold_text, normal_text) in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(f"{bold_text} ").bold = True
    p.add_run(normal_text)

doc.add_paragraph()

# Footer
footer_para = doc.add_paragraph()
footer_run = footer_para.add_run("Prepared by James Kirby III, CSCP, MS-SCM | Department of Procurement Services | February 2026")
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = NAVY_BLUE
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save Word document
word_output = Path("/workspaces/Spend_Analysis_and_Categorization/spend-analysis/outputs/DRINKING_WATER_PILOT_BUSINESS_CASE.docx")
doc.save(word_output)
print(f"✓ Saved: {word_output.name}")
print()

# ============================================================================
# DELIVERABLE 2: COMPREHENSIVE SAVINGS ANALYSIS WORKBOOK (EXCEL)
# ============================================================================

print("Creating Deliverable 2: Comprehensive Savings Analysis Workbook (.xlsx)...")
print()

# Create new workbook
wb = openpyxl.Workbook()
wb.remove(wb.active)  # Remove default sheet

# Styling
navy_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
light_blue_fill = PatternFill(start_color="0071BC", end_color="0071BC", fill_type="solid")
light_gray_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
yellow_fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
green_fill = PatternFill(start_color="D9EAD3", end_color="D9EAD3", fill_type="solid")

white_font = Font(color="FFFFFF", bold=True, size=11)
bold_font = Font(bold=True, size=11)
navy_font = Font(color="003366", bold=True, size=11)

center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
right_align = Alignment(horizontal="right", vertical="center")

thin_border = Border(
    left=Side(style='thin'),
    right=Side(style='thin'),
    top=Side(style='thin'),
    bottom=Side(style='thin')
)

def format_header_row(ws, row_num, columns):
    """Format a header row with navy blue background and white text"""
    for col_num in range(1, columns + 1):
        cell = ws.cell(row=row_num, column=col_num)
        cell.fill = navy_fill
        cell.font = white_font
        cell.alignment = center_align
        cell.border = thin_border

def format_currency(ws, row, col):
    """Format cell as currency"""
    cell = ws.cell(row=row, column=col)
    cell.number_format = '$#,##0.00'
    cell.border = thin_border

def format_number(ws, row, col):
    """Format cell as number with commas"""
    cell = ws.cell(row=row, column=col)
    cell.number_format = '#,##0'
    cell.border = thin_border

def format_percent(ws, row, col):
    """Format cell as percentage"""
    cell = ws.cell(row=row, column=col)
    cell.number_format = '0.0%'
    cell.border = thin_border

# ============================================================================
# SHEET 1: Executive Summary
# ============================================================================
print("  Creating Sheet 1: Executive Summary...")

ws1 = wb.create_sheet("Executive Summary")

# Title
ws1['A1'] = "DRINKING WATER SERVICES VENDOR CONSOLIDATION PILOT"
ws1['A1'].font = Font(color="003366", bold=True, size=14)
ws1.merge_cells('A1:D1')

ws1['A2'] = "Executive Summary"
ws1['A2'].font = navy_font
ws1.merge_cells('A2:D2')

ws1.append([])  # blank row

# Key Metrics
ws1.append(["KEY METRICS", "", "", ""])
ws1['A4'].font = bold_font
ws1.merge_cells('A4:D4')

metrics = [
    ["Metric", "Value"],
    ["Total 5-Year Spend", TOTAL_5YR_SPEND],
    ["Annual Run Rate", ANNUAL_RUN_RATE],
    ["Total Transactions", TOTAL_TRANSACTIONS],
    ["Total Vendors", UNIQUE_VENDORS],
    ["Total Departments (Using Areas)", UNIQUE_DEPARTMENTS],
    ["Average Transaction Size", TOTAL_5YR_SPEND / TOTAL_TRANSACTIONS],
]

start_row = 5
for i, row in enumerate(metrics, start=start_row):
    ws1.cell(row=i, column=1, value=row[0])
    ws1.cell(row=i, column=2, value=row[1])
    if i > start_row:  # Format values (not header)
        format_currency(ws1, i, 2)

format_header_row(ws1, start_row, 2)

ws1.append([])
ws1.append([])

# Recommendation
ws1.append(["RECOMMENDATION", "", "", ""])
ws1[f'A{ws1.max_row}'].font = bold_font
ws1.merge_cells(f'A{ws1.max_row}:D{ws1.max_row}')

ws1.append(["Recommended Action:", "Consolidate to 1-2 vendors via single citywide contract"])
ws1[f'A{ws1.max_row}'].font = bold_font

ws1.append(["Implementation Timeline:", "90 days"])
ws1[f'A{ws1.max_row}'].font = bold_font

ws1.append(["Projected Annual Benefit Range:", "$11,353 - $20,907"])
ws1[f'A{ws1.max_row}'].font = bold_font

# Adjust column widths
ws1.column_dimensions['A'].width = 35
ws1.column_dimensions['B'].width = 20

# ============================================================================
# SHEET 2: Vendor Detail
# ============================================================================
print("  Creating Sheet 2: Vendor Detail...")

ws2 = wb.create_sheet("Vendor Detail")

ws2['A1'] = "VENDOR-LEVEL DETAIL"
ws2['A1'].font = Font(color="003366", bold=True, size=14)
ws2.merge_cells('A1:G1')

ws2.append([])

# Create vendor detail data
vendor_detail_data = [
    ["Vendor Name", "Total Spend", "% of Category", "Transactions", "Departments Served", "Department Names", "Consolidation Recommendation"],
    ["DS Services of America, Inc.", 488674.29, 0.698, 30, 3, "51-Public Safety Admin, 57-Chicago Police, 59-Chicago Fire", "Primary vendor candidate - largest share, serves public safety cluster"],
    ["BLUETRITON BRANDS INC", 197043.87, 0.281, 174, 12, "35-Procurement, 50-Family Support, 51-Public Safety, 57-Police, 59-Fire, 60-COPA, 70-Business Affairs, 73-Animal Care, 81-Streets, 84-Transportation, 85-Aviation, 88-Water Mgmt", "Primary or secondary vendor - widest coverage, serves 12 departments"],
    ["HINCKLEY SPRINGS", 14020.36, 0.020, 35, 2, "51-Public Safety Admin, 57-Chicago Police", "Absorb into primary vendor - limited coverage, only 2% of spend"],
    ["BREWSMART BEVERAGE", 666.00, 0.001, 1, 1, "31-Department of Law", "Eliminate - negligible spend (<0.1%), single transaction"],
]

for row in vendor_detail_data:
    ws2.append(row)

# Format header
format_header_row(ws2, 3, 7)

# Format data rows
for row_idx in range(4, 8):
    format_currency(ws2, row_idx, 2)  # Total Spend
    format_percent(ws2, row_idx, 3)   # % of Category
    format_number(ws2, row_idx, 4)    # Transactions
    format_number(ws2, row_idx, 5)    # Departments

# Adjust column widths
ws2.column_dimensions['A'].width = 30
ws2.column_dimensions['B'].width = 15
ws2.column_dimensions['C'].width = 15
ws2.column_dimensions['D'].width = 15
ws2.column_dimensions['E'].width = 18
ws2.column_dimensions['F'].width = 60
ws2.column_dimensions['G'].width = 50

# ============================================================================
# SHEET 3: Department Detail (Using Areas)
# ============================================================================
print("  Creating Sheet 3: Department Detail (Using Areas)...")

ws3 = wb.create_sheet("Department Detail (Using Areas)")

ws3['A1'] = "DEPARTMENT-LEVEL DETAIL (USING AREAS)"
ws3['A1'].font = Font(color="003366", bold=True, size=14)
ws3.merge_cells('A1:G1')

ws3.append([])

# Headers
headers3 = ["Department Code & Name", "Total Spend", "Transactions", "Vendors Used", "Vendor Names", "Avg Transaction", "Fragmentation Level"]
ws3.append(headers3)
format_header_row(ws3, 3, 7)

# Department data (sorted by spend)
dept_detail_data = [
    ["59-CHICAGO FIRE DEPARTMENT", 487713.49, 28, 2, "DS Services, BLUETRITON", 17418.34, "MODERATE"],
    ["85-CHICAGO DEPARTMENT OF AVIATION", 88091.29, 33, 1, "BLUETRITON", 2669.43, "Single vendor"],
    ["81-DEPARTMENT OF STREETS AND SANITATION", 38167.54, 76, 1, "BLUETRITON", 502.20, "Single vendor"],
    ["50-DEPARTMENT OF FAMILY AND SUPPORT SERVICES", 35144.62, 8, 1, "BLUETRITON", 4393.08, "Single vendor"],
    ["51-OFFICE OF PUBLIC SAFETY ADMINISTRATION", 16309.83, 27, 3, "DS Services, BLUETRITON, Hinckley", 604.07, "HIGH"],
    ["88-DEPARTMENT OF WATER MANAGEMENT", 15687.51, 2, 1, "BLUETRITON", 7843.76, "Single vendor"],
    ["57-CHICAGO POLICE DEPARTMENT", 10288.39, 15, 3, "DS Services, BLUETRITON, Hinckley", 685.89, "HIGH"],
    ["84-CHICAGO DEPARTMENT OF TRANSPORTATION", 4568.66, 14, 1, "BLUETRITON", 326.33, "Single vendor"],
    ["73-CHICAGO ANIMAL CARE AND CONTROL", 2447.28, 5, 1, "BLUETRITON", 489.46, "Single vendor"],
    ["31-DEPARTMENT OF LAW", 666.00, 1, 1, "BrewSmart", 666.00, "Single vendor"],
    ["70-DEPARTMENT OF BUSINESS AFFAIRS AND CONSUMER PROTECTION", 501.29, 8, 1, "BLUETRITON", 62.66, "Single vendor"],
    ["35-DEPARTMENT OF PROCUREMENT SERVICES", 421.50, 21, 1, "BLUETRITON", 20.07, "Single vendor"],
    ["60-CIVILIAN OFFICE OF POLICE ACCOUNTABILITY", 397.12, 2, 1, "BLUETRITON", 198.56, "Single vendor"],
]

for row_data in dept_detail_data:
    ws3.append(row_data)

# Format data
for row_idx in range(4, 4 + len(dept_detail_data)):
    format_currency(ws3, row_idx, 2)  # Total Spend
    format_number(ws3, row_idx, 3)    # Transactions
    format_number(ws3, row_idx, 4)    # Vendors Used
    format_currency(ws3, row_idx, 6)  # Avg Transaction

    # Highlight departments using multiple vendors
    vendors_used = ws3.cell(row=row_idx, column=4).value
    if vendors_used > 1:
        for col_idx in range(1, 8):
            ws3.cell(row=row_idx, column=col_idx).fill = yellow_fill

# Summary section
ws3.append([])
ws3.append(["FRAGMENTATION SUMMARY", "", "", "", "", "", ""])
ws3[f'A{ws3.max_row}'].font = bold_font

ws3.append(["Departments using 3 vendors (HIGHEST fragmentation):", "51-Public Safety Admin, 57-Chicago Police", "", "", "", "", ""])
ws3.append(["Departments using 2 vendors (MODERATE fragmentation):", "59-Chicago Fire", "", "", "", "", ""])
ws3.append(["Departments using 1 vendor (Still maverick spend):", "10 departments (all others)", "", "", "", "", ""])

# Adjust column widths
ws3.column_dimensions['A'].width = 45
ws3.column_dimensions['B'].width = 15
ws3.column_dimensions['C'].width = 15
ws3.column_dimensions['D'].width = 15
ws3.column_dimensions['E'].width = 40
ws3.column_dimensions['F'].width = 18
ws3.column_dimensions['G'].width = 20

# ============================================================================
# SHEET 4: Vendor-Department Matrix
# ============================================================================
print("  Creating Sheet 4: Vendor-Department Matrix...")

ws4 = wb.create_sheet("Vendor-Department Matrix")

ws4['A1'] = "VENDOR-DEPARTMENT SPEND MATRIX"
ws4['A1'].font = Font(color="003366", bold=True, size=14)
ws4.merge_cells('A1:F1')

ws4.append([])

# Headers
matrix_headers = ["Department", "DS Services", "BLUETRITON", "Hinckley", "BrewSmart", "TOTAL"]
ws4.append(matrix_headers)
format_header_row(ws4, 3, 6)

# Matrix data
matrix_data = [
    ["59-CHICAGO FIRE DEPARTMENT", 486008.03, 1705.46, 0, 0, 487713.49],
    ["85-CHICAGO DEPARTMENT OF AVIATION", 0, 88091.29, 0, 0, 88091.29],
    ["81-DEPARTMENT OF STREETS AND SANITATION", 0, 38167.54, 0, 0, 38167.54],
    ["50-DEPARTMENT OF FAMILY AND SUPPORT SERVICES", 0, 35144.62, 0, 0, 35144.62],
    ["51-OFFICE OF PUBLIC SAFETY ADMINISTRATION", 2116.23, 1607.67, 12585.93, 0, 16309.83],
    ["88-DEPARTMENT OF WATER MANAGEMENT", 0, 15687.51, 0, 0, 15687.51],
    ["57-CHICAGO POLICE DEPARTMENT", 550.03, 8303.93, 1434.43, 0, 10288.39],
    ["84-CHICAGO DEPARTMENT OF TRANSPORTATION", 0, 4568.66, 0, 0, 4568.66],
    ["73-CHICAGO ANIMAL CARE AND CONTROL", 0, 2447.28, 0, 0, 2447.28],
    ["31-DEPARTMENT OF LAW", 0, 0, 0, 666.00, 666.00],
    ["70-DEPARTMENT OF BUSINESS AFFAIRS AND CONSUMER PROTECTION", 0, 501.29, 0, 0, 501.29],
    ["35-DEPARTMENT OF PROCUREMENT SERVICES", 0, 421.50, 0, 0, 421.50],
    ["60-CIVILIAN OFFICE OF POLICE ACCOUNTABILITY", 0, 397.12, 0, 0, 397.12],
    ["TOTAL", 488674.29, 197043.87, 14020.36, 666.00, 700404.52],
]

for row_data in matrix_data:
    ws4.append(row_data)

# Format all currency cells and add color coding
for row_idx in range(4, 4 + len(matrix_data)):
    for col_idx in range(2, 7):
        format_currency(ws4, row_idx, col_idx)

        # Green fill for cells with spend > 0 (except total row and column)
        if row_idx < 4 + len(matrix_data) - 1 and col_idx < 6:
            if ws4.cell(row=row_idx, column=col_idx).value > 0:
                ws4.cell(row=row_idx, column=col_idx).fill = green_fill

# Bold total row
for col_idx in range(1, 7):
    ws4.cell(row=4 + len(matrix_data) - 1, column=col_idx).font = bold_font

# Adjust column widths
ws4.column_dimensions['A'].width = 45
for col in ['B', 'C', 'D', 'E', 'F']:
    ws4.column_dimensions[col].width = 18

# ============================================================================
# SHEET 5: Addressable Spend Calculation
# ============================================================================
print("  Creating Sheet 5: Addressable Spend Calculation...")

ws5 = wb.create_sheet("Addressable Spend Calculation")

ws5['A1'] = "ADDRESSABLE SPEND CALCULATION"
ws5['A1'].font = Font(color="003366", bold=True, size=14)
ws5.merge_cells('A1:E1')

ws5.append([])

ws5.append(["Step-by-step calculation showing how we determine addressable spend for consolidation savings"])
ws5.merge_cells('A3:E3')

ws5.append([])

# Headers
calc_headers = ["Step", "Description", "Calculation", "Amount", "Explanation"]
ws5.append(calc_headers)
format_header_row(ws5, 5, 5)

# Calculation steps
calc_steps = [
    [1, "Total 5-Year Spend", "From data", TOTAL_5YR_SPEND, "All drinking water/cooler Exhibit B transactions, Jan 2021 - Jan 2026"],
    [2, "Annual Run Rate", "$700,404 ÷ 5", ANNUAL_RUN_RATE, "Annualized baseline — this is what the City spends per year"],
    [3, "Already Optimized (15%)", "$140,081 × 15%", ANNUAL_RUN_RATE * 0.15, "Industry benchmark (McKinsey, NIGP): 10-20% of spend not addressable through consolidation. Includes: mandated vendor requirements, emergency purchases, regulatory constraints. We use 15% as conservative mid-point."],
    [4, "Annual Addressable Spend", "$140,081 - $21,012", ANNUAL_RUN_RATE * 0.85, "This is the spend we CAN improve through vendor consolidation. Addressable spend is standard procurement terminology used by Fortune 500 companies, EY, McKinsey, NIGP, and CIPS."],
    ["", "WHY ANNUALIZE FIRST?", "", "", "We annualize BEFORE applying savings rates because we are projecting forward-looking annual savings from FUTURE contracts, not retroactive savings on historical spend. The 5-year data establishes the annual run rate baseline. Savings come from contracts we negotiate going forward."],
]

row_num = 6
for step_data in calc_steps:
    ws5.append(step_data)
    if step_data[0] != "":  # Not the explanation row
        if isinstance(step_data[0], int):
            format_number(ws5, row_num, 1)
        format_currency(ws5, row_num, 4)
    row_num += 1

# Adjust column widths
ws5.column_dimensions['A'].width = 8
ws5.column_dimensions['B'].width = 30
ws5.column_dimensions['C'].width = 25
ws5.column_dimensions['D'].width = 18
ws5.column_dimensions['E'].width = 80

# Wrap text for explanation column
for row_idx in range(6, row_num):
    ws5.cell(row=row_idx, column=5).alignment = left_align

# ============================================================================
# SHEET 6: Savings Projections
# ============================================================================
print("  Creating Sheet 6: Savings Projections...")

ws6 = wb.create_sheet("Savings Projections")

ws6['A1'] = "SAVINGS PROJECTIONS"
ws6['A1'].font = Font(color="003366", bold=True, size=14)
ws6.merge_cells('A1:D1')

ws6.append([])
ws6.append([])

# Consolidation Savings
ws6.append(["CONSOLIDATION SAVINGS", "", "", ""])
ws6['A4'].font = bold_font
ws6.merge_cells('A4:D4')

addressable_spend = ANNUAL_RUN_RATE * 0.85

cons_headers = ["Scenario", "Savings Rate", "Calculation", "Annual Savings"]
ws6.append(cons_headers)
format_header_row(ws6, 5, 4)

consolidation_scenarios = [
    ["Conservative", "5%", f"${addressable_spend:,.2f} × 5%", addressable_spend * 0.05],
    ["Moderate (RECOMMENDED)", "10%", f"${addressable_spend:,.2f} × 10%", addressable_spend * 0.10],
    ["Aggressive", "15%", f"${addressable_spend:,.2f} × 15%", addressable_spend * 0.15],
]

for row_data in consolidation_scenarios:
    ws6.append(row_data)

for row_idx in range(6, 9):
    format_currency(ws6, row_idx, 4)

ws6.append([])
ws6.append(["WHY 5-15% RANGE:", "Industry benchmarks from McKinsey and NIGP show 5-15% savings typical for vendor consolidation programs. Reducing vendors increases volume per vendor = stronger negotiating position = better pricing."])
ws6.merge_cells(f'A10:D10')

ws6.append(["WHY 10% MODERATE:", "Conservative enough to be achievable, aggressive enough to justify the effort. This is the same rate EY used in their efficiency review analysis."])
ws6.merge_cells(f'A11:D11')

ws6.append([])
ws6.append([])

# Cost Avoidance
ws6.append(["COST AVOIDANCE (INFLATION PROTECTION)", "", "", ""])
ws6['A14'].font = bold_font
ws6.merge_cells('A14:D14')

ws6.append(["Metric", "Calculation", "", "Amount"])
format_header_row(ws6, 15, 4)

cost_avoid = [
    ["Annual Spend (total, not addressable)", f"${TOTAL_5YR_SPEND:,.2f} ÷ 5", "", ANNUAL_RUN_RATE],
    ["Inflation Rate", "BLS Producer Price Index", "", 0.05],
    ["Annual Cost Avoidance", f"${ANNUAL_RUN_RATE:,.2f} × 5%", "", ANNUAL_RUN_RATE * 0.05],
]

for row_data in cost_avoid:
    ws6.append(row_data)

format_currency(ws6, 16, 4)
format_percent(ws6, 17, 4)
format_currency(ws6, 18, 4)

ws6.append([])
ws6.append(["EXPLANATION:", "Cost avoidance uses TOTAL spend (not addressable) because inflation affects ALL contracts whether or not we consolidate them. Strategic contracts with price protection clauses shield the entire spend base from increases. This is different from consolidation savings which only come from contracts we can improve."])
ws6.merge_cells(f'A20:D20')

ws6.append([])
ws6.append([])

# Administrative Savings
ws6.append(["ADMINISTRATIVE COST SAVINGS", "", "", ""])
ws6['A23'].font = bold_font
ws6.merge_cells('A23:D23')

ws6.append(["Metric", "Calculation", "", "Amount"])
format_header_row(ws6, 24, 4)

admin_savings = [
    ["Current Annual Transactions", "240 ÷ 5", "", ANNUAL_TRANSACTIONS],
    ["Projected Transactions (consolidated)", "12 (quarterly per vendor)", "", 12],
    ["Transactions Eliminated", "48 - 12", "", 36],
    ["Cost per Transaction (NIGP benchmark)", "$150 - $250", "", ""],
    ["Annual Admin Savings", "36 × $150 to 36 × $250", "", "$5,400 - $9,000"],
]

for row_data in admin_savings:
    ws6.append(row_data)

for row_idx in [25, 26, 27]:
    format_number(ws6, row_idx, 4)

ws6.append([])
ws6.append(["EXPLANATION:", "Transaction processing costs include staff time for requisition creation, PO generation, invoice matching, payment processing, and vendor management. NIGP benchmarks for government purchase order processing range $150-$250 per transaction."])
ws6.merge_cells(f'A31:D31')

ws6.append([])
ws6.append([])

# Total Benefit Summary
ws6.append(["TOTAL BENEFIT SUMMARY", "", "", ""])
ws6['A34'].font = bold_font
ws6.merge_cells('A34:D34')

ws6.append(["Benefit Type", "Conservative", "Moderate", ""])
format_header_row(ws6, 35, 4)

total_benefit = [
    ["Consolidation Savings", 5953, 11907, ""],
    ["Admin Savings (low)", 5400, 5400, ""],
    ["Annual Total", 11353, 17307, ""],
    ["3-Year Total", 34060, 51920, ""],
]

for row_data in total_benefit:
    ws6.append(row_data)

for row_idx in range(36, 40):
    format_currency(ws6, row_idx, 2)
    format_currency(ws6, row_idx, 3)

# Bold totals
for col_idx in [1, 2, 3]:
    ws6.cell(row=38, column=col_idx).font = bold_font
    ws6.cell(row=39, column=col_idx).font = bold_font

# Adjust column widths
ws6.column_dimensions['A'].width = 40
ws6.column_dimensions['B'].width = 25
ws6.column_dimensions['C'].width = 25
ws6.column_dimensions['D'].width = 18

# ============================================================================
# SHEET 7: Methodology & Assumptions
# ============================================================================
print("  Creating Sheet 7: Methodology & Assumptions...")

ws7 = wb.create_sheet("Methodology & Assumptions")

ws7['A1'] = "METHODOLOGY & ASSUMPTIONS"
ws7['A1'].font = Font(color="003366", bold=True, size=14)
ws7.merge_cells('A1:B1')

ws7.append([])

ws7.append(["Assumption", "Source / Justification"])
format_header_row(ws7, 3, 2)

assumptions = [
    ["15% already optimized exclusion", "McKinsey & Company, NIGP Best Practices in Category Management (10-20% range, 15% mid-point)"],
    ["5-15% consolidation savings range", "Industry benchmark for vendor consolidation (McKinsey, NIGP, CIPS)"],
    ["10% moderate scenario recommended", "Same rate used by EY in City of Chicago Efficiency Review"],
    ["5% inflation/cost avoidance rate", "Bureau of Labor Statistics Producer Price Index for services"],
    ["$150-$250 transaction cost", "NIGP government purchase order processing benchmarks"],
    ["Annualize-first methodology", "Forward-looking projection from future contracts, not retroactive savings on historical spend"],
    ["5-year data period", "January 2021 - January 2026 City of Chicago Exhibit B Spend Report"],
    ["Addressable spend definition", "Standard procurement terminology (Fortune 500, McKinsey, EY, NIGP, CIPS) for spend that can be improved through strategic sourcing"],
    ["90-day implementation", "Realistic timeline for RFQ, vendor selection, contract setup, department communication"],
    ["Zero operational risk", "Drinking water/cooler delivery is commoditized service with multiple qualified vendors in Chicago market"],
]

for row_data in assumptions:
    ws7.append(row_data)

# Adjust column widths
ws7.column_dimensions['A'].width = 40
ws7.column_dimensions['B'].width = 100

# Wrap text
for row_idx in range(4, 4 + len(assumptions)):
    ws7.cell(row=row_idx, column=2).alignment = left_align

# ============================================================================
# SHEET 8: EY Validation
# ============================================================================
print("  Creating Sheet 8: EY Validation...")

ws8 = wb.create_sheet("EY Validation")

ws8['A1'] = "EY EFFICIENCY REVIEW VALIDATION"
ws8['A1'].font = Font(color="003366", bold=True, size=14)
ws8.merge_cells('A1:B1')

ws8.append([])
ws8.append(["How this pilot validates and operationalizes EY Efficiency Review findings"])
ws8.merge_cells('A3:B3')

ws8.append([])

ws8.append(["EY Finding", "This Pilot Validates"])
format_header_row(ws8, 5, 2)

ey_findings = [
    ["Vendor Tail: 97% of vendors = 20% of spend", "4 vendors for one commodity across 13 departments confirms fragmentation"],
    ["Non-Contract (Maverick) Spend", "Exhibit B IS maverick spend — this pilot moves $140K/year to contract"],
    ["Category Management Opportunity", "Drinking water consolidation demonstrates CM methodology on smallest possible scale"],
    ["Vendor Rationalization recommendation", "Reducing 4 vendors to 1-2 is vendor rationalization in action"],
    ["Duplication Across Departments", "13 departments buying same commodity independently — no coordination"],
    ["Administrative Burden", "240 transactions (48/year) reduced to ~12/year saves $5,400-$9,000 annually"],
    ["Savings Opportunity ($55M-$111M citywide)", "This pilot proves the methodology before scaling to $59M tail spend"],
]

for row_data in ey_findings:
    ws8.append(row_data)

ws8.append([])
ws8.append([])

ws8.append(["CRITICAL NOTE:", ""])
ws8['A' + str(ws8.max_row)].font = bold_font
ws8.append(["EY projected $55M-$111M in citywide savings through Category Management. This pilot applies the SAME methodology to the simplest possible category to prove the model before scaling. If consolidation works for drinking water (zero risk, clear commodity), it validates the approach for higher-value categories and the broader vendor tail."])
ws8.merge_cells(f'A{ws8.max_row}:B{ws8.max_row}')

# Adjust column widths
ws8.column_dimensions['A'].width = 50
ws8.column_dimensions['B'].width = 80

# Wrap text
for row_idx in range(6, ws8.max_row + 1):
    ws8.cell(row=row_idx, column=2).alignment = left_align

# ============================================================================
# SHEET 9: Anticipated Questions
# ============================================================================
print("  Creating Sheet 9: Anticipated Questions...")

ws9 = wb.create_sheet("Anticipated Questions")

ws9['A1'] = "ANTICIPATED QUESTIONS & ANSWERS"
ws9['A1'].font = Font(color="003366", bold=True, size=14)
ws9.merge_cells('A1:B1')

ws9.append([])
ws9.append(["Preparation for CPO presentation and stakeholder questions"])
ws9.merge_cells('A3:B3')

ws9.append([])

ws9.append(["Question", "Answer"])
format_header_row(ws9, 5, 2)

qa_pairs = [
    ["Why start with drinking water?", "Zero operational risk — commodity product, interchangeable vendors. Proves methodology before applying to higher-value categories."],

    ["Which departments are affected?", "13 departments currently buying independently. Three (Public Safety Admin, CPD, CFD) use multiple vendors for the same commodity — the worst fragmentation."],

    ["Why annualize first?", "We project forward-looking annual savings from future contracts, not retroactive savings on historical spend. The 5-year data establishes the baseline annual run rate."],

    ["Where does 15% already optimized come from?", "Industry benchmarks (McKinsey, NIGP) suggest 10-20% not addressable. 15% is conservative mid-point for mandated vendors, emergencies, regulatory constraints."],

    ["Why 10% savings rate?", "Industry benchmarks show 5-15% for consolidation. 10% moderate scenario is what EY also used in their efficiency review."],

    ["Is addressable spend a real term?", "Yes — standard procurement terminology used by Fortune 500, McKinsey, EY, NIGP, CIPS for spend improvable through strategic sourcing."],

    ["What if vendors can't serve all locations?", "RFQ will require citywide service capability. Multiple qualified vendors exist in Chicago market (Culligan, Primo, DS Services, etc.)."],

    ["How does this scale?", "Same methodology applies to all 508 tail vendors ($59M). This pilot proves the model on a simple, low-risk category first."],

    ["What about MBE/WBE requirements?", "RFQ will include standard City diversity requirements. Consolidation doesn't eliminate diverse vendors — it creates managed relationships."],

    ["Why are 3 departments using multiple water vendors?", "No central coordination — each department procures independently via Exhibit B. This is exactly the fragmentation pattern EY identified citywide."],

    ["What's the risk if this fails?", "Minimal. Worst case: we return to current state. But failure is unlikely — water delivery is commoditized with multiple qualified vendors."],

    ["How long to see results?", "90 days to implement. First quarterly savings report at 120 days (after first full quarter on contract)."],

    ["What about sustainability and bottled water phase-out?", "This is Phase 1 (consolidation). Phase 2 will use consolidated data to evaluate cost-benefit of water filtration systems vs. bottled water, aligning with municipal trends toward eliminating single-use plastics. Can't properly evaluate alternatives without baseline usage data, which this pilot establishes."],
]

for row_data in qa_pairs:
    ws9.append(row_data)

# Adjust column widths
ws9.column_dimensions['A'].width = 50
ws9.column_dimensions['B'].width = 100

# Wrap text
for row_idx in range(6, 6 + len(qa_pairs)):
    ws9.cell(row=row_idx, column=1).alignment = left_align
    ws9.cell(row=row_idx, column=2).alignment = left_align

# Save Excel workbook
excel_output = Path("/workspaces/Spend_Analysis_and_Categorization/spend-analysis/outputs/DRINKING_WATER_PILOT_ANALYSIS_WORKBOOK.xlsx")
wb.save(excel_output)
print(f"✓ Saved: {excel_output.name}")
print()

print("=" * 80)
print("✓ BOTH DELIVERABLES COMPLETED SUCCESSFULLY")
print("=" * 80)
print()
print("Files ready for download:")
print(f"  1. {word_output}")
print(f"  2. {excel_output}")
print()
print("Next steps:")
print("  • Review both files")
print("  • Verify all numbers trace correctly")
print("  • Present to CPO for pilot approval")
print()
