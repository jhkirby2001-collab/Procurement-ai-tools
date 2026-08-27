"""Build the leadership-facing "Spend Report — How-To & Methodology" Word document
for the NIGP-Sourced Procurement Category Mapper.

Generates:
  outputs/HOW_TO_Spend_Report_JHK3.docx

This is the Word companion to the app's "Spend Report Methodology & How-To" page —
written for leadership and staff who upload a spend file, read the report, and need
to stand behind it. Style/palette are reused from build_leadership_deliverables_JHK3.py
so it matches the other leadership artifacts. DE-BRANDED (independent-tool framing).

Portable: output directory is derived from this file's location, so it runs both in a
codespace and in a fresh clone.

    python spend-analysis/scripts/build_spend_report_howto_docx_JHK3.py
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

# Reuse helpers + brand palette from the leadership-deliverables script.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from build_leadership_deliverables_JHK3 import (  # noqa: E402
    CHI_LT_BLUE_HEX,
    add_body,
    add_bullet,
    add_callout_box,
    add_h1,
    add_h2,
    add_title_banner,
    set_cell_shading,
    set_doc_margins,
    set_table_borders,
    style_table_header,
)

REPO_ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = REPO_ROOT / "outputs"
OUT_PATH = OUT_DIR / "HOW_TO_Spend_Report_JHK3.docx"

DOC_VERSION = "Version 2026-08-20  •  Savings & cost-avoidance edition"
AUTHOR_LINE = (
    "Independent public-sector procurement tool  •  "
    "Prepared by James H. Kirby III, CSCP, MS-SCM  •  " + DOC_VERSION
)


def _shade_alt_rows(table, color_hex=CHI_LT_BLUE_HEX):
    for i, row in enumerate(table.rows[1:]):
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, color_hex)
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)


def _table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    style_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    _shade_alt_rows(table)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    return table


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_doc_margins(doc, cm=1.8)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ------------------------------------------------------------------ Banner
    add_title_banner(
        doc,
        title="Spend Report — How-To & Methodology",
        subtitle="NIGP-Sourced Procurement Category Mapper",
        author_line=AUTHOR_LINE,
    )
    add_body(
        doc,
        "Who this is for: leadership who need to read and stand behind a spend report, and "
        "procurement staff who upload a file and produce one. Everything here is deterministic — "
        "plain arithmetic on your data. No AI runs, and your file never leaves the session.",
    )

    # ------------------------------------------------------------ 0. Summary
    add_callout_box(
        doc,
        label="THE ONE-PARAGRAPH SUMMARY",
        body_text=(
            "Upload a procurement spend file in any format. The tool inspects every column, "
            "figures out what each one is (description, amount, vendor, department, date, or a "
            "breakdown flag), classifies every line into a Business Category using the mapper's "
            "rules, and then runs every spend analysis the data supports — spend by category, "
            "Pareto 80/20, top vendors, consolidation, concentration, tail spend, trend, and more. "
            "It writes a data-driven executive summary and a formatted, multi-tab Excel workbook "
            "with charts. Every row lands in a real Business Category, so nothing reads as a gap, "
            "and every number is reproducible by hand."
        ),
    )

    # ------------------------------------------------------------ 1. What you need
    add_h1(doc, "1.  What you need to run it")
    add_body(
        doc,
        "The only thing the report truly requires is a description column — that is what gets "
        "classified. Everything else unlocks more analysis when present, and is simply skipped "
        "when absent, with a note telling you why.",
    )
    _table(
        doc,
        ["Column", "What it unlocks", "Required?"],
        [
            ["Description", "Classification into Business Categories — the backbone of the report", "Yes"],
            ["Amount / spend", "All dollar analysis (spend by category, Pareto, top vendors by $)", "No"],
            ["Vendor", "Top vendors, consolidation, concentration (HHI), tail spend", "No"],
            ["Department", "Spend by department and the category × department matrix", "No"],
            ["Date", "Spend trend over time with year-over-year change", "No"],
            ["Any flag column", "A 'Spend by …' breakdown (on/off contract, diversity, status)", "No"],
        ],
        col_widths_cm=[4.5, 9.5, 2.5],
    )
    add_body(
        doc,
        "Without an amount column, the report gracefully falls back to transaction counts instead "
        "of dollars — every analysis still runs.",
    )

    # ------------------------------------------------------------ 2. How to run
    add_h1(doc, "2.  How to run it — step by step")
    add_bullet(doc, "Open the Spend Report page in the web app.", bold_lead="1")
    add_bullet(doc, "Upload a CSV or Excel spend file.", bold_lead="2")
    add_bullet(
        doc,
        "Confirm the detected columns. The tool guesses each column's role; in most cases you can "
        "just click Generate. Only change a box if a guess is wrong. Description is required.",
        bold_lead="3",
    )
    add_bullet(
        doc,
        "Choose the analyses you need. A short list is ticked to start with — the savings "
        "opportunity, spend by category, Pareto 80/20 and top vendors. Add any of the deeper "
        "cuts you want; anything your file cannot support is not offered. The Executive "
        "Summary and Sources & Methodology are always included.",
        bold_lead="4",
    )
    add_bullet(
        doc,
        "Click Generate spend report. (Large files show a row cap you can adjust for a fast "
        "first look.)",
        bold_lead="5",
    )
    add_bullet(
        doc,
        "Read the report on screen, then click Download full spend report (Excel) for the "
        "formatted workbook — it contains exactly the analyses you selected, each with its "
        "chart. Change the selection at any time and both update instantly.",
        bold_lead="6",
    )
    add_callout_box(
        doc,
        label="IT ADAPTS TO THE DATA, NOT THE OTHER WAY AROUND",
        body_text=(
            "Drop in any procurement spend file — the tool profiles what's there and offers "
            "whatever the data supports. It never invents data it doesn't have; if a column "
            "isn't present, the dependent analysis isn't offered, with a note saying why. Your "
            "file decides what is available; you decide what is produced. The report also stays "
            "put while you move between pages in the app — it is saved for the session until "
            "you upload a new file or clear it."
        ),
    )

    # ------------------------------------------------------------ 3. What it produces
    add_h1(doc, "3.  The reports it produces — and the decision each supports")
    add_body(
        doc,
        "Your file decides which of these are available; you decide which are produced. Pick "
        "only what you will use — a shorter report gets read. Every analysis you select brings "
        "its chart with it, on screen and on its Excel tab.",
    )
    _table(
        doc,
        ["Report", "What it shows", "Decision it supports"],
        [
            ["Executive Summary", "Headline findings in plain language + recommended first steps", "Read this first"],
            ["Spend by Category", "Dollars & share by the 17 Business Categories, with Examples", "Where money goes, by what is bought"],
            ["Pareto 80/20", "The few categories that drive ~80% of spend", "Where to focus sourcing first"],
            ["Top Vendors", "Largest vendors by spend and share", "Vendor rationalization targets"],
            ["Consolidation / Fragmentation", "Categories bought from many vendors (and departments)", "Where to consolidate demand — savings"],
            ["Same Item — Multiple Vendors / Depts", "The same item bought from >1 vendor and/or across >1 department, with names", "Line-level maverick buying — sharpest targets"],
            ["Same Commodity (NIGP code)", "Groups by NIGP commodity code so differently-worded descriptions of the same commodity roll together", "Catches split buying hidden behind wording differences"],
            ["Top Consolidation Opportunities", "Each fragmented commodity with the vendors and departments named, the wordings that rolled together, spend, estimated savings and a recommended action", "Who to call, and what to put on one contract"],
            ["Spend Trend", "Spend by year with year-over-year change", "Is spend rising or falling; when it peaked"],
            ["Vendor Concentration", "Top-1/5/10 vendor share and an HHI score", "How dependent are we on a few suppliers"],
            ["Tail-Spend", "The many small vendors making up the last ~20%", "The classic consolidation target"],
            ["Single- vs Multi-Source", "Categories bought from one vendor vs. many", "Vendor rationalization & supply risk"],
            ["Spend by Department", "Which departments spend, and on what", "Cross-department demand"],
            ["Category × Dept Matrix", "Same commodity bought across many departments", "Cross-department leverage"],
            ["Spend by Dimension", "Breakdown by any flag (contract, diversity, status)", "Compliance & policy questions"],
        ],
        col_widths_cm=[4.2, 7.3, 5.0],
    )
    add_callout_box(
        doc,
        label="THE OPPORTUNITIES TAB NAMES NAMES",
        body_text=(
            "Top Consolidation Opportunities does not stop at \u201cbought from 12 vendors.\u201d It "
            "lists the vendors and the departments by name, ordered by spend, so the suppliers "
            "worth a conversation come first. It also shows every description that rolled "
            "together under that row \u2014 which lets you check the grouping rather than take it "
            "on trust. If a row was matched on a 3-digit NIGP class, that column is where you "
            "would see two things that are not really the same commodity sitting together, and "
            "judge for yourself. Headers carry sort/filter dropdowns and the item name stays "
            "pinned as you scroll right. Long vendor lists are capped, with a \u201c+N more\u201d tail."
        ),
    )
    add_callout_box(
        doc,
        label="THE NUMBERS DO NOT DEPEND ON WHAT YOU TICK",
        body_text=(
            "Every analysis is calculated on every run, whatever you select. This is deliberate: "
            "the savings figures are derived from the consolidation analysis, so if unticking a "
            "box changed the maths, the headline would move depending on who was reading it — "
            "and nobody could defend that in a budget hearing. Your selection changes what is "
            "shown and exported, never how anything is computed. It is also why changing your "
            "mind after the report has run costs nothing: nothing is re-classified or "
            "recalculated, so the report and the workbook update on the spot."
        ),
    )

    # ------------------------------------------------------------ 3b. How much we could save
    add_h1(doc, "4.  How much we could save — in plain words")
    add_callout_box(
        doc,
        label="THE SIMPLE IDEA",
        body_text=(
            "The report finds where the city buys the same thing from more than one vendor or "
            "department, and estimates what buying it together could save. Estimated savings = the "
            "spend we can combine × a savings rate. The rate grows the more spread-out the buying "
            "is: +3% for each extra vendor + 2% for each extra department, capped at 15%; combining "
            "many tiny vendors adds a 5% lever. Every rate is adjustable in the app."
        ),
    )
    add_bullet(
        doc,
        "Cash savings = money that comes back to the budget. Avoided costs = money we keep from "
        "spending later. Each opportunity is split into both — default 40% cash / 60% avoided — "
        "because finance treats them differently. (The procurement terms are “hard savings” and "
        "“cost avoidance.”)",
        bold_lead="Cash savings vs avoided costs",
    )
    add_bullet(
        doc,
        "“How much we could save” leads the report: the spend we can combine, the estimated savings "
        "(cash + avoided), the amount per year, and the 3-year total. “What to combine” lists each "
        "item bought from many vendors/departments with its estimated savings and what to do.",
        bold_lead="Where it appears",
    )
    add_body(
        doc,
        "These are estimates to show where to look first, not guaranteed savings. The spend we can "
        "combine is measured exactly from your data; the percentage is an adjustable estimate. Files "
        "carry totals, not unit prices, so this is not an exact price comparison. The plain-language "
        "definitions and real sources are on the Sources & Methodology tab.",
    )

    # ------------------------------------------------------------ 4b. Every row in a real category
    add_h1(doc, "5.  Every row lands in a real category")
    add_body(
        doc,
        "There is no \"Unclassified\" or catch-all line in the report. Rows whose description "
        "doesn't match a specific-commodity rule are assigned to their closest of the 17 Business "
        "Categories (best-fit), so every row lands in a real category and the report total "
        "reconciles to your file exactly. The Spend by Category table carries an Examples column "
        "showing representative descriptions per category.",
    )
    add_body(
        doc,
        "One deliberate exception keeps the report clean on any file: the specialized \"Grants & "
        "Pass-Through Funding\" category only appears when a line carries a genuine grant signal "
        "(a program-tag prefix or an explicit grant keyword). It will not fire on an ordinary "
        "commodity purchase in an uploaded file.",
    )

    # ------------------------------------------------------------ 5. How numbers are calculated
    add_h1(doc, "6.  How the numbers are calculated — no black box")
    add_bullet(
        doc,
        "Spend by Category = sum(amount) and row count grouped by Business Category, with each "
        "category's % of total and an Examples column. Every row is included, so the total always "
        "reconciles to your file.",
        bold_lead="Spend by Category",
    )
    add_bullet(
        doc,
        "Pareto 80/20 = the specific-commodity category totals sorted high-to-low with a running "
        "cumulative %; the report names how many categories reach 80%.",
        bold_lead="Pareto",
    )
    add_bullet(
        doc,
        "Top Vendors = sum(amount) grouped by vendor, sorted, with % of total.",
        bold_lead="Top Vendors",
    )
    add_bullet(
        doc,
        "Consolidation / Fragmentation = for each category, the distinct vendor count (and "
        "department count); categories bought from more than one vendor are flagged and ranked by "
        "spend — the biggest consolidation opportunities first.",
        bold_lead="Consolidation",
    )
    add_bullet(
        doc,
        "Same Item — Multiple Vendors / Departments = group rows by the exact item description, "
        "count distinct vendors and departments per item and sum spend, then keep only items "
        "bought from more than one vendor or across more than one department. Vendor and "
        "department names are shown. Grouping is exact text — no fuzzy matching — so a flag is "
        "real, not inferred.",
        bold_lead="Same Item",
    )
    add_bullet(
        doc,
        "Same Commodity (NIGP code) = the same idea as Same Item, but grouped by the NIGP "
        "commodity code (5-digit item where the rule assigned one, otherwise 3-digit class) "
        "rather than the exact wording — so two differently-worded descriptions of the same "
        "commodity roll together. It is laid out as a matrix: each department is its own column "
        "(spend), shaded cells show which departments bought the commodity, and any commodity "
        "bought across two or more departments is highlighted as a consolidation target. The "
        "top consolidatable items also appear in the Executive Summary.",
        bold_lead="Same Commodity",
    )
    add_bullet(
        doc,
        "Vendor Concentration = the spend share of the top 1/5/10 vendors, plus an HHI (sum of "
        "squared vendor shares × 10,000; higher = more concentrated).",
        bold_lead="Concentration",
    )
    add_bullet(
        doc,
        "Tail Spend = the vendors outside the top group that makes up 80% — counted and summed to "
        "show how little of total spend the long tail represents.",
        bold_lead="Tail Spend",
    )
    add_bullet(
        doc,
        "Spend Trend = sum(amount) grouped by the year of the date column, with year-over-year % "
        "change.",
        bold_lead="Trend",
    )
    add_bullet(
        doc,
        "Coverage = how many rows mapped to a specific commodity category vs. fell into General & "
        "Other. This is the report's honesty check — and the Examples column shows what's in the "
        "catch-all.",
        bold_lead="Coverage",
    )
    add_body(
        doc,
        "Currency text is cleaned automatically ($, commas, and parenthesized negatives are "
        "handled). Amounts that can't be parsed are treated as zero, never guessed.",
    )

    # ------------------------------------------------------------ 6. Why trust it
    add_h1(doc, "7.  Why you can trust it")
    add_bullet(doc, "Deterministic — every figure is plain arithmetic you can reproduce by hand.", bold_lead="Deterministic")
    add_bullet(doc, "No AI at runtime — classification uses the same frozen rule library as the batch engine.", bold_lead="Rules-only")
    add_bullet(doc, "Private — your file is processed in the session and is not stored or sent anywhere.", bold_lead="Private")
    add_bullet(doc, "Reconciles — every row is counted, so the report total equals your file's total.", bold_lead="Reconciles")
    add_bullet(doc, "Auditable — the Examples column and the downloadable rule library show why any line landed where it did.", bold_lead="Auditable")

    # ------------------------------------------------------------ 8. Sources & methodology
    add_h1(doc, "8.  Sources & methodology — who uses this math")
    add_body(
        doc,
        "The report is built on standard spend-analysis and strategic-sourcing methods used across "
        "both public and private procurement. The addressable spend is computed from your data; "
        "the savings percentages are transparent, adjustable planning assumptions calibrated to "
        "conservative, commonly-cited ranges — not figures attributed to any single proprietary "
        "report.",
    )
    try:
        import spend_report_JHK3 as _sr
        for heading, lines in _sr.SOURCES_SECTIONS:
            add_h2(doc, heading)
            for ln in lines:
                add_bullet(doc, ln)
        add_h2(doc, "Reference organizations & sources")
        for name, url, note in _sr.SOURCES_REFERENCES:
            add_bullet(doc, f"{name} — {note}  ({url})")
    except Exception:  # noqa: BLE001
        add_body(doc, "See the Sources & Methodology tab in the Excel report for full references.")

    # ------------------------------------------------------------ 9. Limits
    add_h1(doc, "9.  Honest limits")
    add_body(
        doc,
        "The rules were tuned on one agency's vocabulary, so on a very different dataset more rows "
        "start out best-fit rather than directly rule-matched. That is expected and improves as "
        "curated rules are added. Savings are ADDRESSABLE-spend, benchmark-rate ESTIMATES (cost "
        "avoidance plus a cashable hard portion), not price-variance math — your files carry totals, "
        "not unit price × quantity. The tool classifies from description text only — not vendor "
        "names or prior codes — which keeps the classification independent and defensible.",
    )

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
