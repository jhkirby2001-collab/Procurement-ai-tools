import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from pathlib import Path

print("="*80)
print("CREATING MICROSOFT WORD AND EXCEL FILES")
print("="*80)

# Paths
data_file = "/workspaces/Spend_Analysis_and_Categorization/spend-analysis/data/raw/BottledWater raw data.xlsx"
output_dir = Path("/workspaces/Spend_Analysis_and_Categorization/spend-analysis/outputs")

# Read the Excel file
print("\n📥 Reading source data...")
excel_file = pd.ExcelFile(data_file)
df1 = pd.read_excel(data_file, sheet_name='SPC49756A-PO14954')
df2 = pd.read_excel(data_file, sheet_name='SPC82197E-PO30043')
df3 = pd.read_excel(data_file, sheet_name='EXHIBIT B')

print(f"   Sheet 1: {len(df1)} rows")
print(f"   Sheet 2: {len(df2)} rows")
print(f"   Sheet 3: {len(df3)} rows")

# ============================================================================
# CREATE EXECUTIVE SUMMARY - WORD DOCUMENT
# ============================================================================
print("\n📄 Creating Executive Summary (Word)...")

doc = Document()

# Title
title = doc.add_heading('EXECUTIVE SUMMARY', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Bottled Water Spend Analysis')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(16)
subtitle_format.font.bold = True

doc.add_paragraph()

# Date and basics
p = doc.add_paragraph()
p.add_run('Report Date: ').bold = True
p.add_run('January 20, 2026')
doc.add_paragraph()

# Overview section
doc.add_heading('Overview', 1)
overview_text = """This analysis examines bottled water procurement across three data sources spanning 2007-2026. The analysis covers two consecutive bottled water contracts and general procurement requisitions totaling $278.97 Million across 4,345 records."""
doc.add_paragraph(overview_text)

doc.add_page_break()

# Key Findings
doc.add_heading('Key Findings', 1)

doc.add_heading('Contract Performance', 2)
findings_table = doc.add_table(rows=4, cols=3)
findings_table.style = 'Light Grid Accent 1'

# Header row
header_cells = findings_table.rows[0].cells
header_cells[0].text = 'Metric'
header_cells[1].text = 'Contract 1 (2007-2013)'
header_cells[2].text = 'Contract 2 (2014-2022)'

# Data rows
row1 = findings_table.rows[1].cells
row1[0].text = 'Vendor'
row1[1].text = 'Ice Mountain'
row1[2].text = 'Bluetriton Brands Inc'

row2 = findings_table.rows[2].cells
row2[0].text = 'Contract Limit'
row2[1].text = '$211,523.70'
row2[2].text = '$925,006.00'

row3 = findings_table.rows[3].cells
row3[0].text = 'Utilization Rate'
row3[1].text = '92.30% (A grade)'
row3[2].text = '74.56% (C grade)'

doc.add_paragraph()

# Critical Issues
doc.add_heading('Critical Issues Identified', 1)

doc.add_heading('1. Single Vendor Dependency (15 Years)', 2)
doc.add_paragraph('• 100% dependency on Ice Mountain/Bluetriton for 15 consecutive years')
doc.add_paragraph('• No competitive bidding - estimated 15-25% price premium')
doc.add_paragraph('• Estimated overpayment: $132,736 over 15 years')
doc.add_paragraph('• High supply chain risk with single source')

doc.add_heading('2. Contract Expiration Crisis (Overall Procurement)', 2)
doc.add_paragraph('• 67.1% of all requisitions ($182.2M) cite expired contracts')
doc.add_paragraph('• Systematic contract renewal failures')
doc.add_paragraph('• 2,913 requisitions affected by expired contracts')
doc.add_paragraph('• Additional 15% ($40.8M) spent on post-expiration invoice reconciliation')

doc.add_heading('3. Contract 2 Poor Utilization', 2)
doc.add_paragraph('• Only 74.56% utilization vs. 92.30% in Contract 1')
doc.add_paragraph('• $235,346 wasted capacity (25.4% of contract)')
doc.add_paragraph('• Contract limit increased 337% but spend only increased 253%')
doc.add_paragraph('• Poor demand forecasting')

doc.add_page_break()

# Financial Summary
doc.add_heading('Financial Summary', 1)

doc.add_heading('Bottled Water Contracts', 2)
bw_table = doc.add_table(rows=6, cols=2)
bw_table.style = 'Light List Accent 1'

bw_data = [
    ('Total Contract Value', '$1,136,529.70'),
    ('Total Actual Spend', '$884,906.53'),
    ('Contract 1 Spend', '$195,246.71'),
    ('Contract 2 Spend', '$689,659.82'),
    ('Wasted Capacity (Contract 2)', '$235,346.18'),
    ('Estimated Overpayment (15 yrs)', '$132,736.00')
]

for i, (label, value) in enumerate(bw_data):
    row = bw_table.rows[i].cells
    row[0].text = label
    row[1].text = value

doc.add_paragraph()

doc.add_heading('Overall Procurement (All Sources)', 2)
overall_table = doc.add_table(rows=5, cols=2)
overall_table.style = 'Light List Accent 1'

overall_data = [
    ('Total Analyzed Spend', '$278,969,275.62'),
    ('Total Records', '4,345'),
    ('Departments', '32'),
    ('Vendors', '511'),
    ('Expired Contract Spend', '$182,210,400 (67.1%)')
]

for i, (label, value) in enumerate(overall_data):
    row = overall_table.rows[i].cells
    row[0].text = label
    row[1].text = value

doc.add_page_break()

# Cost Savings Opportunities
doc.add_heading('Cost Savings Opportunities', 1)

savings_table = doc.add_table(rows=6, cols=3)
savings_table.style = 'Light Grid Accent 1'

# Header
header = savings_table.rows[0].cells
header[0].text = 'Initiative'
header[1].text = 'Estimated Savings'
header[2].text = 'Timeline'

# Data
savings_data = [
    ('Competitive Bidding (Bottled Water)', '$90,000 - $150,000', '3 years'),
    ('Product Mix Optimization', '$150,000 - $250,000/year', 'Year 2-3'),
    ('Right-Size Contracting', '$20,000 - $40,000/year', 'Year 1'),
    ('Contract Lifecycle Management', '$4M - $6M/year', 'Year 1-2'),
    ('TOTAL (3-Year Cumulative)', '$6M - $10M', '3 years')
]

for i, (initiative, savings, timeline) in enumerate(savings_data, 1):
    row = savings_table.rows[i].cells
    row[0].text = initiative
    row[1].text = savings
    row[2].text = timeline

doc.add_paragraph()

# Top Recommendations
doc.add_heading('Top 3 Recommendations', 1)

doc.add_heading('1. Implement Contract Lifecycle Management System', 2)
doc.add_paragraph('PRIORITY: CRITICAL | TIMELINE: Immediate (30-60 days)')
doc.add_paragraph('• Conduct complete contract inventory')
doc.add_paragraph('• Implement 180-day advance expiration alerts')
doc.add_paragraph('• Deploy contract management software')
doc.add_paragraph('• Expected impact: Reduce expired contract requisitions from 67% to <5%')

doc.add_heading('2. Launch Competitive Bidding for Bottled Water', 2)
doc.add_paragraph('PRIORITY: HIGH | TIMELINE: 60-90 days')
doc.add_paragraph('• Issue RFP to 5+ qualified vendors')
doc.add_paragraph('• Award primary (70%) and secondary (30%) vendors')
doc.add_paragraph('• Expected savings: 15-25% ($90K-$150K over 3 years)')

doc.add_heading('3. Optimize Product Mix & Sustainability', 2)
doc.add_paragraph('PRIORITY: HIGH | TIMELINE: 6-12 months')
doc.add_paragraph('• Install water filtration systems in high-volume locations')
doc.add_paragraph('• Reduce individual bottle usage by 50%')
doc.add_paragraph('• Expected savings: $150K-$250K/year + environmental benefits')

doc.add_page_break()

# Appendix
doc.add_heading('Supporting Documents', 1)
doc.add_paragraph('For detailed analysis, please refer to:')
doc.add_paragraph('• BW_RAW_DATA_CONTRACTS_ANALYSIS.md - Full contract analysis')
doc.add_paragraph('• BW_RAW_DATA_COMPREHENSIVE_ANALYSIS.md - All data analysis')
doc.add_paragraph('• BW_RAW_DATA_PARETO_VISUAL_REPORT.md - Pareto 80/20 analysis')
doc.add_paragraph('• bw_raw_data_charts/ - Bottled water visualizations')
doc.add_paragraph('• bw_raw_data_pareto_charts/ - Pareto analysis charts')

doc.add_paragraph()
doc.add_paragraph('Analysis prepared by: Claude Code AI Assistant')
doc.add_paragraph('Date: January 20, 2026')

# Save Word document
word_file = output_dir / "BW_RAW_DATA_EXECUTIVE_SUMMARY.docx"
doc.save(word_file)
print(f"   ✓ Saved: {word_file.name}")

# ============================================================================
# CREATE EXCEL ANALYSIS FILE
# ============================================================================
print("\n📊 Creating Excel Analysis File...")

# Create Excel workbook
wb = openpyxl.Workbook()
wb.remove(wb.active)  # Remove default sheet

# Header styles
header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
header_font = Font(bold=True, color="FFFFFF", size=12)
header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

# Title style
title_font = Font(bold=True, size=14, color="FFFFFF")
title_fill = PatternFill(start_color="203864", end_color="203864", fill_type="solid")

# ============================================================================
# Sheet 1: Executive Summary
# ============================================================================
ws1 = wb.create_sheet("Executive Summary")

# Title
ws1.merge_cells('A1:E1')
ws1['A1'] = 'BOTTLED WATER SPEND ANALYSIS - EXECUTIVE SUMMARY'
ws1['A1'].font = title_font
ws1['A1'].fill = title_fill
ws1['A1'].alignment = Alignment(horizontal='center', vertical='center')
ws1.row_dimensions[1].height = 30

ws1['A2'] = 'Report Date:'
ws1['B2'] = 'January 20, 2026'
ws1['A3'] = 'Analysis Period:'
ws1['B3'] = '2007-2026'

# Key Metrics
ws1['A5'] = 'KEY METRICS'
ws1['A5'].font = Font(bold=True, size=12)
ws1.merge_cells('A5:E5')

metrics_data = [
    ['Metric', 'Value', '', 'Metric', 'Value'],
    ['Total Contract Value', '$1,136,529.70', '', 'Total Records', '4,345'],
    ['Total Actual Spend', '$884,906.53', '', 'Departments', '32'],
    ['Contract 1 Utilization', '92.30%', '', 'Vendors', '511'],
    ['Contract 2 Utilization', '74.56%', '', 'Approval Rate', '95.3%'],
    ['Wasted Capacity', '$235,346.18', '', 'Expired Contracts', '67.1% of spend']
]

row = 6
for data_row in metrics_data:
    for col, value in enumerate(data_row, 1):
        cell = ws1.cell(row=row, column=col, value=value)
        if row == 6:  # Header row
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
    row += 1

# Critical Issues
ws1['A13'] = 'CRITICAL ISSUES'
ws1['A13'].font = Font(bold=True, size=12)
ws1.merge_cells('A13:E13')

issues_data = [
    ['Issue', 'Impact', 'Priority'],
    ['Single Vendor (15 years)', 'Estimated $132K overpayment', 'CRITICAL'],
    ['Expired Contracts (67%)', '$182.2M affected', 'CRITICAL'],
    ['Contract 2 Poor Utilization', '$235K wasted capacity', 'HIGH'],
    ['Product Mix Inefficiency', 'High cost per gallon', 'MEDIUM']
]

row = 14
for data_row in issues_data:
    for col, value in enumerate(data_row, 1):
        cell = ws1.cell(row=row, column=col, value=value)
        if row == 14:  # Header row
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
    row += 1

# Savings Opportunities
ws1['A20'] = 'SAVINGS OPPORTUNITIES'
ws1['A20'].font = Font(bold=True, size=12)
ws1.merge_cells('A20:E20')

savings_data = [
    ['Initiative', 'Estimated Savings', 'Timeline', 'ROI'],
    ['Competitive Bidding', '$90,000 - $150,000', '3 years', 'Very High'],
    ['Product Mix Optimization', '$150,000 - $250,000/year', 'Year 2-3', 'Very High'],
    ['Right-Size Contracting', '$20,000 - $40,000/year', 'Year 1', 'Medium'],
    ['Contract Lifecycle Mgmt', '$4M - $6M/year', 'Year 1-2', 'Very High'],
    ['TOTAL', '$6M - $10M', '3 years', 'Excellent']
]

row = 21
for data_row in savings_data:
    for col, value in enumerate(data_row, 1):
        cell = ws1.cell(row=row, column=col, value=value)
        if row == 21:  # Header row
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        elif row == 26:  # Total row
            cell.font = Font(bold=True)
    row += 1

# Adjust column widths
ws1.column_dimensions['A'].width = 30
ws1.column_dimensions['B'].width = 25
ws1.column_dimensions['C'].width = 15
ws1.column_dimensions['D'].width = 25
ws1.column_dimensions['E'].width = 20

# ============================================================================
# Sheet 2: Contract 1 Analysis
# ============================================================================
ws2 = wb.create_sheet("Contract 1 (2007-2013)")

# Title
ws2.merge_cells('A1:F1')
ws2['A1'] = 'CONTRACT 1: SPC49756A-PO14954 (Ice Mountain, 2007-2013)'
ws2['A1'].font = title_font
ws2['A1'].fill = title_fill
ws2['A1'].alignment = Alignment(horizontal='center', vertical='center')
ws2.row_dimensions[1].height = 25

# Contract Info
ws2['A3'] = 'Contract Number:'
ws2['B3'] = '49756A / PO14954'
ws2['A4'] = 'Vendor:'
ws2['B4'] = 'Ice Mountain'
ws2['A5'] = 'Period:'
ws2['B5'] = 'July 1, 2007 - December 28, 2013 (6.5 years)'
ws2['A6'] = 'Contract Limit:'
ws2['B6'] = '$211,523.70'
ws2['A7'] = 'Actual Spend:'
ws2['B7'] = '$195,246.71'
ws2['A8'] = 'Utilization:'
ws2['B8'] = '92.30%'
ws2['A9'] = 'Grade:'
ws2['B9'] = 'A (Excellent)'

# Product breakdown
ws2['A11'] = 'PRODUCT MIX'
ws2['A11'].font = Font(bold=True, size=12)
ws2.merge_cells('A11:E11')

product_header = ['Product', 'Spend', '% of Total', 'Orders', 'Avg Order']
row = 12
for col, header in enumerate(product_header, 1):
    cell = ws2.cell(row=row, column=col, value=header)
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = header_alignment

product_data = [
    ['Drinking Water 5-Gal', '$1,521.95', '0.71%', '1', '$1,521.95'],
    ['Distilled Water 1-Gal', '$394.47', '0.18%', '1', '$394.47'],
    ['Distilled Water 5-Gal', '$324.65', '0.15%', '1', '$324.65'],
    ['16.9oz Bottles (24/case)', '$7.50', '0.00%', '1', '$7.50'],
    ['Cooler Rental (No Fridge)', '$68.42', '0.03%', '1', '$68.42'],
    ['Cooler Rental (w/Fridge)', '$9.96', '0.00%', '1', '$9.96']
]

row = 13
for data_row in product_data:
    for col, value in enumerate(data_row, 1):
        ws2.cell(row=row, column=col, value=value)
    row += 1

ws2.column_dimensions['A'].width = 30
ws2.column_dimensions['B'].width = 15
ws2.column_dimensions['C'].width = 12
ws2.column_dimensions['D'].width = 10
ws2.column_dimensions['E'].width = 15

# ============================================================================
# Sheet 3: Contract 2 Analysis
# ============================================================================
ws3 = wb.create_sheet("Contract 2 (2014-2022)")

# Title
ws3.merge_cells('A1:F1')
ws3['A1'] = 'CONTRACT 2: SPC82197E-PO30043 (Bluetriton Brands, 2014-2022)'
ws3['A1'].font = title_font
ws3['A1'].fill = title_fill
ws3['A1'].alignment = Alignment(horizontal='center', vertical='center')
ws3.row_dimensions[1].height = 25

# Contract Info
ws3['A3'] = 'Contract Number:'
ws3['B3'] = '82197E / PO30043'
ws3['A4'] = 'Vendor:'
ws3['B4'] = 'Bluetriton Brands Inc'
ws3['A5'] = 'Period:'
ws3['B5'] = 'June 3, 2014 - December 1, 2022 (8.5 years)'
ws3['A6'] = 'Contract Limit:'
ws3['B6'] = '$925,006.00'
ws3['A7'] = 'Actual Spend:'
ws3['B7'] = '$689,659.82'
ws3['A8'] = 'Utilization:'
ws3['B8'] = '74.56%'
ws3['A9'] = 'Grade:'
ws3['B9'] = 'C (Below Target)'
ws3['A10'] = 'Wasted Capacity:'
ws3['B10'] = '$235,346.18 (25.4%)'

# Product breakdown
ws3['A12'] = 'PRODUCT MIX'
ws3['A12'].font = Font(bold=True, size=12)
ws3.merge_cells('A12:E12')

row = 13
for col, header in enumerate(product_header, 1):
    cell = ws3.cell(row=row, column=col, value=header)
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = header_alignment

product_data2 = [
    ['Drinking Water 5-Gal', '$2,087.53', '0.53%', '1', '$2,087.53'],
    ['Distilled Water 5-Gal', '$2,020.98', '0.52%', '1', '$2,020.98'],
    ['Distilled Water 1-Gal', '$1,525.71', '0.39%', '1', '$1,525.71'],
    ['16.9oz Bottles (24/case)', '$199.60', '0.05%', '1', '$199.60'],
    ['Cooler Rental (No Fridge)', '$2.99', '0.00%', '1', '$2.99']
]

row = 14
for data_row in product_data2:
    for col, value in enumerate(data_row, 1):
        ws3.cell(row=row, column=col, value=value)
    row += 1

ws3.column_dimensions['A'].width = 30
ws3.column_dimensions['B'].width = 15
ws3.column_dimensions['C'].width = 12
ws3.column_dimensions['D'].width = 10
ws3.column_dimensions['E'].width = 15

# ============================================================================
# Sheet 4: Department Spend (from EXHIBIT B)
# ============================================================================
ws4 = wb.create_sheet("Department Spend")

# Title
ws4.merge_cells('A1:E1')
ws4['A1'] = 'DEPARTMENT SPEND ANALYSIS (EXHIBIT B)'
ws4['A1'].font = title_font
ws4['A1'].fill = title_fill
ws4['A1'].alignment = Alignment(horizontal='center', vertical='center')
ws4.row_dimensions[1].height = 25

# Header
dept_header = ['Department', 'Total Spend', '% of Total', 'Count', 'Avg Requisition']
row = 3
for col, header in enumerate(dept_header, 1):
    cell = ws4.cell(row=row, column=col, value=header)
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = header_alignment

# Calculate department spend
dept_spend = df3.groupby('Department Description')['Requisition Amount'].agg(['sum', 'count', 'mean']).sort_values('sum', ascending=False).head(15)
total_spend = df3['Requisition Amount'].sum()

row = 4
for dept, values in dept_spend.iterrows():
    ws4.cell(row=row, column=1, value=dept)
    ws4.cell(row=row, column=2, value=f"${values['sum']:,.2f}")
    ws4.cell(row=row, column=3, value=f"{values['sum']/total_spend*100:.2f}%")
    ws4.cell(row=row, column=4, value=int(values['count']))
    ws4.cell(row=row, column=5, value=f"${values['mean']:,.2f}")
    row += 1

ws4.column_dimensions['A'].width = 50
ws4.column_dimensions['B'].width = 18
ws4.column_dimensions['C'].width = 12
ws4.column_dimensions['D'].width = 10
ws4.column_dimensions['E'].width = 18

# ============================================================================
# Sheet 5: Recommendations
# ============================================================================
ws5 = wb.create_sheet("Recommendations")

# Title
ws5.merge_cells('A1:E1')
ws5['A1'] = 'TOP RECOMMENDATIONS'
ws5['A1'].font = title_font
ws5['A1'].fill = title_fill
ws5['A1'].alignment = Alignment(horizontal='center', vertical='center')
ws5.row_dimensions[1].height = 25

# Recommendations
rec_data = [
    ['Priority', 'Recommendation', 'Timeline', 'Expected Impact', 'Cost'],
    ['CRITICAL', 'Implement Contract Lifecycle Management System', '30-60 days', 'Reduce expired contracts from 67% to <5%', '$50K-$100K'],
    ['HIGH', 'Launch Competitive Bidding for Bottled Water', '60-90 days', 'Save $90K-$150K over 3 years', '$10K-$20K'],
    ['HIGH', 'Optimize Product Mix & Install Filtration', '6-12 months', 'Save $150K-$250K/year', '$50K-$100K'],
    ['MEDIUM', 'Right-Size Contract Values', '90 days', 'Achieve 85-95% utilization', 'Minimal'],
    ['MEDIUM', 'Establish Cost Benchmarking Program', '30-60 days', 'Market-validated pricing', '$5K/year'],
    ['MEDIUM', 'Implement Demand Forecasting Model', '60-90 days', 'Improve forecast accuracy to +/-10%', '$20K-$30K']
]

row = 3
for data_row in rec_data:
    for col, value in enumerate(data_row, 1):
        cell = ws5.cell(row=row, column=col, value=value)
        if row == 3:  # Header
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        ws5.row_dimensions[row].height = 30
    row += 1

ws5.column_dimensions['A'].width = 12
ws5.column_dimensions['B'].width = 45
ws5.column_dimensions['C'].width = 15
ws5.column_dimensions['D'].width = 35
ws5.column_dimensions['E'].width = 15

# Save Excel file
excel_file = output_dir / "BW_RAW_DATA_ANALYSIS.xlsx"
wb.save(excel_file)
print(f"   ✓ Saved: {excel_file.name}")

print("\n" + "="*80)
print("✅ COMPLETE!")
print("="*80)
print(f"\nCreated files in: {output_dir}")
print("\n📄 Microsoft Word:")
print("   • BW_RAW_DATA_EXECUTIVE_SUMMARY.docx")
print("\n📊 Microsoft Excel:")
print("   • BW_RAW_DATA_ANALYSIS.xlsx")
print("     - Executive Summary sheet")
print("     - Contract 1 (2007-2013) sheet")
print("     - Contract 2 (2014-2022) sheet")
print("     - Department Spend sheet")
print("     - Recommendations sheet")
print("="*80)
